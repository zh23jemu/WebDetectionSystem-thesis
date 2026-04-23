from pathlib import Path

from docx import Document


# 说明：
# 1. 本脚本只修正论文中“代码展示段落”的换行与缩进；
# 2. 不改动正文说明、图题、表题、章节标题等其他内容；
# 3. 使用段落索引精确替换，尽量将影响范围控制在代码段本身。


DOC_PATH = Path(
    r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的木质板材缺陷检测系统的设计与实现.docx"
)


CODE_REPLACEMENTS = {
    180: """# app.py
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        login_role = request.form.get('login_role')
        user = User.query.filter_by(username=username).first()

        if user and check_password_hash(user.password, password):
            if user.role != login_role:
                flash('身份不匹配，请切换到对应标签页登录。', 'warning')
                return redirect(url_for('login'))
            login_user(user)
            return redirect(url_for('dashboard'))
    return render_template('login.html')""",
    187: """# app.py
@app.route('/history', methods=['GET', 'POST'])
@login_required
def history():
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '')
    query = DetectionHistory.query

    if current_user.role != 'admin':
        query = query.filter_by(user_id=current_user.id)
    if search:
        query = query.filter(DetectionHistory.defect_type.contains(search))

    records = query.order_by(
        DetectionHistory.upload_date.desc()
    ).paginate(page=page, per_page=12)""",
    193: """部分关键代码如下：

# app.py
@app.route('/logs')
@login_required
def logs():
    today_count = hist_query.filter(
        func.date(DetectionHistory.upload_date) == date.today()
    ).count()
    stats_data = [[k, v] for k, v in counter.items()]
    trend_data = []
    for i in range(6, -1, -1):
        target_date = date.today() - timedelta(days=i)
        count = hist_query.filter(
            func.date(DetectionHistory.upload_date) == target_date
        ).count()
        trend_data.append(count)""",
    199: """# app.py
total_seconds = 0
for l in log_query.filter(WorkLog.logout_time != None).all():
    total_seconds += (l.logout_time - l.login_time).total_seconds()

if session.get('current_log_id'):
    curr = db.session.get(WorkLog, session['current_log_id'])
    if curr:
        total_seconds += (datetime.now() - curr.login_time).total_seconds()

work_time_str = (
    f"{int(total_seconds // 3600)}小时 "
    f"{int((total_seconds % 3600) // 60)}分钟"
)""",
    205: """# app.py
@app.route('/sales', methods=['GET', 'POST'])
@login_required
def sales():
    if request.method == 'POST':
        grade = request.form.get('grade')
        qty = int(request.form.get('quantity'))
        inv = db.session.get(Inventory, grade)
        if not inv or inv.count < qty:
            flash(f'{grade} 库存不足！', 'danger')
        else:
            inv.count -= qty
            sale = SalesRecord(
                customer_name=request.form.get('customer'),
                product_grade=grade,
                quantity=qty
            )
            db.session.add(sale)""",
    211: """# app.py / templates/help.html
@app.route('/help')
@login_required
def help_center():
    return render_template('help.html')

<img src="{{ url_for('static', filename='img/live_knot.png') }}"
     class="card-img-top"
     alt="活节样张">
<img src="{{ url_for('static', filename='img/dead_knot.png') }}"
     class="card-img-top"
     alt="死节样张">
<img src="{{ url_for('static', filename='img/crack.png') }}"
     class="card-img-top"
     alt="裂纹样张">""",
    218: """# app.py
@app.route('/admin/users', methods=['GET', 'POST'])
@login_required
def admin_users():
    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'add':
            new_user = User(
                username=request.form.get('username'),
                password=generate_password_hash(request.form.get('password')),
                role=request.form.get('role')
            )
            db.session.add(new_user)
            db.session.commit()""",
    223: """部分关键代码如下：

# app.py
@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    val = request.form.get('conf_threshold')
    s = SystemSettings.query.filter_by(key_name='conf_threshold').first()
    s.value = val

    f = request.files.get('model_file')
    if f and f.filename.endswith('.pt'):
        fname = secure_filename(f.filename)
        f.save(os.path.join(app.root_path, fname))
        m = SystemSettings.query.filter_by(key_name='current_model').first()
        m.value = fname

5.3 算法检测模块的实现""",
    229: """# templates/detect.html
<ul class="nav nav-tabs mb-3" id="detectTab">
    <li class="nav-item">
        <button class="nav-link active" data-bs-target="#upload-pane"
                type="button" onclick="stopCamera()">本地上传</button>
    </li>
    <li class="nav-item">
        <button class="nav-link" data-bs-target="#camera-pane"
                type="button" onclick="initCamera()">摄像头采集</button>
    </li>
</ul>
<form id="uploadForm">
    <input class="form-control" type="file" id="images"
           name="images" multiple accept="image/*">
    <div class="row g-3 mb-3" id="previewArea"></div>
</form>

5.3.2 检测模型设计与训练""",
    235: """# utils/detection.py
def process_image(image_path, save_dir, model_path, conf_threshold=0.25):
    model = load_model(model_path)
    img = cv2.imread(image_path)
    results = model(img, conf=float(conf_threshold))

    detected_objects = []
    max_conf = 0.0
    for r in results:
        res_plotted = r.plot()
        for box in r.boxes:
            cls_id = int(box.cls[0])
            detected_objects.append(model.names[cls_id])
            max_conf = max(max_conf, float(box.conf[0]))
    return {
        'defects': ','.join(list(set(detected_objects))),
        'confidence': round(max_conf, 2)
    }""",
    241: """# app.py
detected_grade = 'A级'
if res['is_clean']:
    detected_grade = 'A级'
else:
    for d in res['defects'].split(','):
        zh_name = DEFECT_MAPPING.get(d.strip(), d.strip())
        if zh_name in ['裂纹', '死节']:
            detected_grade = 'C级'
            break
        elif zh_name == '活节':
            detected_grade = 'B级'

inv = db.session.get(Inventory, detected_grade)
if inv:
    inv.count += 1""",
    247: """# app.py
history = DetectionHistory(
    filename=filename,
    result_image=res['filepath'],
    defect_type=final_defect_str,
    confidence=res.get('confidence', 0.0),
    user_id=current_user.id,
    grade=detected_grade
)
db.session.add(history)
results.append(res)
db.session.commit()""",
}


def main() -> None:
    document = Document(DOC_PATH)

    for index, content in CODE_REPLACEMENTS.items():
        if index >= len(document.paragraphs):
            raise IndexError(f"段落索引越界：{index}")
        document.paragraphs[index].text = content

    document.save(DOC_PATH)
    print(f"已修正 {len(CODE_REPLACEMENTS)} 处代码段格式。")


if __name__ == "__main__":
    main()
