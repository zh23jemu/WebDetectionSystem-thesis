from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")


def set_para_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    # 先清空原有边框配置
    for child in list(borders):
        borders.remove(child)

    spec = {
        "top": {"val": "single", "sz": "12", "space": "0", "color": "000000"},
        "left": {"val": "nil"},
        "bottom": {"val": "single", "sz": "12", "space": "0", "color": "000000"},
        "right": {"val": "nil"},
        "insideH": {"val": "single", "sz": "8", "space": "0", "color": "000000"},
        "insideV": {"val": "nil"},
    }
    for edge, attrs in spec.items():
        elem = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            elem.set(qn(f"w:{k}"), v)
        borders.append(elem)


def format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "宋体"
                    run.font.size = Pt(10.5)
                if row_idx == 0:
                    for run in para.runs:
                        run.bold = True


def main():
    doc = Document(DOC_PATH)
    paras = doc.paragraphs

    # 改写测试段落
    replacements = {
        "第六章测试更强调“能否稳定完成完整流程”，而不是单独追求某个理论指标。": (
            "系统测试的主要目的是验证木材表面缺陷智能检测系统在当前运行环境下能否稳定完成各项核心业务流程，"
            "并检查系统在正常输入和异常输入条件下的响应情况。"
            "本章测试重点包括用户登录、缺陷检测、检测记录保存、日志报表统计、Excel 数据导出、销售出库以及后台设置等功能。"
            "通过测试可以进一步说明系统是否满足毕业设计阶段的基本使用要求。"
        ),
        "测试环境基本沿用了平时开发和调试时的配置。": (
            "本次测试所使用的环境与系统日常开发环境保持一致。"
            "测试平台为 Windows 操作系统，后端程序运行在项目本地 `.venv` 虚拟环境中，Web 框架采用 Flask，数据库采用 SQLite，"
            "检测模块调用当前系统配置的 `.pt` 权重文件。"
            "浏览器端主要用于验证登录、图片上传、摄像头采集、图表显示和销售管理等页面功能。"
        ),
        "测试过程中分别使用普通用户和管理员两种身份进入系统，": (
            "测试过程中分别以普通用户和管理员身份进入系统，按照实际业务流程对主要功能模块进行逐项验证。"
            "除正常测试外，还对未选图片直接检测、角色入口不匹配、库存不足继续提交订单等异常情况进行了检查，以观察系统是否能够给出明确提示。"
        ),
        "异常情况的检查虽然不复杂，但很有必要。": (
            "从测试结果来看，系统在正常使用场景下能够完成登录、检测、记录保存、统计报表、数据导出和库存销售联动等操作。"
            "在异常场景下，系统对角色不匹配、未选择图片、库存不足等情况也能够返回对应提示，说明系统具备基本的异常处理能力。"
        ),
        "从实际运行结果来看，论文前面提到的主要功能现在基本都已经串了起来。": (
            "综合测试结果可知，当前系统已经能够较完整地实现论文前文设计的主要功能。"
            "系统能够完成木材表面缺陷图像检测、结果等级判定、检测历史追溯、日志报表统计、库存数量更新、销售出库和管理员后台管理等业务。"
            "整体运行过程较为稳定，能够满足毕业设计演示和课程实践的基本要求。"
        ),
        "不过，系统也还有一些比较明显的改进空间。": (
            "但从测试情况看，系统仍存在一些后续可完善的内容。"
            "例如，当前历史记录删除后尚未同步清理对应图片文件，系统数据库仍以 SQLite 单机部署为主，"
            "并且测试工作主要集中在功能验证层面，对高并发访问、异常恢复和长期运行稳定性的覆盖还不够充分。"
            "这些问题可作为系统后续优化的重要方向。"
        ),
    }

    for para in paras:
        text = para.text.strip()
        for prefix, new_text in replacements.items():
            if text.startswith(prefix):
                set_para_text(para, new_text)
                break

    # 定位 6.3 并插入测试表标题与测试截图小节
    paras = doc.paragraphs
    h63 = next(p for p in paras if p.text.strip() == "6.3 测试用例")
    p_after_h63 = paras[paras.index(h63) + 1]
    if p_after_h63.text.strip() == "":
        set_para_text(p_after_h63, "表6-1 系统功能测试用例")
        p_after_h63.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_after_h63.paragraph_format.space_before = Pt(0)
        p_after_h63.paragraph_format.space_after = Pt(6)

    # 最后一张测试表统一为三线表
    if doc.tables:
        format_table(doc.tables[-1])

    # 若尚未插入测试截图占位，则在 6.3 后面追加
    existing = {p.text.strip() for p in doc.paragraphs}
    if "6.3.1 正常功能测试截图" not in existing:
        # 找到测试表后第一个段落位置
        anchor = None
        for p in doc.paragraphs:
            if p.text.strip().startswith("从测试结果来看，系统在正常使用场景下能够完成登录"):
                anchor = p
                break
        if anchor is None:
            anchor = h63

        h631 = insert_paragraph_after(anchor, "6.3.1 正常功能测试截图", "Heading 3")
        p1 = insert_paragraph_after(h631, "在正常测试过程中，系统能够完成用户登录、图片检测、结果显示以及历史记录保存等功能。正常功能测试截图如图6-1所示。")
        fig1 = insert_paragraph_after(p1, "图6-1 正常功能测试截图")
        note1 = insert_paragraph_after(fig1, "后期补充说明：此处建议放置用户成功登录、图片检测成功返回结果图或历史记录正常显示的系统截图。")

        h632 = insert_paragraph_after(note1, "6.3.2 异常功能测试截图", "Heading 3")
        p2 = insert_paragraph_after(h632, "在异常测试过程中，系统能够对未选图片直接检测、身份不匹配登录以及库存不足出库等情况给出明确提示。异常功能测试截图如图6-2所示。")
        fig2 = insert_paragraph_after(p2, "图6-2 异常功能测试截图")
        insert_paragraph_after(fig2, "后期补充说明：此处建议放置角色不匹配提示、未选图片提示或库存不足提示等界面截图。")

        for para in [p1, note1, p2]:
            para.paragraph_format.space_after = Pt(6)
        for para in [fig1, fig2]:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(6)

    doc.save(DOC_PATH)
    print("done")


if __name__ == "__main__":
    main()
