import shutil
from pathlib import Path

from docx import Document


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_AI率定点优化前备份.docx")


REWRITE = {
    69: "摘要：木材表面缺陷检测直接关系到板材分级、库存统计和后续销售处理。在很多中小型加工或教学演示场景中，这项工作仍然依赖人工观察，不仅效率受操作人员经验影响较大，而且检测记录零散，后续查询和统计也不方便。本文以本人整理和调试的 `WebDetectionSystem` 项目为基础，对一套基于 Python 的 Web 木材表面缺陷智能检测系统进行设计与实现分析。系统后端采用 Flask，检测部分调用 Ultralytics YOLO 模型，数据库使用 SQLite，前端页面由 Jinja2 模板配合 Bootstrap 5、Chart.js 和 Viewer.js 搭建。结合当前程序实现情况，系统已经能够完成用户登录、图片上传检测、摄像头拍照检测、缺陷中文映射、等级判定、检测历史保存、报表展示、Excel 导出、库存联动、销售出库、帮助中心展示以及管理员参数设置和模型切换等功能。与单纯展示模型识别结果的课程作品相比，本系统进一步把检测结果接入了等级、库存和销售流程，形成了较完整的业务闭环。论文围绕该系统的研究背景、关键技术、需求分析、系统设计、功能实现和测试情况展开说明，并结合实际代码结构对主要模块进行梳理。",
    75: "Abstract: Wood surface defect detection is closely related to timber grading, stock statistics and later sales processing. In many small-scale processing or teaching scenarios, this task is still completed manually, which means the final result depends heavily on human experience and the records are often inconvenient to trace. Based on the `WebDetectionSystem` project that has been organized and debugged in this work, this paper analyzes the design and implementation of a Python-based web wood surface defect detection system. The backend is built with Flask, the detection module uses Ultralytics YOLO, the database is SQLite, and the pages are constructed with Jinja2 templates, Bootstrap 5, Chart.js and Viewer.js. According to the current implementation, the system already supports user login, image upload detection, camera capture, defect mapping, timber grading, history storage, report visualization, Excel export, inventory linkage, sales outbound processing, help-center presentation, and administrator-side parameter and model management. Instead of showing raw model output only, the system further connects the detection result with grade judgement, inventory update and sales workflow, which makes it closer to a complete business application.",
    80: "木材表面缺陷检测并不是一个孤立的识别问题，它和板材分级、库存数量、销售报价以及后续质量追溯都有直接关系。以裂纹、死节和活节为例，这些缺陷不仅影响木材外观，还会影响材料强度和实际使用价值。如果前端检测环节出现误判，后续库存和销售数据也会跟着失真。因此，缺陷识别结果是否稳定、是否能够及时沉淀为可查询的数据，在木材加工场景中都很关键。",
    81: "近几年，目标检测算法在工业视觉中的应用越来越普遍。与传统人工检查方式相比，图像识别方案至少有两个明显优势：一是处理速度更快，二是检测结果可以直接留痕并进入数据库，便于后续统计和追溯。尤其是 YOLO 这类单阶段检测模型，在推理速度、工程接入难度和部署灵活性之间取得了较好的平衡，因此非常适合拿来做教学演示和原型系统。",
    82: "本项目中的 `WebDetectionSystem` 就是在这样的背景下搭建起来的。从实际代码来看，它并不只负责“识别图片中有没有缺陷”，还把检测结果继续用于等级判定、历史记录保存、日志统计、库存增加和销售出库。也就是说，这个系统尝试把“算法识别”向“业务流程处理”再往前推进一步。因此，以这个项目为基础撰写论文，比单纯讨论模型精度更能体现系统设计和工程实现的完整过程。",
    85: "从相关研究脉络来看，国外在木材缺陷检测和工业视觉识别方面起步较早。早期研究更依赖纹理特征、边缘信息和图像分割，再结合支持向量机、决策树等传统方法进行分类。随着卷积神经网络和目标检测框架的发展，研究重点逐渐转向端到端识别、实时检测和复杂场景下的稳定输出。",
    86: "在工程应用层面，国外不少制造行业已经把机器视觉系统嵌入实际业务流程中，用于在线分拣、异常报警和质量统计。这说明当前国外研究并不只关注“能不能识别”，而是更关注模型在复杂环境中的可靠性、跨设备泛化能力以及与生产业务的连接方式。",
    88: "国内在木材缺陷识别、木材质量分级和智能制造系统建设方面的发展同样很快。一方面，Python、PyTorch、OpenCV 等工具降低了实验和开发门槛；另一方面，越来越多的课程设计、毕业设计和科研项目都开始把目标检测算法引入木材表面检测任务。",
    89: "不过，从当前常见的课程设计和毕业设计作品来看，很多系统仍然停留在“上传图片、调用模型、显示结果”的层面。真正和业务流程结合得较紧的内容，例如检测记录沉淀、按角色区分权限、库存变化、销售出库和统计报表，往往写得不够细。本文围绕 `WebDetectionSystem` 展开，正是因为这个项目在这些环节上已经做了比较完整的连接。",
    100: "本系统选择 Python，并不只是因为它使用范围广，更重要的是它在当前项目里确实能把多个模块放在同一套环境下完成。Flask 路由处理、YOLO 模型调用、OpenCV 图像读写、pandas 报表导出，都可以直接在 Python 中实现，开发和联调效率都比较高。",
    101: "从代码结构来看，`app.py` 负责页面路由、权限判断、表单提交和数据库写入；`utils/detection.py` 负责模型加载与检测流程；`models.py` 描述用户、检测历史、工作日志、销售记录和库存等实体；`config.py` 管理数据库、上传目录和模型路径。对毕业设计项目而言，这样的组织方式比较直观，也便于在论文中对应说明。",
    103: "Flask 在本项目中承担的是“把各个功能粘合在一起”的角色。登录、检测、历史记录、日志报表、销售管理和后台设置等业务都通过路由组织在一个应用中，逻辑关系比较清楚，适合课程设计阶段快速搭出原型。",
    104: "项目还通过 Flask-Login 完成用户认证和会话控制。普通用户和管理员虽然共用同一个系统，但菜单入口、可访问页面和可执行操作并不完全相同，这一点让系统不再只是简单的页面集合，而具备了基本的角色管理能力。",
    106: "YOLO 是本系统的核心检测模块。选择它的原因很实际：推理速度较快，Python 接口成熟，和 Flask 项目整合时额外工作量不大。在当前实现中，系统通过 Ultralytics 接口加载 `.pt` 权重文件，并在 `process_image` 函数中完成图像读取、模型推理、检测框解析和结果图保存。",
    107: "更重要的是，系统没有停留在模型原始输出层面。模型给出的英文类别会先映射为中文缺陷名称，再依据裂纹、死节和活节的业务规则转换成 A 级、B 级或 C 级板材等级。这样用户在页面中看到的就不是生硬的模型标签，而是能够直接进入业务流程的结果。",
    109: "数据库部分采用 SQLite，主要考虑的是部署简单和演示方便。对于当前这个单机运行的毕业设计项目来说，它已经足够承载用户信息、检测记录、工作日志、销售记录、系统参数和库存数据，不需要额外搭建数据库服务。",
    110: "在这个项目里，数据库并不是简单地“存一份结果”。历史记录分页、缺陷类型搜索、近七天趋势统计、Excel 导出、库存扣减和销售记录保存，全部依赖同一套结构化数据。因此，数据库实际上承担了业务衔接的底层支撑作用。",
    112: "前端实现上，项目没有采用复杂的前后端分离方案，而是直接使用 Jinja2 模板配合 Bootstrap 5 来组织页面。这种方式的优点是实现快、改动直观，比较适合毕业设计环境下边写边调。",
    113: "从最终效果来看，这套前端方案虽然不算新，但足够稳定。它既能完成基础页面展示，也能满足结果预览、图表统计和大图查看等需求，和当前项目的规模是匹配的。",
    117: "从经济角度看，本系统的主要成本在于开发时间、测试图片准备和运行环境调试，而不是商业软件采购。项目使用的 Python、Flask、Ultralytics、OpenCV、Bootstrap 和 Chart.js 都可以直接获取，这一点对于课程设计和毕业设计来说非常重要。",
    119: "从操作层面来看，这个系统的门槛并不高。登录后左侧侧边栏已经把主要功能集中展示出来，用户基本不需要反复查找入口。检测页面同时支持上传图片和摄像头拍照，管理员页面则通过表单方式完成用户维护和参数设置，整体使用路径比较清楚。",
    121: "技术可行性主要体现在两点。第一，项目使用的组件都比较成熟，出现问题时容易查资料和定位；第二，当前代码已经能够跑通从登录、检测到库存和销售记录更新的完整流程，说明技术路线本身没有明显阻塞点。",
    122: "另外，系统核心逻辑主要集中在 `app.py` 与 `utils/detection.py` 两部分，这使得项目后续无论是替换模型、扩展业务，还是继续做部署优化，都有比较明确的修改入口。",
    124: "结合当前程序实现，系统目前围绕普通用户和管理员两类角色展开。普通用户主要负责登录、检测、查看历史记录和浏览报表；管理员在此基础上还承担用户维护、参数调整和模型切换等工作。这样的角色划分和当前代码是一致的，并不是停留在论文层面的空设定。",
    128: "如果按程序结构来拆分，这个系统可以理解为四层：页面展示层、业务控制层、模型推理层和数据存储层。这种划分方式和项目目录能够一一对应，因此论文中的模块分析和实际代码实现之间没有明显脱节。",
    132: "性能需求在本项目里更多体现为“够用且稳定”。单张图片检测需要在较短时间内给出结果，批量检测时页面不能长时间卡住；历史记录查询、图表统计和 Excel 导出也应保持基本流畅。对毕业设计演示场景来说，这样的目标比单纯追求理论高并发更实际。",
    134: "从项目目录可以比较直观地看出，这个系统并没有做复杂拆分，而是采用了“单应用入口 + 多模板页面 + 工具模块 + 静态资源目录”的组织方式。这样的结构优点是清楚，老师在答辩或检查代码时也更容易顺着目录理解系统实现。",
    274: "本文不是从一个完全抽象的系统方案出发，而是围绕已经完成的 `WebDetectionSystem` 项目来展开分析。从代码、页面到数据库结构，论文中的大部分内容都能在当前程序中找到对应实现。这一点使论文和项目之间的对应关系比较明确，也减少了空泛描述。",
    275: "后续如果继续完善该系统，我认为最值得推进的有几个方向。首先是模型本身，目前识别类别仍然较少，后续可以继续扩展数据集并提升复杂场景下的稳定性；其次是数据库和部署方式，当前项目采用 SQLite 单机运行，更适合教学演示，若要支持多人协作，就需要考虑迁移数据库和拆分部署；再次是业务细节，例如真正删除对应图片文件、补充更细的权限控制、完善日志审计和销售统计。如果条件允许，还可以继续尝试接入移动端拍照采集或现场摄像设备，使系统更贴近真实检测场景。"
}


def replace_text(para, new_text):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def main():
    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(DOC_PATH)
    for idx, new_text in REWRITE.items():
        if idx < len(doc.paragraphs):
            replace_text(doc.paragraphs[idx], new_text)
    doc.save(DOC_PATH)
    print(len(REWRITE))


if __name__ == "__main__":
    main()
