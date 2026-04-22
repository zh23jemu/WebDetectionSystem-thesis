import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


TARGET_DOC = Path(
    r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx"
)
REFERENCE_DOC = Path(
    r"C:\Coding\WebDetectionSystem-thesis\13.张海娇+基于Python的个性化服装推荐系统的设计与实现（6）.docx"
)
BACKUP_DOC = TARGET_DOC.with_name(TARGET_DOC.stem + "_字体统一前备份.docx")


def set_style_font(style, west_name=None, east_name=None, size_pt=None, bold=None):
    """统一样式级别的字体，优先修复继承链问题。"""
    if west_name:
        style.font.name = west_name
    if east_name:
        style.element.rPr.rFonts.set(qn("w:eastAsia"), east_name)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold


def set_run_font(run, west_name=None, east_name=None, size_pt=None, bold=None):
    """对已有文本执行直接格式化，修正被手工改乱的 run。"""
    if west_name:
        run.font.name = west_name
    if east_name:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def format_paragraph_runs(paragraph, west_name=None, east_name=None, size_pt=None, bold=None):
    for run in paragraph.runs:
        if run.text:
            set_run_font(run, west_name=west_name, east_name=east_name, size_pt=size_pt, bold=bold)


def main():
    shutil.copy2(TARGET_DOC, BACKUP_DOC)

    # 读取参考文档只是为了确认采用的版式口径，真正修复直接写到目标论文。
    _ = Document(str(REFERENCE_DOC))
    doc = Document(str(TARGET_DOC))

    # 参考论文口径：
    # 一级标题：14pt，中文仿宋_GB2312
    # 二级标题：15pt，中文黑体，加粗
    # 三级标题：12pt，中文仿宋_GB2312，加粗可保留以增强层级
    # 正文：12pt，西文 Times New Roman，中文宋体
    set_style_font(doc.styles["Heading 1"], west_name="Times New Roman", east_name="仿宋_GB2312", size_pt=14, bold=True)
    set_style_font(doc.styles["Heading 2"], west_name="Times New Roman", east_name="黑体", size_pt=15, bold=True)
    set_style_font(doc.styles["Heading 3"], west_name="Times New Roman", east_name="仿宋_GB2312", size_pt=12, bold=True)
    set_style_font(doc.styles["Normal"], west_name="Times New Roman", east_name="宋体", size_pt=12, bold=False)

    # 目录样式也一并修正，避免目录更新后再次出现字体混乱。
    for toc_name in ["toc 1", "toc 2", "toc 3"]:
        if toc_name in [s.name for s in doc.styles]:
            set_style_font(doc.styles[toc_name], west_name="Times New Roman", east_name="宋体", size_pt=10.5, bold=False)

    # 对已有段落做一次直接修正，确保手工改乱的 run 被统一回来。
    for para in doc.paragraphs:
        style_name = para.style.name
        text = para.text.strip()

        if style_name == "Heading 1":
            format_paragraph_runs(para, west_name="Times New Roman", east_name="仿宋_GB2312", size_pt=14, bold=True)
        elif style_name == "Heading 2":
            format_paragraph_runs(para, west_name="Times New Roman", east_name="黑体", size_pt=15, bold=True)
        elif style_name == "Heading 3":
            format_paragraph_runs(para, west_name="Times New Roman", east_name="仿宋_GB2312", size_pt=12, bold=True)
        elif style_name.startswith("toc "):
            format_paragraph_runs(para, west_name="Times New Roman", east_name="宋体", size_pt=10.5, bold=False)
        elif style_name == "Normal":
            # 图题、表题保持宋体小四更自然；普通正文沿用正文样式。
            if text.startswith("图") or text.startswith("表"):
                format_paragraph_runs(para, west_name="Times New Roman", east_name="宋体", size_pt=10.5, bold=False)
            else:
                format_paragraph_runs(para, west_name="Times New Roman", east_name="宋体", size_pt=12, bold=False)

    doc.save(str(TARGET_DOC))
    print("done")


if __name__ == "__main__":
    main()
