import re
from pathlib import Path

from docx import Document


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")


def normalize_captions():
    doc = Document(DOC_PATH)
    pattern = re.compile(r"^【(图\d+-\d+\s*.+)】$")

    changed = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        match = pattern.match(text)
        if not match:
            continue

        new_text = match.group(1)
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)
        changed += 1

    doc.save(DOC_PATH)
    return changed


if __name__ == "__main__":
    print(normalize_captions())
