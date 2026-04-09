import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_格式调整前备份.docx")


def set_run_font(run, east_asia, latin="Times New Roman", size=12, bold=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def apply_para_runs(paragraph, east_asia, latin="Times New Roman", size=12, bold=None):
    for run in paragraph.runs:
        if run.text:
            set_run_font(run, east_asia=east_asia, latin=latin, size=size, bold=bold)


def reset_paragraph_format(paragraph, *, align=None, first_line_indent=None, line_spacing=1.5, space_before=0, space_after=0):
    fmt = paragraph.paragraph_format
    paragraph.alignment = align
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = line_spacing
    fmt.first_line_indent = first_line_indent


def is_heading1(text):
    return bool(re.match(r"^\d+\s*", text)) or text in {"参考文献", "致谢"}


def is_heading2(text):
    return bool(re.match(r"^\d+\.\d+\s*", text))


def is_heading3(text):
    return bool(re.match(r"^\d+\.\d+\.\d+\s*", text))


def format_doc():
    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(DOC_PATH)

    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        sec.header_distance = Cm(0)
        sec.footer_distance = Cm(1.0)

    title_text = "基于Python的Web木材表面缺陷智能检测系统的设计与实现"
    english_title = "Design and Implementation of a Python-based Web Wood Surface Defect Detection System"

    body_started = False
    in_toc = False

    for p in doc.paragraphs:
        text = p.text.strip()
        style_name = p.style.name if p.style else ""

        if not text:
            reset_paragraph_format(p, align=None, first_line_indent=None, line_spacing=1.5)
            continue

        if text == "目 录":
            in_toc = True
            reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=Pt(18))
            apply_para_runs(p, east_asia="黑体", size=14, bold=False)
            continue

        if in_toc and style_name.startswith("toc"):
            reset_paragraph_format(p, align=None, first_line_indent=None, line_spacing=Pt(18))
            apply_para_runs(p, east_asia="宋体", size=12, bold=False)
            continue

        if text == title_text:
            in_toc = False
            reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=1.5)
            apply_para_runs(p, east_asia="黑体", size=16, bold=False)
            continue

        if text == english_title:
            reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=1.5)
            apply_para_runs(p, east_asia="Times New Roman", latin="Times New Roman", size=16, bold=False)
            continue

        if text.startswith("计算机科学与技术专业") or text.startswith("指导教师") or text.startswith("Computer Science and Technology") or text.startswith("Tutor"):
            reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=1.5)
            apply_para_runs(p, east_asia="宋体", size=12, bold=False)
            continue

        if text.startswith("摘要："):
            reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None, line_spacing=1.0)
            for idx, run in enumerate(p.runs):
                if idx == 0:
                    set_run_font(run, east_asia="黑体", size=10.5, bold=False)
                else:
                    set_run_font(run, east_asia="楷体_GB2312", size=10.5, bold=False)
            continue

        if text.startswith("关键词："):
            reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None, line_spacing=1.0)
            for idx, run in enumerate(p.runs):
                if idx == 0:
                    set_run_font(run, east_asia="黑体", size=10.5, bold=False)
                else:
                    set_run_font(run, east_asia="楷体_GB2312", size=10.5, bold=False)
            continue

        if text.startswith("Abstract:") or text.startswith("Abstract："):
            reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None, line_spacing=1.0)
            apply_para_runs(p, east_asia="Times New Roman", latin="Times New Roman", size=10.5, bold=False)
            continue

        if text.startswith("Keywords:"):
            reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None, line_spacing=1.0)
            apply_para_runs(p, east_asia="Times New Roman", latin="Times New Roman", size=10.5, bold=False)
            continue

        if is_heading3(text):
            body_started = True
            reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None, line_spacing=1.5, space_before=0, space_after=0)
            apply_para_runs(p, east_asia="仿宋_GB2312", size=12, bold=False)
            continue

        if is_heading2(text):
            body_started = True
            reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None, line_spacing=1.5, space_before=0, space_after=0)
            apply_para_runs(p, east_asia="黑体", size=12, bold=True)
            continue

        if is_heading1(text):
            body_started = True
            reset_paragraph_format(p, align=None, first_line_indent=None, line_spacing=1.5, space_before=6, space_after=6)
            apply_para_runs(p, east_asia="仿宋_GB2312", size=14, bold=False)
            continue

        if text.startswith("图") or text.startswith("表") or "截图】" in text or "结构图】" in text or "顺序图】" in text or "E-R 图】" in text:
            reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=1.5)
            apply_para_runs(p, east_asia="宋体", size=10.5, bold=False)
            continue

        if body_started:
            reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Pt(24), line_spacing=1.5)
            apply_para_runs(p, east_asia="宋体", size=12, bold=False)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = p.text.strip()
                    if not txt:
                        continue
                    reset_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=1.5)
                    apply_para_runs(p, east_asia="宋体", size=10.5, bold=False)

    doc.save(DOC_PATH)


if __name__ == "__main__":
    format_doc()
