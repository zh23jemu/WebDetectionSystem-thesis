from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的木质板材缺陷检测系统的设计与实现.docx")
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_补代码前备份.docx")


CODE_MAP = {
    "图5-1 登录界面截图": """# app.py
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        login_role = request.form.get('login_role')
        user = User.query.filter_by(username=username).first()

        if user and check_password_hash(user.password, password):
            if user.role != login_role:
                flash('身份不匹配，请切换到对应标签页登录。', 'warning')
                return redirect(url_for('login'))
            login_user(user)
            return redirect(url_for('dashboard'))
    return render_template('login.html')""",
    "图5-3 历史记录页面截图": """# app.py
@app.route('/history', methods=['GET', 'POST'])
@login_required
def history():
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '')
    query = DetectionHistory.query
    if current_user.role != 'admin':
        query = query.filter_by(user_id=current_user.id)
    if search:
        query = query.filter(DetectionHistory.defect_type.contains(search))
    records = query.order_by(
        DetectionHistory.upload_date.desc()
    ).paginate(page=page, per_page=12)
    return render_template('history.html', records=records)""",
    "图5-4 日志与报表页面截图": """# app.py
@app.route('/logs')
@login_required
def logs():
    today_count = hist_query.filter(
        func.date(DetectionHistory.upload_date) == date.today()
    ).count()

    stats_data = [[k, v] for k, v in counter.items()]
    trend_data = []
    for i in range(6, -1, -1):
        target_date = date.today() - timedelta(days=i)
        count = hist_query.filter(
            func.date(DetectionHistory.upload_date) == target_date
        ).count()
        trend_data.append(count)""",
    "图5-5 工作时长统计截图": """# app.py
total_seconds = 0
for l in log_query.filter(WorkLog.logout_time != None).all():
    total_seconds += (l.logout_time - l.login_time).total_seconds()

if session.get('current_log_id'):
    curr = db.session.get(WorkLog, session['current_log_id'])
    if curr:
        total_seconds += (datetime.now() - curr.login_time).total_seconds()

work_time_str = (
    f"{int(total_seconds // 3600)}小时 "
    f"{int((total_seconds % 3600) // 60)}分钟"
)""",
    "图5-6 销售与库存页面截图": """# app.py
@app.route('/sales', methods=['GET', 'POST'])
@login_required
def sales():
    if request.method == 'POST':
        grade = request.form.get('grade')
        qty = int(request.form.get('quantity'))
        inv = db.session.get(Inventory, grade)
        if not inv or inv.count < qty:
            flash(f'{grade} 库存不足！', 'danger')
        else:
            inv.count -= qty
            sale = SalesRecord(
                customer_name=request.form.get('customer'),
                product_grade=grade, quantity=qty
            )
            db.session.add(sale)""",
    "图5-7 帮助中心页面截图": """# app.py / templates/help.html
@app.route('/help')
@login_required
def help_center():
    return render_template('help.html')

<img src="{{ url_for('static', filename='img/live_knot.png') }}"
     class="card-img-top" alt="活节样张">
<img src="{{ url_for('static', filename='img/dead_knot.png') }}"
     class="card-img-top" alt="死节样张">
<img src="{{ url_for('static', filename='img/crack.png') }}"
     class="card-img-top" alt="裂纹样张">""",
    "图5-8 用户管理页面截图": """# app.py
@app.route('/admin/users', methods=['GET', 'POST'])
@login_required
def admin_users():
    if current_user.role != 'admin':
        return redirect(url_for('dashboard'))
    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'add':
            new_user = User(
                username=request.form.get('username'),
                password=generate_password_hash(request.form.get('password')),
                role=request.form.get('role')
            )
            db.session.add(new_user)
            db.session.commit()""",
    "图5-9 系统设置页面截图": """# app.py
@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    val = request.form.get('conf_threshold')
    s = SystemSettings.query.filter_by(key_name='conf_threshold').first()
    s.value = val

    f = request.files.get('model_file')
    if f and f.filename.endswith('.pt'):
        fname = secure_filename(f.filename)
        f.save(os.path.join(app.root_path, fname))
        m = SystemSettings.query.filter_by(key_name='current_model').first()
        m.value = fname""",
    "图5-10 检测页面截图": """# templates/detect.html
<ul class="nav nav-tabs mb-3" id="detectTab" role="tablist">
    <li class="nav-item">
        <button class="nav-link active" id="upload-tab"
                data-bs-target="#upload-pane" type="button">
            本地上传
        </button>
    </li>
    <li class="nav-item">
        <button class="nav-link" id="camera-tab"
                data-bs-target="#camera-pane" type="button">
            摄像头采集
        </button>
    </li>
</ul>""",
    "图5-12 等级判定与库存联动截图": """# app.py
detected_grade = 'A级'
if res['is_clean']:
    detected_grade = 'A级'
else:
    for d in res['defects'].split(','):
        zh_name = DEFECT_MAPPING.get(d.strip(), d.strip())
        if zh_name in ['裂纹', '死节']:
            detected_grade = 'C级'
            break
        elif zh_name == '活节':
            detected_grade = 'B级'

inv = db.session.get(Inventory, detected_grade)
if inv:
    inv.count += 1""",
    "图5-13 检测结果持久化截图": """# app.py
history = DetectionHistory(
    filename=filename,
    result_image=res['filepath'],
    defect_type=final_defect_str,
    confidence=res.get('confidence', 0.0),
    user_id=current_user.id,
    grade=detected_grade
)
db.session.add(history)
results.append(res)
db.session.commit()""",
}


def insert_paragraph_after(paragraph, text=""):
    """在指定段落后插入新段落。"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def format_code_paragraph(paragraph):
    """仅设置新增代码段的样式，不触碰其他现有段落格式。"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(0.74)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.runs[0]
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)


def find_prev_caption(paragraphs, idx):
    """从占位位置向前找到对应图题。"""
    for j in range(idx - 1, -1, -1):
        text = paragraphs[j].text.strip()
        if text.startswith("图5-"):
            return text
    return None


def next_nonempty_text(paragraphs, idx):
    """检查占位后是否已经存在代码，避免重复插入。"""
    for j in range(idx + 1, min(idx + 4, len(paragraphs))):
        text = paragraphs[j].text.strip()
        if text:
            return text
    return ""


def main():
    if not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(DOC_PATH)
    inserted = 0

    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        if para.text.strip() == "部分关键代码如下：":
            caption = find_prev_caption(doc.paragraphs, i)
            if caption in CODE_MAP:
                nearby_text = next_nonempty_text(doc.paragraphs, i)
                if not nearby_text.startswith("#"):
                    code_para = insert_paragraph_after(para)
                    code_para.add_run(CODE_MAP[caption])
                    format_code_paragraph(code_para)
                    inserted += 1
                    # 段落集合变化后重新抓取，避免索引错位。
                    doc = Document(DOC_PATH) if False else doc
        i += 1

    doc.save(DOC_PATH)
    print(inserted)


if __name__ == "__main__":
    main()
