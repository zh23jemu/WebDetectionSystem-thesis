from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")
IMG_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\paper_diagrams\reference_style_generated\图4-5_销售出库顺序图_参考风格.png")
CAPTION = "图4-5 销售出库顺序图"


def replace():
    doc = Document(DOC_PATH)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() != CAPTION:
            continue
        img_para = doc.paragraphs[i - 1]
        # clear existing runs / image
        for child in list(img_para._p):
            img_para._p.remove(child)
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt = img_para.paragraph_format
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(3)
        fmt.line_spacing = 1.0
        run = img_para.add_run()
        run.add_picture(str(IMG_PATH), width=Cm(15.5))
        doc.save(DOC_PATH)
        return True
    return False


if __name__ == "__main__":
    print(replace())
