import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOC_PATH = Path(
    r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx"
)
TEMPLATE_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\论文格式参考模版.docx")
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_按模板格式调整前备份.docx")


def set_run_font(run, east_name=None, west_name=None, size_pt=None, bold=None):
    """统一单个 run 的中英文字体与字号。"""
    rpr = run._element.get_or_add_rPr()
    if east_name:
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), east_name)
    if west_name:
        run.font.name = west_name
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), west_name)
        rfonts.set(qn("w:hAnsi"), west_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def format_runs(paragraph, east_name=None, west_name=None, size_pt=None, bold=None):
    for run in paragraph.runs:
        if run.text is not None:
            set_run_font(run, east_name=east_name, west_name=west_name, size_pt=size_pt, bold=bold)


def set_paragraph_common(paragraph, align=None, first_line_chars=None, line_15=True):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE if line_15 else WD_LINE_SPACING.SINGLE
    if line_15:
        pf.line_spacing = 1.5
    else:
        pf.line_spacing = 1.0
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    if first_line_chars is not None:
        # Word 中首行缩进 2 字符近似处理为 24 磅。
        pf.first_line_indent = Pt(24 * (first_line_chars / 2))
    if align is not None:
        paragraph.alignment = align


def apply_style_font(style, east_name=None, west_name=None, size_pt=None, bold=None):
    if east_name:
        if style.element.rPr is None:
            style.element.append(OxmlElement("w:rPr"))
        rpr = style.element.rPr
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), east_name)
    if west_name:
        style.font.name = west_name
        if style.element.rPr is None:
            style.element.append(OxmlElement("w:rPr"))
        rpr = style.element.rPr
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), west_name)
        rfonts.set(qn("w:hAnsi"), west_name)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold


def paragraph_text_equals(paragraph, text):
    return paragraph.text.strip() == text


def rebuild_labeled_paragraph(paragraph, label, body, label_east, body_east, west_name="Times New Roman", size_pt=10.5):
    """将摘要/关键词等标签段重建为“标签+内容”双 run，便于分别控制字体。"""
    for run in paragraph.runs:
        run.text = ""
    if not paragraph.runs:
        paragraph.add_run("")
    paragraph.runs[0].text = ""
    label_run = paragraph.add_run(label)
    body_run = paragraph.add_run(body)
    set_run_font(label_run, east_name=label_east, west_name=west_name, size_pt=size_pt, bold=False)
    set_run_font(body_run, east_name=body_east, west_name=west_name, size_pt=size_pt, bold=False)


def is_heading1(text):
    return text.startswith(("1 ", "2 ", "3 ", "4 ", "5 ", "6 ", "7 ")) and "." not in text.split(" ", 1)[0]


def is_heading2(text):
    prefix = text.split(" ", 1)[0]
    return prefix.count(".") == 1 and prefix.replace(".", "").isdigit()


def is_heading3(text):
    prefix = text.split(" ", 1)[0]
    return prefix.count(".") == 2 and prefix.replace(".", "").isdigit()


def is_figure_caption(text):
    return text.startswith("图") and "-" in text[:8]


def is_table_caption(text):
    return text.startswith("表") and "-" in text[:8]


def is_code_paragraph(text):
    if not text:
        return False
    markers = [
        "@app.route",
        "def ",
        "return ",
        "if ",
        "for ",
        "while ",
        "db.session",
        "render_template",
        "login_required",
        "class ",
        "import ",
        "<img ",
        "<ul ",
        "<form ",
        "switch (",
        "public void ",
        "params.put",
    ]
    return any(text.lstrip().startswith(m) for m in markers) or text.startswith("# ")


def apply_template_format():
    shutil.copy2(DOC_PATH, BACKUP_PATH)
    template_doc = Document(str(TEMPLATE_PATH))
    doc = Document(str(DOC_PATH))

    # 页面设置按模板同步；当前论文仅 1 个 section，直接套首节设置即可。
    tpl_sec = template_doc.sections[0]
    for sec in doc.sections:
        sec.top_margin = tpl_sec.top_margin
        sec.bottom_margin = tpl_sec.bottom_margin
        sec.left_margin = tpl_sec.left_margin
        sec.right_margin = tpl_sec.right_margin
        sec.header_distance = tpl_sec.header_distance
        sec.footer_distance = tpl_sec.footer_distance
        sec.page_width = tpl_sec.page_width
        sec.page_height = tpl_sec.page_height

    # 样式层统一，尽量贴模板说明。
    apply_style_font(doc.styles["Normal"], east_name="宋体", west_name="Times New Roman", size_pt=12, bold=False)
    apply_style_font(doc.styles["Heading 1"], east_name="仿宋_GB2312", west_name="Times New Roman", size_pt=14, bold=False)
    apply_style_font(doc.styles["Heading 2"], east_name="黑体", west_name="Times New Roman", size_pt=12, bold=True)
    apply_style_font(doc.styles["Heading 3"], east_name="仿宋_GB2312", west_name="Times New Roman", size_pt=12, bold=False)

    for toc_name in ["toc 1", "toc 2", "toc 3"]:
        if toc_name in [s.name for s in doc.styles]:
            apply_style_font(doc.styles[toc_name], east_name="宋体", west_name="Times New Roman", size_pt=10.5, bold=False)

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name

        # 封面信息
        if text == "基于Python的Web木材表面缺陷智能检测系统的设计与实现":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            format_runs(para, east_name="黑体", west_name="Times New Roman", size_pt=16, bold=False)
            continue
        if text.startswith("计算机科学与技术专业"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            format_runs(para, east_name="宋体", west_name="Times New Roman", size_pt=12, bold=False)
            continue
        if text.startswith("指导教师"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            format_runs(para, east_name="宋体", west_name="Times New Roman", size_pt=12, bold=False)
            continue

        # 摘要 / 关键词 / 英文摘要
        if text.startswith("摘要：") or text.startswith("关键词："):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_common(para, first_line_chars=None, line_15=False)
            if text.startswith("摘要："):
                rebuild_labeled_paragraph(para, "摘要：", text[len("摘要："):], "黑体", "楷体_GB2312")
            else:
                rebuild_labeled_paragraph(para, "关键词：", text[len("关键词："):], "黑体", "楷体_GB2312")
            continue

        if text.startswith("Design and Implementation"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_common(para, first_line_chars=None, line_15=False)
            format_runs(para, east_name="Times New Roman", west_name="Times New Roman", size_pt=16, bold=False)
            continue
        if text.startswith("Computer Science") or text.startswith("Tutor"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_common(para, first_line_chars=None, line_15=False)
            format_runs(para, east_name="Times New Roman", west_name="Times New Roman", size_pt=12, bold=False)
            continue
        if text.startswith("Abstract:") or text.startswith("Abstract："):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_common(para, first_line_chars=None, line_15=False)
            sep = "Abstract:" if text.startswith("Abstract:") else "Abstract："
            rebuild_labeled_paragraph(para, sep, text[len(sep):], "Times New Roman", "Times New Roman")
            continue
        if text.startswith("Keywords:"):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_common(para, first_line_chars=None, line_15=False)
            rebuild_labeled_paragraph(para, "Keywords:", text[len("Keywords:"):], "Times New Roman", "Times New Roman")
            continue

        # 目录
        if text == "目 录":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_common(para, first_line_chars=None, line_15=False)
            format_runs(para, east_name="黑体", west_name="Times New Roman", size_pt=14, bold=False)
            continue
        if style_name.startswith("toc"):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_common(para, first_line_chars=None, line_15=False)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(18)
            format_runs(para, east_name="宋体", west_name="Times New Roman", size_pt=10.5, bold=False)
            continue

        # 标题
        if is_heading1(text):
            para.style = doc.styles["Heading 1"]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_common(para, first_line_chars=None, line_15=True)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            format_runs(para, east_name="仿宋_GB2312", west_name="Times New Roman", size_pt=14, bold=False)
            continue
        if is_heading2(text):
            para.style = doc.styles["Heading 2"]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_common(para, first_line_chars=None, line_15=True)
            format_runs(para, east_name="黑体", west_name="Times New Roman", size_pt=12, bold=True)
            continue
        if is_heading3(text):
            para.style = doc.styles["Heading 3"]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_common(para, first_line_chars=None, line_15=True)
            format_runs(para, east_name="仿宋_GB2312", west_name="Times New Roman", size_pt=12, bold=False)
            continue

        # 图题、表题、代码说明
        if is_figure_caption(text) or is_table_caption(text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_figure_caption(text) else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_common(para, first_line_chars=None, line_15=False)
            format_runs(para, east_name="黑体", west_name="Times New Roman", size_pt=10.5, bold=False)
            continue
        if text.startswith("部分关键代码如下"):
            set_paragraph_common(para, first_line_chars=2, line_15=True)
            format_runs(para, east_name="宋体", west_name="Times New Roman", size_pt=12, bold=False)
            continue
        if is_code_paragraph(text):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_common(para, first_line_chars=None, line_15=True)
            format_runs(para, east_name="Times New Roman", west_name="Times New Roman", size_pt=12, bold=False)
            continue

        # 参考文献标题/致谢标题
        if paragraph_text_equals(para, "参考文献") or paragraph_text_equals(para, "致谢"):
            para.style = doc.styles["Heading 1"]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_common(para, first_line_chars=None, line_15=True)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            format_runs(para, east_name="仿宋_GB2312", west_name="Times New Roman", size_pt=14, bold=False)
            continue

        # 默认正文
        para.style = doc.styles["Normal"]
        set_paragraph_common(para, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_chars=2, line_15=True)
        format_runs(para, east_name="宋体", west_name="Times New Roman", size_pt=12, bold=False)

    # 表格内容按模板：中文宋体五号；数字字母 TNR 五号
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_common(para, first_line_chars=None, line_15=False)
                    format_runs(para, east_name="宋体", west_name="Times New Roman", size_pt=10.5, bold=False)

    doc.save(str(DOC_PATH))
    print("done")


if __name__ == "__main__":
    apply_template_format()
