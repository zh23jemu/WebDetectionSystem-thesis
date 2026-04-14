import shutil
from pathlib import Path

from docx import Document


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_AI率优化前备份.docx")


REPLACEMENTS = {
    "摘要：木材表面缺陷检测是木材加工、质量分级和销售流转中的关键环节。": (
        "摘要：木材表面缺陷检测直接关系到板材分级、库存统计和后续销售处理。"
        "在很多中小型加工或教学演示场景中，这项工作往往还是靠人工观察来完成，"
        "不仅效率受人员经验影响较大，而且检测记录分散，后续查询和统计也不方便。"
        "本文结合本人整理和调试的 `WebDetectionSystem` 项目，对一套基于 Python 的 Web 木材表面缺陷智能检测系统进行了设计与实现分析。"
        "系统后端采用 Flask 框架，检测部分调用 Ultralytics YOLO 模型，数据库使用 SQLite，前端页面基于 Jinja2 模板、Bootstrap 5、Chart.js 和 Viewer.js 构建。"
        "从当前程序实现情况看，系统已经能够完成用户登录、图片上传检测、摄像头拍照检测、缺陷类别中文映射、板材等级判定、检测历史保存、统计报表展示、Excel 导出、库存联动、销售出库、帮助中心展示以及管理员参数设置和模型切换等功能。"
        "其中，项目并不是简单地把模型推理结果直接显示出来，而是把检测结果继续映射到等级、库存和销售流程中，使识别模块和业务模块形成了较完整的闭环。"
        "论文围绕该系统的研究背景、关键技术、需求分析、系统设计、功能实现和测试情况展开说明，并结合实际代码结构对主要模块进行梳理。"
        "从课程设计和毕业设计的角度看，该系统具有功能完整、演示直观、实现路径清晰的特点，能够较好地体现 Python Web 开发与目标检测技术结合的工程实践过程。"
    ),
    "Abstract: Wood surface defect inspection is a critical procedure in timber processing, quality grading and sales circulation.": (
        "Abstract: Wood surface defect inspection affects timber grading, stock management and subsequent sales processing. "
        "In many small-scale processing and teaching scenarios, this work is still completed manually, which makes the final result highly dependent on experience and also causes inconvenience in record tracing and data statistics. "
        "Based on the actual implementation of the `WebDetectionSystem` project, this paper studies the design and implementation of a Python-based web wood surface defect detection system. "
        "The system uses Flask as the backend framework, Ultralytics YOLO as the detection engine, SQLite as the database, and Jinja2 templates with Bootstrap 5, Chart.js and Viewer.js to build the web pages. "
        "According to the current program structure, the system already supports user login, image upload detection, camera capture, Chinese defect mapping, timber grading, detection history storage, report visualization, Excel export, inventory linkage, sales outbound processing, help-center presentation, and administrator-side parameter and model management. "
        "A key feature of this project is that the model output is not displayed alone; instead, it is further connected with grade judgement, inventory update and sales records, which makes the whole system closer to a complete business workflow rather than a simple algorithm demo. "
        "This paper explains the system from the aspects of background, key technologies, requirement analysis, system design, implementation and testing, and summarizes the main modules according to the actual codebase. "
        "The final result shows that the project is suitable for course design, teaching demonstration and graduation project presentation."
    ),
    "随着木材加工、家具制造、建筑装饰以及板材贸易等行业不断发展，木材表面质量检测的重要性日益凸显。": (
        "木材表面缺陷检测并不是一个孤立的识别问题，它和板材的分级、入库、销售报价乃至后续质量追溯都有直接关系。"
        "在家具制造、建筑装饰和板材流通等场景中，裂纹、死节、活节这类缺陷会明显影响木材外观和使用价值。"
        "如果前端检测环节出现误判，后面的库存统计和销售处理也会受到影响。"
    ),
    "近年来，深度学习和计算机视觉技术在工业检测中的应用逐渐成熟，目标检测算法能够从图像中快速定位缺陷区域并输出类别结果，为传统人工检测模式提供了新的升级路径。": (
        "近几年，目标检测算法在工业视觉领域的应用越来越成熟。"
        "与传统依赖人工经验的检查方式相比，基于图像的自动检测至少有两个明显优势：一是结果输出速度更快，二是检测记录可以直接保存为结构化数据，便于后续查询、统计和追溯。"
        "特别是 YOLO 这类单阶段检测模型，在推理速度和工程接入难度之间取得了比较合适的平衡，很适合拿来做课程设计原型系统。"
    ),
    "本项目的 `WebDetectionSystem` 正是在上述需求背景下构建的一个木材表面缺陷智能检测系统。": (
        "本项目中的 `WebDetectionSystem` 就是在这样的背景下搭建起来的。"
        "从实际代码来看，它并不只负责“识别一张图片里有没有缺陷”，还把检测结果继续用于等级判定、历史记录保存、报表统计、库存更新和销售出库。"
        "也就是说，这个系统试图把“算法识别”向“业务流程处理”再往前推进一步。"
        "因此，以这个项目为基础撰写论文，比单纯讨论模型精度更能体现完整的系统设计思路。"
    ),
    "国外在木材缺陷检测和工业视觉识别领域起步较早，研究重点主要集中在纹理分析、特征提取、图像分割、传统机器学习分类以及深度神经网络识别等方向。": (
        "从公开资料和相关研究脉络来看，国外在木材缺陷检测和工业视觉识别方面起步更早。"
        "早期工作更多依赖纹理特征、边缘信息和图像分割，再结合支持向量机、决策树等传统方法完成分类。"
        "后来随着卷积神经网络和目标检测框架的发展，研究重点逐步转向端到端识别和实时检测。"
    ),
    "在工程应用层面，国外木材、钢材、纺织和电子制造行业较早尝试将机器视觉系统嵌入生产线，实现缺陷自动报警、在线分级和品质统计。": (
        "在工程落地方面，国外不少制造行业已经把机器视觉系统嵌入实际流程中，用于在线分拣、异常报警和质量统计。"
        "因此，当前国外研究已经不只关注“能不能识别”，而是更关注模型在复杂场景中的稳定性、泛化能力以及与生产业务的衔接方式。"
    ),
    "国内在木材缺陷识别、木材质量分级以及智能制造系统建设方面同样发展迅速。": (
        "国内在木材缺陷识别和木材质量分级方面的研究发展也很快。"
        "一方面，Python、PyTorch、OpenCV 等工具降低了实验门槛；另一方面，越来越多的课程设计、毕业设计和科研项目都开始把目标检测算法引入木材表面检测任务。"
    ),
    "在课程设计和毕业设计实践中，基于 Flask 或 Django 的智能检测系统数量逐年增加，项目往往能够完成图像上传、模型推理、结果显示等基础功能。": (
        "不过，从当前常见的课程设计和毕业设计作品来看，很多系统仍然停留在“上传图片、调用模型、显示结果”这一层。"
        "真正和业务流程结合得比较紧的内容，比如检测记录沉淀、按角色区分权限、库存变化、销售出库和数据报表，往往写得比较少。"
        "这也是本文选择围绕 `WebDetectionSystem` 展开的原因，因为这个项目在这些环节上已经做了相对完整的连接。"
    ),
    "Python 语法简洁、生态丰富、开发效率高，广泛应用于 Web 开发、数据分析、人工智能和自动化运维等领域。": (
        "本系统之所以选择 Python，不只是因为它“流行”，更关键的是它在当前项目里确实能把几个原本分散的功能串起来。"
        "例如，Flask 路由处理、YOLO 模型调用、OpenCV 图像读写、pandas 数据导出，都可以直接在同一套语言环境下完成，开发和调试都比较顺手。"
    ),
    "从项目代码来看，`app.py` 统一处理路由、表单提交、数据库交互和 JSON 响应；": (
        "从代码结构上看，`app.py` 集中处理页面路由、表单提交、权限判断、数据库写入和异步响应；"
        "`utils/detection.py` 单独负责模型加载与检测流程；`models.py` 描述数据库实体；`config.py` 则维护上传目录、结果目录和模型路径。"
        "这种组织方式虽然不算复杂，但对毕业设计项目来说比较直接，便于展示整体实现思路。"
    ),
    "Flask 是一个轻量级 Web 框架，具备路由清晰、扩展灵活、上手成本低等特点。": (
        "Flask 在这个项目里承担的是“把各个功能粘合在一起”的角色。"
        "登录、检测、历史记录、日志报表、销售管理和后台设置等功能都通过路由组织在一个应用里，逻辑关系清楚，适合课程设计阶段快速搭出原型。"
    ),
    "此外，系统还使用 Flask-Login 完成用户身份认证和会话控制。": (
        "项目还通过 Flask-Login 处理用户认证和会话控制。"
        "从实际效果看，普通用户和管理员的入口、菜单和操作权限都做了区分，这一点比单纯把所有页面开放出来更符合系统化设计的要求。"
    ),
    "YOLO（You Only Look Once）是一种典型的单阶段目标检测算法，其核心思想是将目标定位与分类统一到一次前向推理中完成，兼顾较快推理速度与较高检测精度。": (
        "YOLO 是本系统的核心检测模块。"
        "选择它的原因并不复杂：一方面推理速度较快，另一方面 Python 接口比较成熟，拿来和 Flask 项目整合的工作量相对可控。"
        "在当前实现中，系统通过 Ultralytics 接口加载 `.pt` 权重文件，并在 `process_image` 中完成模型推理、检测框解析和结果图输出。"
    ),
    "在业务层面，系统并不直接将模型输出作为最终结果，而是进一步完成英文类别到中文缺陷名称的映射，并通过判定规则将检测结果转换为 A 级、B 级或 C 级板材等级。": (
        "更重要的是，系统没有停留在模型原始输出层面。"
        "检测结果先被映射为中文缺陷名称，再根据裂纹、死节和活节的业务规则转换成板材等级。"
        "这样一来，用户在页面上看到的就不是生硬的模型类别，而是能够直接用于业务判断的结果。"
    ),
    "SQLite 是轻量级关系型数据库，部署方便、无需单独数据库服务进程，适合课程设计、单机原型和教学演示场景。": (
        "数据库部分选用 SQLite，主要考虑的是部署简单和演示方便。"
        "对于当前这个单机运行的毕业设计项目来说，它已经足够承载用户信息、检测记录、工作日志、销售记录、系统参数和库存数据。"
    ),
    "由于检测历史、报表统计与销售管理都依赖结构化数据支撑，数据库不仅承担存储功能，还为后续分页查询、趋势统计、数据导出和业务联动提供基础。": (
        "在这个项目里，数据库并不是简单的“存一份结果”。"
        "历史记录分页、缺陷类型搜索、近七天趋势统计、Excel 导出、库存扣减和销售记录保存，都依赖同一套结构化数据。"
        "因此，数据库实际上承担了业务衔接的底层支撑作用。"
    ),
    "系统前端主要采用 Jinja2 模板 + Bootstrap 5 方式实现页面构建，辅以 Font Awesome 图标库优化可视化表达。": (
        "前端实现上，项目没有采用复杂的前后端分离方案，而是直接使用 Jinja2 模板配合 Bootstrap 5 来组织页面。"
        "这种方式的优点是实现快、改动直观，比较适合毕业设计环境下边写边调。"
    ),
    "这种技术组合实现简单、学习成本低，且与 Flask 模板渲染方式兼容良好，适合毕业设计场景下快速搭建完整的可视化页面体系。": (
        "从最终效果来看，这套前端方案虽然不算新，但足够稳定。"
        "它既能完成基础页面展示，也能满足结果预览、图表统计和大图查看等需求，和当前项目的体量是匹配的。"
    ),
    "系统开发过程中主要使用 Python、Flask、Ultralytics、OpenCV、Bootstrap、Chart.js 等开源技术，能够显著降低软件采购成本。": (
        "从经济角度看，本系统的开发成本主要体现在时间投入和测试环境准备上，而不是商业软件授权费用。"
        "项目使用的 Python、Flask、Ultralytics、OpenCV、Bootstrap、Chart.js 等工具都可以直接获取，比较适合课程设计预算有限的情况。"
    ),
    "从页面交互看，系统采用统一侧边栏导航组织功能，用户可以在仪表盘、开始检测、历史记录、日志报表、销售管理和帮助中心之间快速切换。": (
        "操作层面上，这个系统的门槛并不高。"
        "登录后左侧侧边栏会把主要功能集中展示出来，用户基本不需要反复查找入口。"
        "检测页面既支持上传图片，也支持摄像头拍照，管理员页面则通过表单方式完成用户维护和参数设置，整体操作路径比较清楚。"
    ),
    "本系统所使用的关键技术均较为成熟。": (
        "技术可行性主要体现在两点。"
        "第一，项目使用的组件都比较成熟，出现问题时容易查资料和定位；第二，当前代码已经能够跑通完整流程，说明技术路线本身没有明显阻塞点。"
    ),
    "根据当前程序实现情况，系统面向普通用户和管理员两类角色。": (
        "结合实际程序实现，系统目前围绕普通用户和管理员两类角色展开。"
        "普通用户主要负责检测、查看记录和浏览报表；管理员在此基础上还承担用户维护、阈值调整和模型切换等工作。"
    ),
    "结合当前项目目录结构，系统功能可划分为界面展示层、业务逻辑层、模型推理层和数据存储层四个部分。": (
        "如果按程序结构来拆分，这个系统可以理解为四层：页面展示层、业务控制层、模型推理层和数据存储层。"
        "这种划分方式和项目目录是能一一对应上的，因此论文里写起来也比较自然，不会出现“设计一套、代码又是另一套”的情况。"
    ),
    "在性能方面，系统需要保证单张图片上传后能够在可接受时间内返回结果，批量图片检测时页面仍能保持基本响应；": (
        "性能需求在本项目里更多体现为“够用且稳定”。"
        "单张图片检测需要在较短时间内返回结果，批量检测时页面不能长时间卡死；历史记录查询和报表生成也应保持基本流畅。"
        "对于课程设计场景来说，这样的性能目标是比较现实的。"
    ),
    "结合项目目录遍历结果可以看出，程序结构围绕“单应用入口 + 多模板页面 + 工具模块 + 静态资源”展开。": (
        "从项目目录可以比较直观地看出，这个系统并没有做复杂拆分，而是采用了“单应用入口 + 多模板页面 + 工具模块 + 静态资源目录”的组织方式。"
        "这种结构的优点是清楚，答辩时老师也更容易顺着目录理解系统实现。"
    ),
    "本文围绕当前 `WebDetectionSystem` 项目的实际程序结构，对基于 Python 的 Web 木材表面缺陷智能检测系统进行了完整整理与说明。": (
        "本文并不是从一个完全抽象的系统方案出发，而是围绕已经完成的 `WebDetectionSystem` 项目来展开分析。"
        "从代码、页面到数据库结构，论文内容基本都能在当前程序里找到对应实现。"
        "这一点使得论文和项目之间的对应关系比较明确，也减少了空泛描述。"
    ),
    "后续展望主要包括以下几个方向：其一，进一步优化模型训练数据和类别体系，提高在复杂木纹、光照变化和多缺陷叠加场景下的泛化能力；": (
        "后续如果继续完善这个系统，我认为最值得推进的有几个方向。"
        "首先是模型本身，目前识别类别还比较有限，后续可以继续扩展数据集并提升复杂场景下的稳定性；"
        "其次是数据库和部署方式，当前项目采用 SQLite 单机运行，更适合教学演示，若要支持多人协作，就需要考虑迁移数据库和拆分部署；"
        "再次是业务细节，例如真正删除对应图片文件、补充更细的权限控制、完善日志审计和销售统计。"
        "如果条件允许，还可以继续尝试接入移动端拍照采集或现场摄像设备，使系统更贴近真实检测场景。"
    ),
}


def replace_text(para, new_text):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def main():
    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(DOC_PATH)
    changed = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        for prefix, new_text in REPLACEMENTS.items():
            if text.startswith(prefix):
                replace_text(para, new_text)
                changed += 1
                break
    doc.save(DOC_PATH)
    print(changed)


if __name__ == "__main__":
    main()
