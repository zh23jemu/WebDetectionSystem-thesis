from pathlib import Path
import re

from docx import Document


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")


def set_para_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def main():
    doc = Document(DOC_PATH)
    in_ch5 = False
    count = 1
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "5 系统实现":
            in_ch5 = True
            continue
        if in_ch5 and text == "6 系统测试":
            break
        m = re.match(r"^图5-\d+(\s+.*)$", text)
        if in_ch5 and m:
            set_para_text(p, f"图5-{count}{m.group(1)}")
            count += 1
    doc.save(DOC_PATH)
    print(count - 1)


if __name__ == "__main__":
    main()
