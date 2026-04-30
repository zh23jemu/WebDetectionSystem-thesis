from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


# 该脚本用于把论文 docx 内容抽取为 Markdown，便于后续直接修改正文。
# 转换时只读取 Word 文件，不修改原论文。
DOCX_PATH = Path(
    r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的木质板材缺陷检测系统的设计与实现.docx"
)
OUT_PATH = DOCX_PATH.with_suffix(".md")


def iter_block_items(document):
    """按 Word 文档中的原始顺序遍历段落和表格。"""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", child
        elif child.tag == qn("w:tbl"):
            yield "table", child


def paragraph_from_element(document, element):
    """由底层 XML 段落元素找到 python-docx 段落对象。"""
    for paragraph in document.paragraphs:
        if paragraph._element is element:
            return paragraph
    return None


def table_from_element(document, element):
    """由底层 XML 表格元素找到 python-docx 表格对象。"""
    for table in document.tables:
        if table._element is element:
            return table
    return None


def heading_level(paragraph):
    """根据 Word 样式判断 Markdown 标题层级。"""
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading"):
        parts = style_name.split()
        if len(parts) == 2 and parts[1].isdigit():
            return min(int(parts[1]), 6)
    text = paragraph.text.strip()
    if text and text[0].isdigit():
        dot_count = text.split()[0].count(".")
        if text[0].isdigit() and dot_count == 0:
            return 1
        if dot_count == 1:
            return 2
        if dot_count >= 2:
            return 3
    return None


def has_picture(paragraph):
    """判断段落内是否包含图片、绘图对象或嵌入对象。"""
    xml = paragraph._element.xml
    return any(tag in xml for tag in ("<w:drawing", "<w:pict", "<v:shape", "<v:imagedata"))


def is_code_like(text):
    """粗略识别论文里的代码段，保留其原始换行和缩进。"""
    markers = [
        "# app.py",
        "# templates/",
        "# utils/",
        "@app.route",
        "@login_required",
        "def ",
        "<ul ",
        "<form ",
        "<img ",
        "DetectionHistory(",
        "db.session",
    ]
    return any(marker in text for marker in markers)


def clean_cell_text(text):
    """清理表格单元格文本，避免破坏 Markdown 表格。"""
    return " ".join(text.replace("|", "\\|").split())


def table_to_markdown(table):
    """把 Word 表格转为 Markdown 表格。"""
    rows = []
    for row in table.rows:
        rows.append([clean_cell_text(cell.text) for cell in row.cells])

    if not rows:
        return []

    max_cols = max(len(row) for row in rows)
    normalized = [row + [""] * (max_cols - len(row)) for row in rows]
    header = normalized[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * max_cols) + " |",
    ]
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def paragraph_to_markdown(paragraph, image_index):
    """把单个段落转为 Markdown 行，并返回新的图片序号。"""
    lines = []
    text = paragraph.text.rstrip()

    if has_picture(paragraph):
        lines.append(f"[图片占位 {image_index}]")
        image_index += 1
        if text.strip():
            lines.append(text.strip())
        return lines, image_index

    if not text.strip():
        return [""], image_index

    level = heading_level(paragraph)
    if level:
        lines.append(f"{'#' * level} {text.strip()}")
        return lines, image_index

    if text.strip().startswith("部分关键代码如下：") and "\n" in text:
        lead, code = text.split("\n", 1)
        lines.append(lead.strip())
        lines.append("")
        language = "html" if "<" in code and ">" in code else "python"
        lines.append(f"```{language}")
        lines.extend(code.strip("\n").splitlines())
        lines.append("```")
        return lines, image_index

    if is_code_like(text):
        language = "html" if "<" in text and ">" in text else "python"
        lines.append(f"```{language}")
        lines.extend(text.strip("\n").splitlines())
        lines.append("```")
        return lines, image_index

    lines.append(text.strip())
    return lines, image_index


def count_embedded_images(docx_path):
    """统计 docx 包内的图片数量，作为导出结果的核对信息。"""
    with ZipFile(docx_path) as archive:
        return len([name for name in archive.namelist() if name.startswith("word/media/")])


def main():
    document = Document(DOCX_PATH)
    output = []
    image_index = 1

    output.append(f"<!-- 来源：{DOCX_PATH} -->")
    output.append("")

    for kind, element in iter_block_items(document):
        if kind == "paragraph":
            paragraph = paragraph_from_element(document, element)
            if paragraph is None:
                continue
            lines, image_index = paragraph_to_markdown(paragraph, image_index)
            output.extend(lines)
            output.append("")
        elif kind == "table":
            table = table_from_element(document, element)
            if table is None:
                continue
            output.extend(table_to_markdown(table))
            output.append("")

    image_count = count_embedded_images(DOCX_PATH)
    output.append(f"<!-- 共检测到 Word 内嵌图片 {image_count} 张，Markdown 中以图片占位标记。 -->")
    OUT_PATH.write_text("\n".join(output).rstrip() + "\n", encoding="utf-8")
    print(OUT_PATH)


if __name__ == "__main__":
    main()
