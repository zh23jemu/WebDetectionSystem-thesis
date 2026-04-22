import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_按模板最终套版前备份.docx")


TITLE_TEXT = "基于Python的Web木材表面缺陷智能检测系统的设计与实现"
EN_TITLE = "Design and Implementation of a Python-based Web Wood Surface Defect Detection System"


def set_run_font(run, east_asia, latin="Times New Roman", size=12, bold=None):
    run.font.name = latin
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def apply_para_runs(paragraph, east_asia, latin="Times New Roman", size=12, bold=None):
    for run in paragraph.runs:
        set_run_font(run, east_asia=east_asia, latin=latin, size=size, bold=bold)


def set_paragraph_format(paragraph, *, align=None, first_line_indent=None, line_spacing=1.5, space_before=0, space_after=0):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = first_line_indent
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = line_spacing


def is_heading1(text):
    return bool(re.match(r"^\d+\s+", text)) or text in {"参考文献", "致谢"}


def is_heading2(text):
    return bool(re.match(r"^\d+\.\d+\s+", text))


def is_heading3(text):
    return bool(re.match(r"^\d+\.\d+\.\d+\s+", text))


def is_figure_caption(text):
    return bool(re.match(r"^图\d+-\d+\s+", text))


def is_table_caption(text):
    return bool(re.match(r"^表\d+-\d+\s+", text))


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    for child in list(borders):
        borders.remove(child)

    spec = {
        "top": {"val": "single", "sz": "12", "space": "0", "color": "000000"},
        "left": {"val": "nil"},
        "bottom": {"val": "single", "sz": "12", "space": "0", "color": "000000"},
        "right": {"val": "nil"},
        "insideH": {"val": "single", "sz": "8", "space": "0", "color": "000000"},
        "insideV": {"val": "nil"},
    }
    for edge, attrs in spec.items():
        elem = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            elem.set(qn(f"w:{k}"), v)
        borders.append(elem)


def format_tables(doc):
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table)
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.first_line_indent = None
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1.0
                    for run in para.runs:
                        set_run_font(run, east_asia="宋体", latin="Times New Roman", size=10.5, bold=(r_idx == 0))


def clean_empty_runs(paragraph):
    for run in paragraph.runs:
        if run.text == "":
            set_run_font(run, east_asia="宋体", latin="Times New Roman", size=12, bold=False)


def format_document():
    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(DOC_PATH)

    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        sec.header_distance = Cm(0)
        sec.footer_distance = Cm(1.0)

    in_toc = False
    for p in doc.paragraphs:
        text = p.text.strip()
        style_name = p.style.name if p.style else ""
        clean_empty_runs(p)

        if not text:
            # 正文空行不保留首行缩进
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None, line_spacing=1.5)
            continue

        if text == "目 录":
            in_toc = True
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=Pt(18))
            apply_para_runs(p, east_asia="黑体", latin="Times New Roman", size=14, bold=False)
            continue

        if in_toc and style_name.startswith("toc"):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None, line_spacing=Pt(18))
            apply_para_runs(p, east_asia="宋体", latin="Times New Roman", size=12, bold=False)
            continue

        if text == TITLE_TEXT:
            in_toc = False
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=1.5)
            apply_para_runs(p, east_asia="黑体", latin="Times New Roman", size=16, bold=False)
            continue

        if text == EN_TITLE:
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=1.5)
            apply_para_runs(p, east_asia="Times New Roman", latin="Times New Roman", size=16, bold=False)
            continue

        if text.startswith("计算机科学与技术专业") or text.startswith("指导教师"):
            # 与模板一致：小四居中，标题页有一定段前段后
            if text.startswith("计算机科学与技术专业"):
                set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=1.5, space_before=6, space_after=6)
            else:
                set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=1.5, space_before=6, space_after=24)
            apply_para_runs(p, east_asia="宋体", latin="Times New Roman", size=12, bold=False)
            continue

        if text.startswith("Computer Science and Technology") or text.startswith("Tutor"):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=1.5)
            apply_para_runs(p, east_asia="Times New Roman", latin="Times New Roman", size=12, bold=False)
            continue

        if text.startswith("摘要："):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None, line_spacing=1.0)
            # 清空后重建 runs，避免历史空 run 干扰字体
            full_text = p.text
            for run in p.runs:
                run.text = ""
            r1 = p.runs[0] if p.runs else p.add_run()
            r1.text = "摘要："
            set_run_font(r1, east_asia="黑体", latin="Times New Roman", size=10.5, bold=False)
            r2 = p.add_run(full_text[len("摘要："):])
            set_run_font(r2, east_asia="楷体_GB2312", latin="Times New Roman", size=10.5, bold=False)
            continue

        if text.startswith("关键词："):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None, line_spacing=1.0)
            full_text = p.text
            for run in p.runs:
                run.text = ""
            r1 = p.runs[0] if p.runs else p.add_run()
            r1.text = "关键词："
            set_run_font(r1, east_asia="黑体", latin="Times New Roman", size=10.5, bold=False)
            r2 = p.add_run(full_text[len("关键词："):])
            set_run_font(r2, east_asia="楷体_GB2312", latin="Times New Roman", size=10.5, bold=False)
            continue

        if text.startswith("Abstract:") or text.startswith("Abstract："):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None, line_spacing=1.0)
            full_text = p.text
            for run in p.runs:
                run.text = ""
            tag = "Abstract:" if full_text.startswith("Abstract:") else "Abstract："
            r1 = p.runs[0] if p.runs else p.add_run()
            r1.text = tag
            set_run_font(r1, east_asia="Times New Roman", latin="Times New Roman", size=10.5, bold=False)
            r2 = p.add_run(full_text[len(tag):])
            set_run_font(r2, east_asia="Times New Roman", latin="Times New Roman", size=10.5, bold=False)
            continue

        if text.startswith("Keywords:"):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None, line_spacing=1.0)
            full_text = p.text
            for run in p.runs:
                run.text = ""
            r1 = p.runs[0] if p.runs else p.add_run()
            r1.text = "Keywords:"
            set_run_font(r1, east_asia="Times New Roman", latin="Times New Roman", size=10.5, bold=False)
            r2 = p.add_run(full_text[len("Keywords:"):])
            set_run_font(r2, east_asia="Times New Roman", latin="Times New Roman", size=10.5, bold=False)
            continue

        if is_heading1(text):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None, line_spacing=1.5, space_before=6, space_after=6)
            apply_para_runs(p, east_asia="仿宋_GB2312", latin="Times New Roman", size=14, bold=False)
            continue

        if is_heading2(text):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None, line_spacing=1.5, space_before=0, space_after=0)
            apply_para_runs(p, east_asia="黑体", latin="Times New Roman", size=12, bold=True)
            continue

        if is_heading3(text):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None, line_spacing=1.5, space_before=0, space_after=0)
            apply_para_runs(p, east_asia="仿宋_GB2312", latin="Times New Roman", size=12, bold=False)
            continue

        if is_figure_caption(text):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=1.5, space_before=0, space_after=6)
            apply_para_runs(p, east_asia="黑体", latin="Times New Roman", size=10.5, bold=False)
            continue

        if is_table_caption(text):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, line_spacing=1.5, space_before=0, space_after=6)
            apply_para_runs(p, east_asia="黑体", latin="Times New Roman", size=10.5, bold=False)
            continue

        if text.startswith("部分关键代码如下"):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Pt(21), line_spacing=1.5, space_before=0, space_after=0)
            apply_para_runs(p, east_asia="宋体", latin="Times New Roman", size=12, bold=False)
            continue

        if text.startswith("# ") or text.startswith("if __name__") or text.startswith("@app.route") or text.startswith("history = DetectionHistory(") or text.startswith("<"):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Pt(21), line_spacing=1.5, space_before=0, space_after=0)
            apply_para_runs(p, east_asia="Times New Roman", latin="Times New Roman", size=12, bold=False)
            continue

        if text.startswith("[") and re.match(r"^\[\d+\]", text):
            # 参考文献：五号、单倍行距
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None, line_spacing=1.0, space_before=0, space_after=0)
            apply_para_runs(p, east_asia="宋体", latin="Times New Roman", size=10.5, bold=False)
            continue

        # 正文
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Pt(24), line_spacing=1.5, space_before=0, space_after=0)
        apply_para_runs(p, east_asia="宋体", latin="Times New Roman", size=12, bold=False)

    format_tables(doc)
    doc.save(DOC_PATH)


if __name__ == "__main__":
    format_document()
