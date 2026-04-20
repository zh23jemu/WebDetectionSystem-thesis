import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_插入第五章代码展示前备份.docx")


CODE_MAP = {
    "图5-1 登录界面截图": """# app.py
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        login_role = request.form.get('login_role')
        user = User.query.filter_by(username=username).first()

        if user and check_password_hash(user.password, password):
            if user.role != login_role:
                flash('身份不匹配，请切换到对应标签页登录。', 'warning')
                return redirect(url_for('login'))
            login_user(user)
            return redirect(url_for('dashboard'))
    return render_template('login.html')""",
    "图5-2 仪表盘与个人中心截图": """# app.py
@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html')

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    if request.method == 'POST':
        old_pass = request.form.get('old_password')
        new_pass = request.form.get('new_password')
        if check_password_hash(current_user.password, old_pass):
            current_user.password = generate_password_hash(new_pass)
            db.session.commit()
            return redirect(url_for('logout'))
    return render_template('profile.html')""",
    "图5-3 公共布局与导航截图": """# templates/base.html
<a class="nav-link" href="{{ url_for('dashboard') }}">
    <i class="fas fa-home"></i> 仪表盘
</a>
<a class="nav-link" href="{{ url_for('detect') }}">
    <i class="fas fa-camera"></i> 开始检测
</a>
<a class="nav-link" href="{{ url_for('history') }}">
    <i class="fas fa-images"></i> 历史记录
</a>
{% if current_user.role == 'admin' %}
<a class="nav-link" href="{{ url_for('admin_users') }}">
    <i class="fas fa-users-cog"></i> 用户管理
</a>
{% endif %}""",
    "图5-4 检测页面截图": """# templates/detect.html
<ul class="nav nav-tabs mb-3" id="detectTab">
    <li class="nav-item">
        <button class="nav-link active" data-bs-target="#upload-pane"
                type="button" onclick="stopCamera()">本地上传</button>
    </li>
    <li class="nav-item">
        <button class="nav-link" data-bs-target="#camera-pane"
                type="button" onclick="initCamera()">摄像头采集</button>
    </li>
</ul>
<form id="uploadForm">
    <input class="form-control" type="file" id="images"
           name="images" multiple accept="image/*">
    <div class="row g-3 mb-3" id="previewArea"></div>
</form>""",
    "图5-5 推理结果展示截图": """# utils/detection.py
def process_image(image_path, save_dir, model_path, conf_threshold=0.25):
    model = load_model(model_path)
    img = cv2.imread(image_path)
    results = model(img, conf=float(conf_threshold))

    detected_objects = []
    max_conf = 0.0
    for r in results:
        res_plotted = r.plot()
        for box in r.boxes:
            cls_id = int(box.cls[0])
            detected_objects.append(model.names[cls_id])
            max_conf = max(max_conf, float(box.conf[0]))
    return {'defects': ','.join(list(set(detected_objects))),
            'confidence': round(max_conf, 2)}""",
    "图5-6 等级判定与库存联动截图": """# app.py
detected_grade = 'A级'
if res['is_clean']:
    detected_grade = 'A级'
else:
    for d in res['defects'].split(','):
        zh_name = DEFECT_MAPPING.get(d.strip(), d.strip())
        if zh_name in ['裂纹', '死节']:
            detected_grade = 'C级'
            break
        elif zh_name == '活节':
            detected_grade = 'B级'

inv = db.session.get(Inventory, detected_grade)
if inv:
    inv.count += 1""",
    "图5-7 检测结果持久化截图": """# app.py
history = DetectionHistory(
    filename=filename,
    result_image=res['filepath'],
    defect_type=final_defect_str,
    confidence=res.get('confidence', 0.0),
    user_id=current_user.id,
    grade=detected_grade
)
db.session.add(history)
results.append(res)
db.session.commit()""",
    "图5-8 历史记录页面截图": """# app.py
@app.route('/history', methods=['GET', 'POST'])
@login_required
def history():
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '')
    query = DetectionHistory.query
    if current_user.role != 'admin':
        query = query.filter_by(user_id=current_user.id)
    if search:
        query = query.filter(DetectionHistory.defect_type.contains(search))
    records = query.order_by(
        DetectionHistory.upload_date.desc()
    ).paginate(page=page, per_page=12)""",
    "图5-9 日志与报表页面截图": """# app.py
@app.route('/logs')
@login_required
def logs():
    today_count = hist_query.filter(
        func.date(DetectionHistory.upload_date) == date.today()
    ).count()
    stats_data = [[k, v] for k, v in counter.items()]
    trend_data = []
    for i in range(6, -1, -1):
        target_date = date.today() - timedelta(days=i)
        count = hist_query.filter(
            func.date(DetectionHistory.upload_date) == target_date
        ).count()
        trend_data.append(count)""",
    "图5-10 工作时长统计截图": """# app.py
total_seconds = 0
for l in log_query.filter(WorkLog.logout_time != None).all():
    total_seconds += (l.logout_time - l.login_time).total_seconds()

if session.get('current_log_id'):
    curr = db.session.get(WorkLog, session['current_log_id'])
    if curr:
        total_seconds += (datetime.now() - curr.login_time).total_seconds()

work_time_str = (
    f"{int(total_seconds // 3600)}小时 "
    f"{int((total_seconds % 3600) // 60)}分钟"
)""",
    "图5-11 销售与库存页面截图": """# app.py
@app.route('/sales', methods=['GET', 'POST'])
@login_required
def sales():
    if request.method == 'POST':
        grade = request.form.get('grade')
        qty = int(request.form.get('quantity'))
        inv = db.session.get(Inventory, grade)
        if not inv or inv.count < qty:
            flash(f'{grade} 库存不足！', 'danger')
        else:
            inv.count -= qty
            sale = SalesRecord(customer_name=request.form.get('customer'),
                               product_grade=grade, quantity=qty)
            db.session.add(sale)""",
    "图5-12 用户管理页面截图": """# app.py
@app.route('/admin/users', methods=['GET', 'POST'])
@login_required
def admin_users():
    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'add':
            new_user = User(
                username=request.form.get('username'),
                password=generate_password_hash(request.form.get('password')),
                role=request.form.get('role')
            )
            db.session.add(new_user)
            db.session.commit()""",
    "图5-13 系统设置页面截图": """# app.py
@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    val = request.form.get('conf_threshold')
    s = SystemSettings.query.filter_by(key_name='conf_threshold').first()
    s.value = val

    f = request.files.get('model_file')
    if f and f.filename.endswith('.pt'):
        fname = secure_filename(f.filename)
        f.save(os.path.join(app.root_path, fname))
        m = SystemSettings.query.filter_by(key_name='current_model').first()
        m.value = fname""",
    "图5-14 初始化后系统运行截图": """# app.py
if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['RESULT_FOLDER'], exist_ok=True)
    with app.app_context():
        db.create_all()
        if not User.query.filter_by(username='admin').first():
            db.session.add(User(username='admin',
                                password=generate_password_hash('123456'),
                                role='admin'))
        if not Inventory.query.first():
            db.session.add(Inventory(grade='A级', count=0))
            db.session.add(Inventory(grade='B级', count=0))
            db.session.add(Inventory(grade='C级', count=0))""",
    "图5-15 帮助中心页面截图": """# app.py / templates/help.html
@app.route('/help')
@login_required
def help_center():
    return render_template('help.html')

<img src="{{ url_for('static', filename='img/live_knot.png') }}"
     class="card-img-top" alt="活节样张">
<img src="{{ url_for('static', filename='img/dead_knot.png') }}"
     class="card-img-top" alt="死节样张">
<img src="{{ url_for('static', filename='img/crack.png') }}"
     class="card-img-top" alt="裂纹样张">""",
}


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def format_code_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(0.74)
    paragraph.paragraph_format.right_indent = Cm(0.4)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.name = "Courier New"
        if run._element.rPr is not None and run._element.rPr.rFonts is not None:
            run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        run.font.size = Pt(9)


def format_label_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    for run in paragraph.runs:
        run.font.name = "宋体"
        run.font.size = Pt(12)


def has_existing_block(next_para_text):
    return next_para_text.startswith("部分关键代码如下：") or next_para_text.startswith("部分关键代码如下:")


def main():
    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(DOC_PATH)
    inserted = 0
    paras = doc.paragraphs
    i = 0
    while i < len(paras):
        para = paras[i]
        text = para.text.strip()
        if text in CODE_MAP:
            next_text = paras[i + 1].text.strip() if i + 1 < len(paras) else ""
            if not has_existing_block(next_text):
                label_para = insert_paragraph_after(para, "部分关键代码如下：", para.style)
                format_label_paragraph(label_para)
                code_para = insert_paragraph_after(label_para, CODE_MAP[text], para.style)
                format_code_paragraph(code_para)
                inserted += 1
                paras = doc.paragraphs
        i += 1
    doc.save(DOC_PATH)
    print(inserted)


if __name__ == "__main__":
    main()
