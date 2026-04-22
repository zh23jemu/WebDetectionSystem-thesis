import re
import shutil
from pathlib import Path

from docx import Document


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_按导师目录结构修改前备份.docx")


def set_para_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def heading_level(paragraph):
    name = paragraph.style.name
    if name.startswith("Heading "):
        try:
            return int(name.split()[-1])
        except ValueError:
            return None
    return None


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def delete_section_by_heading(doc, heading_text, max_level=None):
    paras = list(doc.paragraphs)
    start = None
    level = None
    for i, p in enumerate(paras):
        if p.text.strip() == heading_text:
            start = i
            level = heading_level(p)
            break
    if start is None:
        return 0
    if max_level is not None:
        level = max_level

    end = len(paras)
    for j in range(start + 1, len(paras)):
        lv = heading_level(paras[j])
        if lv is not None and lv <= level:
            end = j
            break
    for p in paras[start:end]:
        delete_paragraph(p)
    return end - start


def delete_heading_only(doc, heading_text):
    for p in list(doc.paragraphs):
        if p.text.strip() == heading_text:
            delete_paragraph(p)
            return 1
    return 0


def rename_heading(doc, old, new):
    for p in doc.paragraphs:
        if p.text.strip() == old:
            set_para_text(p, new)
            return True
    return False


def find_heading_index(paras, text):
    for i, p in enumerate(paras):
        if p.text.strip() == text:
            return i
    return None


def get_heading_block(paras, heading_text):
    start = find_heading_index(paras, heading_text)
    if start is None:
        raise ValueError(f"未找到标题：{heading_text}")
    level = heading_level(paras[start])
    end = len(paras)
    for j in range(start + 1, len(paras)):
        lv = heading_level(paras[j])
        if lv is not None and lv <= level:
            end = j
            break
    return paras[start:end]


def as_elements(paragraphs):
    return [p._element for p in paragraphs]


def insert_elements_after(anchor, elements):
    current = anchor
    for element in elements:
        current.addnext(element)
        current = element
    return current


def strip_heading_from_block(block):
    return block[1:] if block else block


def renumber_ch5_figures(doc):
    mapping = {}
    count = 1
    in_ch5 = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "5 系统实现":
            in_ch5 = True
            continue
        if in_ch5 and heading_level(p) == 1:
            break
        m = re.match(r"^(图5-\d+)(\s+.*)$", text)
        if m:
            old = m.group(1)
            new = f"图5-{count}"
            mapping[old] = new
            set_para_text(p, new + m.group(2))
            count += 1

    # 同步正文中的图号引用，避免正文和图题不一致。
    if mapping:
        for p in doc.paragraphs:
            text = p.text
            new_text = text
            for old, new in mapping.items():
                new_text = new_text.replace(old, new)
            if new_text != text:
                set_para_text(p, new_text)
    return count - 1


def rebuild_ch5(doc):
    paras = list(doc.paragraphs)

    h1_5_idx = find_heading_index(paras, "5 系统实现")
    h1_6_idx = find_heading_index(paras, "6 系统测试")
    if h1_5_idx is None or h1_6_idx is None:
        raise ValueError("未找到第5章或第6章标题")

    h2_user = get_heading_block(paras, "5.1 用户功能模块的实现")[0]
    h2_algo = get_heading_block(paras, "5.2 木材缺陷检测模块的实现")[0]
    h2_admin = get_heading_block(paras, "5.4 销售与后台管理模块的实现")[0]

    blocks = {
        "login": get_heading_block(paras, "5.1.1 登录与退出"),
        "dashboard": get_heading_block(paras, "5.1.2 仪表盘与个人中心"),
        "layout": get_heading_block(paras, "5.1.3 公共布局与导航实现"),
        "upload": get_heading_block(paras, "5.2.1 本地上传与摄像头采集"),
        "model": get_heading_block(paras, "5.2.2 模型推理与缺陷映射"),
        "grade": get_heading_block(paras, "5.2.3 等级判定与库存联动"),
        "persist": get_heading_block(paras, "5.2.4 检测结果持久化实现"),
        "history": get_heading_block(paras, "5.3.1 历史记录查询与删除"),
        "logs": get_heading_block(paras, "5.3.2 日志统计与 Excel 导出"),
        "worktime": get_heading_block(paras, "5.3.3 工作时长统计实现"),
        "sales": get_heading_block(paras, "5.4.1 销售与库存管理"),
        "admin_users": get_heading_block(paras, "5.4.2 用户管理"),
        "settings": get_heading_block(paras, "5.4.3 系统设置与模型切换"),
        "init": get_heading_block(paras, "5.4.4 初始化与默认数据实现"),
        "help": get_heading_block(paras, "5.5 缺陷图谱与帮助中心的实现"),
    }

    # 改成导师要求的第5章目录结构。保留截图和代码块，但删除“公共布局与导航”独立目录项。
    set_para_text(h2_user, "5.1 用户功能模块的实现")
    set_para_text(blocks["login"][0], "5.1.1 用户登录界面")
    set_para_text(blocks["history"][0], "5.1.2 检测历史查询界面")
    set_para_text(blocks["logs"][0], "5.1.3 日志报表查看界面")
    set_para_text(blocks["sales"][0], "5.1.4 销售与库存管理界面")
    set_para_text(blocks["help"][0], "5.1.5 缺陷图谱与帮助中心界面")
    blocks["help"][0].style = "Heading 3"

    set_para_text(h2_admin, "5.2 管理员功能模块的实现")
    set_para_text(blocks["admin_users"][0], "5.2.1 用户管理界面")
    set_para_text(blocks["settings"][0], "5.2.2 系统设置与模型切换界面")
    set_para_text(blocks["init"][0], "5.2.3 系统初始化与默认数据实现")

    set_para_text(h2_algo, "5.3 算法检测模块的实现")
    set_para_text(blocks["upload"][0], "5.3.1 数据获取与预处理")
    set_para_text(blocks["model"][0], "5.3.2 检测模型设计与训练")
    set_para_text(blocks["grade"][0], "5.3.3 系统检测与可视化分析")

    # 第5章中仪表盘、公共布局、检测结果持久化、工作时长统计不再作为目录小节出现。
    dashboard_content = strip_heading_from_block(blocks["dashboard"])
    persist_content = strip_heading_from_block(blocks["persist"])
    worktime_content = strip_heading_from_block(blocks["worktime"])

    # 删除第5章原有内容，随后按新结构重新插入已保留的 XML 节点。
    body = doc._body._element
    h1_5_el = paras[h1_5_idx]._element
    for p in paras[h1_5_idx + 1:h1_6_idx]:
        body.remove(p._element)

    anchor = h1_5_el
    sequence = [
        [h2_user._element],
        as_elements(blocks["login"]),
        as_elements(dashboard_content),
        as_elements(blocks["history"]),
        as_elements(blocks["logs"]),
        as_elements(worktime_content),
        as_elements(blocks["sales"]),
        as_elements(blocks["help"]),
        [h2_admin._element],
        as_elements(blocks["admin_users"]),
        as_elements(blocks["settings"]),
        as_elements(blocks["init"]),
        [h2_algo._element],
        as_elements(blocks["upload"]),
        as_elements(blocks["model"]),
        as_elements(blocks["grade"]),
        as_elements(persist_content),
    ]
    for part in sequence:
        anchor = insert_elements_after(anchor, part)


def main():
    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(DOC_PATH)

    # 第2章：按导师意见精简标题，删除前端交互与可视化技术独立小节。
    rename_heading(doc, "2.1 Python 语言技术", "2.1 Python")
    rename_heading(doc, "2.2 Flask 框架", "2.2 Flask")
    delete_section_by_heading(doc, "2.5 前端交互与可视化技术")

    # 第3章：保留可行性、需求、功能、性能与安全，删除程序结构分析。
    delete_section_by_heading(doc, "3.5 程序结构分析")

    # 第4章：取消不必要的 4.1.1、4.1.2；数据库设计调整为 E-R 图和表设计两部分。
    delete_heading_only(doc, "4.1.1 模块划分设计")
    delete_heading_only(doc, "4.1.2 数据流设计")
    rename_heading(doc, "4.3.1 数据库实体设计", "4.3.1 数据库 E-R 图设计")
    delete_heading_only(doc, "4.3.2 数据库 E-R 设计")
    rename_heading(doc, "4.3.3 数据库表设计", "4.3.2 数据库表设计")

    # 第5章：按“用户功能、管理员功能、算法检测模块”重排。
    rebuild_ch5(doc)
    figure_count = renumber_ch5_figures(doc)

    doc.save(DOC_PATH)
    print(f"done, ch5_figures={figure_count}")


if __name__ == "__main__":
    main()
