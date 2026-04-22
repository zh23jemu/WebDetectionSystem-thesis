import shutil
from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Length


SOURCE_DOC = Path(
    r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx"
)
REFERENCE_DOC = Path(
    r"C:\Coding\WebDetectionSystem-thesis\13.张海娇+基于Python的个性化服装推荐系统的设计与实现（6）.docx"
)
OUTPUT_DOC = Path(
    r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现_按参考论文格式重建.docx"
)


def iter_block_items(parent):
    """按文档实际顺序遍历段落和表格。"""
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield child
        elif isinstance(child, CT_Tbl):
            yield child


def paragraph_from_elm(doc, elm):
    for p in doc.paragraphs:
        if p._element == elm:
            return p
    return None


def table_from_elm(doc, elm):
    for t in doc.tables:
        if t._element == elm:
            return t
    return None


def find_title_index(doc, title_text):
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == title_text:
            return i
    raise ValueError(f"未找到标题：{title_text}")


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def clear_target_content_from_title(doc, title_text):
    """从参考论文标题页开始，删除其后所有段落/表格等正文块。"""
    start_para = doc.paragraphs[find_title_index(doc, title_text)]
    body = doc._body._element
    start_remove = False
    for child in list(body.iterchildren()):
        if child == start_para._element:
            start_remove = True
        if start_remove:
            body.remove(child)


def get_source_start_index(doc):
    title = "基于Python的Web木材表面缺陷智能检测系统的设计与实现"
    return find_title_index(doc, title)


def map_style_name(style_name):
    mapping = {
        "Heading 1": "标题01",
        "Heading 2": "标题02",
        "Heading 3": "标题03",
        "Heading 4": "标题03",
        "Heading 5": "标题03",
        "Normal": "Normal",
    }
    return mapping.get(style_name, style_name)


def copy_paragraph_format(src_para, dst_para):
    fmt = src_para.paragraph_format
    dst_fmt = dst_para.paragraph_format
    dst_fmt.alignment = src_para.alignment
    dst_fmt.left_indent = fmt.left_indent
    dst_fmt.right_indent = fmt.right_indent
    dst_fmt.first_line_indent = fmt.first_line_indent
    dst_fmt.space_before = fmt.space_before
    dst_fmt.space_after = fmt.space_after
    dst_fmt.line_spacing = fmt.line_spacing
    dst_fmt.line_spacing_rule = fmt.line_spacing_rule
    dst_fmt.keep_together = fmt.keep_together
    dst_fmt.keep_with_next = fmt.keep_with_next
    dst_fmt.page_break_before = fmt.page_break_before
    dst_fmt.widow_control = fmt.widow_control


def copy_run_format(src_run, dst_run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.size = src_run.font.size
    dst_run.font.name = src_run.font.name
    if src_run._element.rPr is not None:
        rfonts = src_run._element.rPr.rFonts
        if rfonts is not None and rfonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"):
            east = rfonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")
            dst_run._element.get_or_add_rPr().rFonts.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", east
            )


def extract_images_from_paragraph(paragraph):
    images = []
    drawing_runs = paragraph._element.xpath(".//a:blip")
    for blip in drawing_runs:
        rid = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if rid and rid in paragraph.part.related_parts:
            image_part = paragraph.part.related_parts[rid]
            images.append(image_part.blob)
    return images


def append_paragraph(src_para, target_doc):
    target_para = target_doc.add_paragraph()
    style_name = map_style_name(src_para.style.name)
    if style_name in [s.name for s in target_doc.styles]:
        target_para.style = target_doc.styles[style_name]
    copy_paragraph_format(src_para, target_para)

    if src_para.text:
        for src_run in src_para.runs:
            dst_run = target_para.add_run(src_run.text)
            copy_run_format(src_run, dst_run)

    images = extract_images_from_paragraph(src_para)
    if images:
        target_para.alignment = src_para.alignment
        for blob in images:
            run = target_para.add_run()
            run.add_picture(BytesIO(blob))

    return target_para


def append_table(src_table, target_doc):
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    new_table = target_doc.add_table(rows=rows, cols=cols)
    try:
        if src_table.style and src_table.style.name in [s.name for s in target_doc.styles]:
            new_table.style = src_table.style.name
    except Exception:
        pass

    for r_idx, row in enumerate(src_table.rows):
        for c_idx, cell in enumerate(row.cells):
            new_cell = new_table.cell(r_idx, c_idx)
            new_cell.text = ""
            if cell.paragraphs:
                new_cell.paragraphs[0].text = cell.paragraphs[0].text
            for extra_idx, para in enumerate(cell.paragraphs[1:], start=1):
                if extra_idx == 1 and new_cell.paragraphs[0].text == "":
                    new_cell.paragraphs[0].text = para.text
                else:
                    new_cell.add_paragraph(para.text)

    return new_table


def rebuild():
    shutil.copy2(REFERENCE_DOC, OUTPUT_DOC)

    src_doc = Document(str(SOURCE_DOC))
    target_doc = Document(str(OUTPUT_DOC))

    ref_title = "基于Python的个性化服装推荐系统的设计与实现"
    clear_target_content_from_title(target_doc, ref_title)

    source_start = get_source_start_index(src_doc)
    source_block_elms = list(iter_block_items(src_doc))

    start_elm = src_doc.paragraphs[source_start]._element
    start_copy = False

    for elm in source_block_elms:
        if elm == start_elm:
            start_copy = True
        if not start_copy:
            continue

        para = paragraph_from_elm(src_doc, elm)
        if para is not None:
            append_paragraph(para, target_doc)
            continue

        table = table_from_elm(src_doc, elm)
        if table is not None:
            append_table(table, target_doc)

    target_doc.save(str(OUTPUT_DOC))
    print("done")


if __name__ == "__main__":
    rebuild()
