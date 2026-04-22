from pathlib import Path
import re

from docx import Document


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")


def heading_level(paragraph):
    name = paragraph.style.name
    if name.startswith("Heading "):
        try:
            return int(name.split()[-1])
        except ValueError:
            return None
    return None


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def set_para_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def remove_section(doc, heading_text):
    paras = list(doc.paragraphs)
    start = None
    level = None
    for i, p in enumerate(paras):
        if p.text.strip() == heading_text:
            start = i
            level = heading_level(p)
            break
    if start is None:
        return 0
    end = len(paras)
    for j in range(start + 1, len(paras)):
        lv = heading_level(paras[j])
        if lv is not None and lv <= level:
            end = j
            break
    for p in paras[start:end]:
        delete_paragraph(p)
    return end - start


def renumber_figures(doc):
    count = 1
    in_ch5 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "5 系统实现":
            in_ch5 = True
            continue
        if in_ch5 and t == "6 系统测试":
            break
        m = re.match(r"^图5-\d+(\s+.*)$", t)
        if in_ch5 and m:
            set_para_text(p, f"图5-{count}{m.group(1)}")
            count += 1


def main():
    doc = Document(DOC_PATH)
    removed = remove_section(doc, "5.2.3 系统初始化与默认数据实现")
    renumber_figures(doc)
    doc.save(DOC_PATH)
    print(removed)


if __name__ == "__main__":
    main()
