from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")
IMG_DIR = Path(r"C:\Coding\WebDetectionSystem-thesis\paper_diagrams\usecase_generated")
USER_IMG = IMG_DIR / "图3-1_普通用户用例图.png"
ADMIN_IMG = IMG_DIR / "图3-2_管理员用例图.png"


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def format_body(para):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.first_line_indent = Pt(24)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


def format_caption(para):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)


def add_image_after(paragraph, img_path, width_cm=13.8):
    p = insert_paragraph_after(paragraph)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    return p


def main():
    doc = Document(DOC_PATH)
    paras = doc.paragraphs
    anchor = None
    for p in paras:
        if p.text.strip() == "若按照功能入口进一步细分，系统主要包括登录、仪表盘、缺陷检测、历史记录、日志报表、数据导出、销售管理、帮助中心、用户管理和系统设置等功能页面。各功能页面与后端业务处理模块相互对应，形成了较为清晰的页面组织结构，从而便于系统维护与功能扩展。":
            anchor = p
            break

    if anchor is None:
        raise ValueError("未找到插入锚点")

    # 避免重复插入
    existing = {p.text.strip() for p in doc.paragraphs}
    if "图3-1 普通用户用例图" in existing or "图3-2 管理员用例图" in existing:
        print("already_exists")
        return

    p1 = insert_paragraph_after(anchor, "从用户角色角度分析，普通用户主要参与登录认证、缺陷检测、历史记录查询、日志报表查看、销售管理、帮助中心浏览和个人中心维护等功能。普通用户用例图如图3-1所示。")
    format_body(p1)
    img1 = add_image_after(p1, USER_IMG)
    cap1 = insert_paragraph_after(img1, "图3-1 普通用户用例图")
    format_caption(cap1)

    p2 = insert_paragraph_after(cap1, "从管理员角色角度分析，管理员在普通用户功能基础上还具备用户管理、系统设置、模型切换以及日志查看等管理权限。管理员用例图如图3-2所示。")
    format_body(p2)
    img2 = add_image_after(p2, ADMIN_IMG)
    cap2 = insert_paragraph_after(img2, "图3-2 管理员用例图")
    format_caption(cap2)

    doc.save(DOC_PATH)
    print("inserted")


if __name__ == "__main__":
    main()
