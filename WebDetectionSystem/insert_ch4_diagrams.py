import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


DOC_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx")
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_插入第四章图片前备份.docx")
IMG_DIR = Path(r"C:\Coding\WebDetectionSystem-thesis\paper_diagrams\reference_style_generated")

FIGURES = {
    "图4-1 系统总体结构图": IMG_DIR / "图4-1_系统总体结构图_参考风格.png",
    "图4-2 登录顺序图": IMG_DIR / "图4-2_登录顺序图_参考风格.png",
    "图4-3 缺陷检测顺序图": IMG_DIR / "图4-3_缺陷检测顺序图_参考风格.png",
    "图4-4 历史与报表顺序图": IMG_DIR / "图4-4_历史与报表顺序图_参考风格.png",
    "图4-5 销售出库顺序图": IMG_DIR / "图4-5_销售出库顺序图_参考风格.png",
    "图4-6 主要实体属性图": IMG_DIR / "图4-6_主要实体属性图_参考风格.png",
    "图4-7 系统 E-R 图": IMG_DIR / "图4-7_系统E-R图_参考风格.png",
}


def create_empty_before(paragraph):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def normalize_caption_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.5


def insert_images():
    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(DOC_PATH)

    inserted = 0
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if text not in FIGURES:
            continue

        img_path = FIGURES[text]
        if not img_path.exists():
            continue

        prev_para = para._p.getprevious()
        img_para = None
        if prev_para is not None and "graphic" in prev_para.xml:
            from docx.text.paragraph import Paragraph
            img_para = Paragraph(prev_para, para._parent)
            for run in img_para.runs:
                run.text = ""
            # remove any existing drawings in the paragraph xml
            for child in list(img_para._p):
                tag = child.tag.lower()
                if tag.endswith("r"):
                    img_para._p.remove(child)
        else:
            img_para = create_empty_before(para)

        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_fmt = img_para.paragraph_format
        img_fmt.space_before = Pt(6)
        img_fmt.space_after = Pt(3)
        img_fmt.line_spacing = 1.0
        run = img_para.add_run()
        run.add_picture(str(img_path), width=Cm(15.5))

        normalize_caption_paragraph(para)
        inserted += 1

    doc.save(DOC_PATH)
    return inserted


if __name__ == "__main__":
    print(insert_images())
