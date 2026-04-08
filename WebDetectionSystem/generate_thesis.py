from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx"
REFERENCE = ROOT.parent / "13.张海娇+基于Python的个性化服装推荐系统的设计与实现（6）.docx"


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=None, bold=False):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_text(paragraph, text, *, east_asia="宋体", latin="Times New Roman", size=12, bold=False):
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, latin=latin, size=size, bold=bold)
    return run


def add_paragraph(document, text, *, style=None, align=None, first_line_chars=2, line_spacing=1.5, font_size=12):
    p = document.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing = line_spacing
    if first_line_chars:
        fmt.first_line_indent = Pt(font_size * first_line_chars)
    set_paragraph_text(p, text, size=font_size)
    return p


def add_heading(document, text, level):
    p = document.add_paragraph(style=f"Heading {level}")
    set_paragraph_text(p, text, east_asia="黑体", size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_center_note(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    set_paragraph_text(p, text, east_asia="宋体", size=11)
    return p


def add_toc(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录将在 Word 中更新域后自动生成"
    fld_sep.append(text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_begin)
    run = paragraph.add_run()
    run._r.append(instr)
    run = paragraph.add_run()
    run._r.append(fld_sep)
    run = paragraph.add_run()
    run._r.append(fld_end)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_begin)
    run = paragraph.add_run()
    run._r.append(instr)
    run = paragraph.add_run()
    run._r.append(fld_end)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.3
    set_paragraph_text(p, text, size=11, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def clone_reference_cover(doc):
    ref = Document(REFERENCE)
    src_tables = ref.tables[:1]
    if not src_tables:
        return False
    tbl = deepcopy(src_tables[0]._tbl)
    doc._body._element.append(tbl)
    return True


def create_cover(doc):
    if clone_reference_cover(doc):
        table = doc.tables[0]
        replacement_map = {
            "基于Python的个性化服装推荐系统的": "基于Python的Web木材表面缺陷智能检测系统的",
            "设计与实现": "设计与实现",
            "张海娇": "待填写",
            "231053011717": "待填写",
            "计算机与信息工程学院": "待填写",
            "计算机科学与技术": "待填写",
            "黄浩英": "待填写",
            "讲师": "待填写",
        }
        for row in table.rows:
            for cell in row.cells:
                for old, new in replacement_map.items():
                    if old in cell.text:
                        for p in cell.paragraphs:
                            if old in p.text:
                                text = p.text.replace(old, new)
                                p.clear()
                                set_paragraph_text(p, text, east_asia="方正姚体" if "本科毕业设计" in text or "设计与实现" in text else "宋体", size=16)
        doc.add_page_break()
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    set_paragraph_text(p, "本科毕业设计（论文）", east_asia="方正姚体", size=28, bold=True)
    info = [
        "",
        "题    目：基于Python的Web木材表面缺陷智能检测系统的设计与实现",
        "姓    名：待填写",
        "学    号：待填写",
        "二级学院：待填写",
        "专    业：待填写",
        "指导教师：待填写",
        "职    称：待填写",
    ]
    for line in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        set_paragraph_text(p, line, east_asia="宋体", size=14, bold="题    目" in line)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    set_paragraph_text(p, "教务处制", east_asia="方正姚体", size=16)
    doc.add_page_break()


def add_reference_section(doc):
    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Python Software Foundation. Python Documentation[EB/OL].",
        "[2] Flask Documentation Team. Flask Documentation[EB/OL].",
        "[3] SQLAlchemy Authors. SQLAlchemy Documentation[EB/OL].",
        "[4] Ultralytics. Ultralytics YOLO Docs[EB/OL].",
        "[5] OpenCV Team. OpenCV-Python Tutorials[EB/OL].",
        "[6] Bootstrap Team. Bootstrap 5 Documentation[EB/OL].",
        "[7] pandas development team. pandas Documentation[EB/OL].",
        "[8] He K, Gkioxari G, Dollar P, et al. Mask R-CNN[C]. IEEE, 2017.",
        "[9] Redmon J, Farhadi A. YOLOv3: An Incremental Improvement[EB/OL].",
        "[10] Jocher G, Chaurasia A, Qiu J. YOLO by Ultralytics[EB/OL].",
        "[11] 王耀南, 刘少创. 机器视觉原理与应用[M]. 北京: 机械工业出版社.",
        "[12] 周志华. 机器学习[M]. 北京: 清华大学出版社.",
    ]
    for ref in refs:
        add_paragraph(doc, ref, first_line_chars=0, font_size=11)


def add_acknowledgement(doc):
    add_heading(doc, "致谢", 1)
    text = (
        "本论文能够顺利完成，离不开指导教师在选题、系统设计、论文结构组织与修改完善过程中给予的耐心帮助。"
        "在 Web 木材表面缺陷智能检测系统的开发中，从需求梳理、页面实现、数据库模型设计，到缺陷检测流程与业务联动的落地，"
        "均得到了老师的悉心指导。与此同时，也感谢学院提供的实验环境与学习平台，使本人能够在课程学习与毕业设计阶段将 Python Web 开发、"
        "数据库设计和目标检测算法应用到具体工程场景中。还要感谢同学和家人在系统调试、资料查阅、论文撰写期间给予的鼓励与支持。"
        "后续本人将继续完善项目中的模型泛化能力、系统部署方式以及多端协同能力，努力将本次毕业设计积累的经验转化为更扎实的工程实践能力。"
    )
    add_paragraph(doc, text)


def add_database_tables(doc):
    add_heading(doc, "4.3.3 数据库表设计", 3)
    add_paragraph(doc, "根据 `models.py` 中的实体定义，系统主要包含用户、检测历史、工作日志、销售记录、系统设置和库存六张核心业务表。")
    tables = [
        ("表4-1 用户表 User", [("字段名", "类型", "说明"), ("id", "Integer", "主键"), ("username", "String(150)", "用户名，唯一"), ("password", "String(150)", "密码哈希"), ("role", "String(20)", "角色，user/admin")]),
        ("表4-2 检测历史表 DetectionHistory", [("字段名", "类型", "说明"), ("id", "Integer", "主键"), ("filename", "String(300)", "原始文件名"), ("result_image", "String(300)", "结果图文件名"), ("upload_date", "DateTime", "上传检测时间"), ("defect_type", "String(100)", "缺陷类型"), ("confidence", "Float", "最高置信度"), ("user_id", "Integer", "关联用户 ID"), ("grade", "String(10)", "等级结果")]),
        ("表4-3 工作日志表 WorkLog", [("字段名", "类型", "说明"), ("id", "Integer", "主键"), ("user_id", "Integer", "操作用户 ID"), ("login_time", "DateTime", "登录时间"), ("logout_time", "DateTime", "退出时间")]),
        ("表4-4 销售记录表 SalesRecord", [("字段名", "类型", "说明"), ("id", "Integer", "主键"), ("customer_name", "String(100)", "客户名称"), ("product_grade", "String(50)", "销售等级"), ("quantity", "Integer", "数量"), ("total_price", "Float", "订单总价"), ("sale_date", "DateTime", "销售时间")]),
        ("表4-5 系统设置表 SystemSettings", [("字段名", "类型", "说明"), ("id", "Integer", "主键"), ("key_name", "String(50)", "设置键名"), ("value", "String(200)", "设置值")]),
        ("表4-6 库存表 Inventory", [("字段名", "类型", "说明"), ("grade", "String(10)", "等级主键"), ("count", "Integer", "库存数量")]),
    ]
    for title, rows in tables:
        add_center_note(doc, title)
        table = doc.add_table(rows=len(rows), cols=3)
        table.style = "Table Grid"
        for r, row in enumerate(rows):
            for c, value in enumerate(row):
                set_cell_text(table.cell(r, c), value, bold=(r == 0))
        doc.add_paragraph()


def build_document():
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    create_cover(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(p, "目 录", east_asia="黑体", size=16, bold=True)
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.line_spacing = 1.5
    add_toc(toc_p)
    doc.add_page_break()

    title = "基于Python的Web木材表面缺陷智能检测系统的设计与实现"
    add_center_note(doc, title)
    add_center_note(doc, "计算机科学与技术专业  待填写")
    add_center_note(doc, "指导教师  待填写")
    add_paragraph(
        doc,
        "摘要：木材表面缺陷检测是木材加工、质量分级和销售流转中的关键环节。传统人工检测方式虽然直观，但长期存在主观性强、效率受限、记录难以追溯等问题。"
        "结合当前项目 `WebDetectionSystem` 的真实实现情况，本文设计并整理了一套基于 Python 的 Web 木材表面缺陷智能检测系统。系统以后端 Flask 框架为基础，"
        "以 Ultralytics YOLO 模型为核心检测引擎，集成用户登录、图片上传检测、摄像头采集、缺陷中文映射、等级判定、检测历史追溯、日志统计分析、"
        "Excel 数据导出、库存联动、销售出库、知识图谱展示与管理员模型切换等功能。系统识别的缺陷类别主要包括裂纹、死节和活节，并依据业务规则自动划分为 A 级、B 级和 C 级板材。"
        "在工程实现层面，系统利用 SQLAlchemy 完成数据建模，采用 Jinja2 模板、Bootstrap 5、Chart.js 和 Viewer.js 构建交互界面，实现了检测业务与库存销售业务的一体化联动。"
        "论文从系统背景、关键技术、需求分析、总体设计、功能实现和测试验证等方面，对当前系统进行完整梳理。实践表明，该系统能够满足课程设计和教学演示场景下的基本需求，"
        "具备检测流程清晰、操作门槛较低、结果可追溯、业务闭环完整等特点，对木材缺陷检测系统的教学研究和原型验证具有一定参考价值。"
    )
    add_paragraph(doc, "关键词：Python；木材缺陷检测；YOLO；Flask；Web 系统；库存销售联动", first_line_chars=0)
    doc.add_paragraph()
    add_center_note(doc, "Design and Implementation of a Python-based Web Wood Surface Defect Detection System")
    add_center_note(doc, "Computer Science and Technology  To be filled")
    add_center_note(doc, "Tutor  To be filled")
    add_paragraph(
        doc,
        "Abstract: Wood surface defect inspection is a critical procedure in timber processing, quality grading and sales circulation. Conventional manual inspection is intuitive, "
        "but it often suffers from unstable standards, low efficiency and poor traceability. Based on the current implementation of the WebDetectionSystem project, this paper presents "
        "a Python-based web wood surface defect detection system. The system adopts Flask as the backend framework and Ultralytics YOLO as the core detection engine, and integrates user login, "
        "image upload detection, camera capture, Chinese defect mapping, automatic grading, history tracing, statistical reporting, Excel export, inventory linkage, sales management, help center "
        "and administrator-side model switching. The system mainly recognizes Crack, Dead Knot and Live Knot, and maps them to grade A, grade B and grade C according to business rules. "
        "On the engineering side, SQLAlchemy is used for data modeling, while Jinja2 templates, Bootstrap 5, Chart.js and Viewer.js are used to build the front-end pages. "
        "This paper discusses the project from the aspects of research background, key technologies, requirement analysis, system design, implementation and testing. "
        "The practice shows that the system can satisfy the basic needs of course design and teaching demonstration, and has practical value in process standardization, visual inspection and business loop integration."
    )
    add_paragraph(doc, "Keywords: Python; Wood Defect Detection; YOLO; Flask; Web System; Inventory and Sales", first_line_chars=0)
    doc.add_page_break()

    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    add_paragraph(
        doc,
        "随着木材加工、家具制造、建筑装饰以及板材贸易等行业不断发展，木材表面质量检测的重要性日益凸显。裂纹、死节和活节等缺陷会直接影响木材的结构强度、"
        "外观等级以及销售价格。在传统生产流程中，木材缺陷识别大多依赖质检人员人工肉眼观察，再结合经验完成等级划分。这种方式虽然可以在小批量场景中使用，"
        "但在连续作业、批量分拣和多人协作时容易受到疲劳、经验差异和现场环境的影响，导致检测标准难以统一，管理数据难以沉淀。"
    )
    add_paragraph(
        doc,
        "近年来，深度学习和计算机视觉技术在工业检测中的应用逐渐成熟，目标检测算法能够从图像中快速定位缺陷区域并输出类别结果，为传统人工检测模式提供了新的升级路径。"
        "尤其是 YOLO 系列算法在推理速度和部署便捷性方面具备明显优势，适合构建轻量级的教学演示系统和企业原型系统。与此同时，仅有检测模型并不能完全满足业务需求，"
        "还需要配套登录权限、检测记录、统计报表、库存联动与销售管理等信息化功能，才能形成真正可用的 Web 业务闭环。"
    )
    add_paragraph(
        doc,
        "本项目的 `WebDetectionSystem` 正是在上述需求背景下构建的一个木材表面缺陷智能检测系统。系统不仅完成图像识别任务，还将检测结果映射到等级判定、历史追溯、日志统计、库存入库和销售出库流程中，"
        "体现出“检测算法 + Web 平台 + 业务管理”的综合应用特征。因此，围绕该系统整理毕业论文，既能够总结项目实现经验，也有助于展示 Python Web 开发与智能视觉算法结合的完整工程过程。"
    )

    add_heading(doc, "1.2 研究现状", 2)
    add_heading(doc, "1.2.1 国外研究现状", 3)
    add_paragraph(
        doc,
        "国外在木材缺陷检测和工业视觉识别领域起步较早，研究重点主要集中在纹理分析、特征提取、图像分割、传统机器学习分类以及深度神经网络识别等方向。"
        "早期研究多采用灰度共生矩阵、边缘算子和手工纹理特征描述木材表面，再通过支持向量机、决策树等方法完成缺陷分类。随着卷积神经网络的发展，"
        "目标检测框架逐步替代传统特征工程，在实际应用中显著提升了检测效率和鲁棒性。"
    )
    add_paragraph(
        doc,
        "在工程应用层面，国外木材、钢材、纺织和电子制造行业较早尝试将机器视觉系统嵌入生产线，实现缺陷自动报警、在线分级和品质统计。"
        "尤其是单阶段检测算法在速度方面的优势，使其更加适合实时质检和边缘部署场景。研究趋势已从单纯提高识别准确率，逐渐扩展到可解释检测、跨设备泛化、小样本训练和端到端业务协同。"
    )
    add_heading(doc, "1.2.2 国内研究现状", 3)
    add_paragraph(
        doc,
        "国内在木材缺陷识别、木材质量分级以及智能制造系统建设方面同样发展迅速。依托 Python 生态、深度学习框架和开源视觉库，越来越多的研究工作将目标检测算法应用于木材节子、裂纹、虫孔等表面缺陷识别。"
        "不少研究已经实现从数据集采集、模型训练到可视化演示平台搭建的完整流程，但在系统化工程落地方面，仍存在重模型、轻业务的现象，即检测结果与实际管理流程缺少深度耦合。"
    )
    add_paragraph(
        doc,
        "在课程设计和毕业设计实践中，基于 Flask 或 Django 的智能检测系统数量逐年增加，项目往往能够完成图像上传、模型推理、结果显示等基础功能。"
        "然而若要满足真实业务需求，还需要进一步解决检测结果长期存储、操作行为追踪、检测报表生成、库存数量变化和多角色权限划分等问题。"
        "因此，将检测算法与 Web 后台管理、数据库设计和前端交互结合起来，构建一个适配教学与演示场景的完整系统，具有较强的实践价值。"
    )
    add_heading(doc, "1.3 论文组织结构", 2)
    add_paragraph(doc, "第1章为绪论，主要介绍木材表面缺陷检测系统的研究背景、国内外研究现状以及本文的整体组织结构。")
    add_paragraph(doc, "第2章为关键技术，围绕项目中实际采用的 Python、Flask、YOLO、SQLite 与前端可视化技术进行说明。")
    add_paragraph(doc, "第3章为系统分析，结合当前代码结构和业务功能，对系统可行性、功能需求、性能需求和角色需求进行梳理。")
    add_paragraph(doc, "第4章为系统设计，对系统架构、核心业务流程以及数据库表结构进行设计说明，并为后续实现与测试提供理论依据。")
    add_paragraph(doc, "第5章为系统实现，按照登录、检测、报表、库存销售、后台设置和帮助中心等模块展开描述，同时预留后期系统截图插入位置。")
    add_paragraph(doc, "第6章为系统测试，对系统测试目标、环境、典型测试用例与测试结论进行总结。")
    add_paragraph(doc, "第7章为总结与展望，对项目建设成果、存在不足以及未来优化方向进行归纳。")

    add_heading(doc, "2 关键技术", 1)
    add_heading(doc, "2.1 Python 语言技术", 2)
    add_paragraph(
        doc,
        "Python 语法简洁、生态丰富、开发效率高，广泛应用于 Web 开发、数据分析、人工智能和自动化运维等领域。对于本系统而言，Python 一方面负责承载 Flask Web 业务逻辑，"
        "另一方面又能够无缝衔接 YOLO 模型推理、OpenCV 图像处理、pandas 数据导出等功能，因此非常适合构建中小型智能检测平台。"
    )
    add_paragraph(
        doc,
        "从项目代码来看，`app.py` 统一处理路由、表单提交、数据库交互和 JSON 响应；`utils/detection.py` 封装模型加载与图片推理流程；`models.py` 定义数据库实体；"
        "`config.py` 管理路径与数据库配置。上述模块均使用 Python 实现，体现了语言在快速原型开发和功能集成方面的优势。"
    )
    add_heading(doc, "2.2 Flask 框架", 2)
    add_paragraph(
        doc,
        "Flask 是一个轻量级 Web 框架，具备路由清晰、扩展灵活、上手成本低等特点。系统采用 Flask 作为后端框架，将登录、检测、历史记录、日志报表、销售、后台设置等业务路由组织在同一入口文件中。"
        "通过 `render_template` 渲染前端页面，通过 `request` 接收表单和文件上传，通过 `jsonify` 返回异步检测结果，通过 `send_file` 完成 Excel 数据导出，实现了较完整的 Web 应用闭环。"
    )
    add_paragraph(
        doc,
        "此外，系统还使用 Flask-Login 完成用户身份认证和会话控制。用户登录后可访问仪表盘、检测、历史记录、日志报表、销售管理等页面；管理员在此基础上额外拥有用户管理和系统设置权限。"
        "这种基于扩展组件构建的方式，既保持了 Flask 本身的简洁性，也便于后期继续扩展。"
    )
    add_heading(doc, "2.3 YOLO 目标检测算法", 2)
    add_paragraph(
        doc,
        "YOLO（You Only Look Once）是一种典型的单阶段目标检测算法，其核心思想是将目标定位与分类统一到一次前向推理中完成，兼顾较快推理速度与较高检测精度。"
        "本系统通过 Ultralytics 提供的 Python 接口加载 `.pt` 权重文件，在 `process_image` 函数中完成图像读取、模型推理、检测框解析、最大置信度计算和结果图保存。"
    )
    add_paragraph(
        doc,
        "在业务层面，系统并不直接将模型输出作为最终结果，而是进一步完成英文类别到中文缺陷名称的映射，并通过判定规则将检测结果转换为 A 级、B 级或 C 级板材等级。"
        "这表明项目对检测算法结果进行了业务化包装，使模型输出更适合现场使用者理解。"
    )
    add_heading(doc, "2.4 SQLite 与 SQLAlchemy", 2)
    add_paragraph(
        doc,
        "SQLite 是轻量级关系型数据库，部署方便、无需单独数据库服务进程，适合课程设计、单机原型和教学演示场景。本项目默认数据库文件为 `wood_system.db`，使用 SQLAlchemy 作为 ORM 层完成实体建模与数据读写。"
        "通过对象映射方式，系统可以较为直观地处理用户、检测记录、工作日志、销售记录、系统设置和库存等业务数据，降低了直接书写 SQL 的复杂度。"
    )
    add_paragraph(
        doc,
        "由于检测历史、报表统计与销售管理都依赖结构化数据支撑，数据库不仅承担存储功能，还为后续分页查询、趋势统计、数据导出和业务联动提供基础。"
        "从系统的使用场景看，SQLite 已能满足本阶段需求，若后期面向多人并发访问，可进一步迁移至 MySQL 或 PostgreSQL。"
    )
    add_heading(doc, "2.5 前端交互与可视化技术", 2)
    add_paragraph(
        doc,
        "系统前端主要采用 Jinja2 模板 + Bootstrap 5 方式实现页面构建，辅以 Font Awesome 图标库优化可视化表达。对历史结果大图浏览使用 Viewer.js 实现，对日志与报表中的趋势图和缺陷占比图使用 Chart.js 进行图形绘制。"
        "检测页面通过原生 JavaScript 调用 `fetch` 提交表单和摄像头采集结果，使前端能够异步刷新检测卡片并展示结果图、缺陷标签和等级判定。"
    )
    add_paragraph(
        doc,
        "这种技术组合实现简单、学习成本低，且与 Flask 模板渲染方式兼容良好，适合毕业设计场景下快速搭建完整的可视化页面体系。"
    )

    add_heading(doc, "3 系统分析", 1)
    add_heading(doc, "3.1 可行性研究", 2)
    add_heading(doc, "3.1.1 经济可行性", 3)
    add_paragraph(
        doc,
        "系统开发过程中主要使用 Python、Flask、Ultralytics、OpenCV、Bootstrap、Chart.js 等开源技术，能够显著降低软件采购成本。硬件方面，演示环境仅需普通开发电脑即可完成模型加载与 Web 服务运行，"
        "无需部署高成本工业控制平台。对于课程设计和中小规模原型验证而言，其投入成本较低，具有较好的经济可行性。"
    )
    add_heading(doc, "3.1.2 操作可行性", 3)
    add_paragraph(
        doc,
        "从页面交互看，系统采用统一侧边栏导航组织功能，用户可以在仪表盘、开始检测、历史记录、日志报表、销售管理和帮助中心之间快速切换。检测页面同时支持本地图片上传与摄像头采集，"
        "管理页面使用表单式交互完成用户添加、阈值设置和模型切换，整体操作逻辑直观。即便是非技术背景的木材检测人员，也可以在简单培训后完成基本使用。"
    )
    add_heading(doc, "3.1.3 技术可行性", 3)
    add_paragraph(
        doc,
        "本系统所使用的关键技术均较为成熟。后端 Web 能力由 Flask 提供，数据库能力由 SQLite 与 SQLAlchemy 提供，目标检测能力由 Ultralytics YOLO 提供，图片读写能力由 OpenCV 提供，"
        "前端页面和图表展示能力由 Bootstrap 5、Chart.js 和 Viewer.js 提供。现有项目已经实现代码级联调，说明技术路线是可行的。"
    )
    add_paragraph(
        doc,
        "同时，系统代码结构相对清晰，核心逻辑集中在 `app.py` 与 `utils/detection.py` 两部分，便于后期进行模型替换、业务扩展和部署优化。因此，从开发技术、学习成本和维护角度看，该系统具备良好的技术可行性。"
    )
    add_heading(doc, "3.2 需求分析", 2)
    add_paragraph(
        doc,
        "根据当前程序实现情况，系统面向普通用户和管理员两类角色。普通用户主要完成登录、图片检测、历史查看、日志统计、库存销售查看和个人密码修改等操作；管理员除拥有普通用户功能外，"
        "还承担用户维护、参数配置和模型切换任务。系统需求既包括图像检测类需求，也包括管理类需求和辅助学习类需求。"
    )
    add_paragraph(
        doc,
        "在检测业务上，系统需要支持单张与多张图片检测，支持浏览器摄像头采集，支持输出缺陷类别、置信度和板材等级，并将结果图保存在服务器指定目录。在业务管理上，系统需要记录每次检测的文件名、时间、操作人、"
        "结果图路径和等级信息，能够按照缺陷类型进行搜索和分页查询。在报表分析上，系统需要统计今日检测量、累计工作时长、缺陷类别占比和近 7 天检测趋势，并支持导出 Excel。"
    )
    add_paragraph(
        doc,
        "在流程联动方面，系统还需要在检测完成后自动更新对应等级库存，在销售页面创建订单并自动扣减库存，从而形成从检测到入库再到出库的闭环。帮助中心则用于展示常见缺陷图谱和等级判定规则，满足新用户学习需求。"
    )
    add_heading(doc, "3.3 系统功能分析", 2)
    add_paragraph(
        doc,
        "结合当前项目目录结构，系统功能可划分为界面展示层、业务逻辑层、模型推理层和数据存储层四个部分。`templates/` 目录负责页面展示，`app.py` 负责路由与业务控制，"
        "`utils/detection.py` 负责模型推理，`models.py` 与数据库文件负责数据持久化。"
    )
    add_paragraph(
        doc,
        "从角色角度分析，普通用户的核心功能包括：登录认证、木材缺陷检测、历史记录查询、日志报表查看、销售信息查看、帮助中心学习和个人密码修改。管理员在此基础上新增用户管理、系统参数调节和模型文件切换功能。"
        "其中，检测模块是系统核心；历史、报表和销售模块体现系统的业务闭环；帮助中心和后台配置则增强了系统的实用性和可维护性。"
    )
    add_paragraph(
        doc,
        "若按照程序入口进一步细分，系统主要包含 `/login` 登录、`/dashboard` 仪表盘、`/detect` 检测、`/history` 历史记录、`/logs` 日志报表、`/export_data` 数据导出、"
        "`/sales` 销售管理、`/help` 帮助中心、`/admin/users` 用户管理和 `/settings` 系统设置等路由。各路由分别对应 `templates` 目录中的页面模板，"
        "形成了“一个路由对应一个业务页面”的清晰结构，这也是当前系统便于阅读和快速维护的重要原因。"
    )
    add_heading(doc, "3.4 性能与安全需求分析", 2)
    add_paragraph(
        doc,
        "在性能方面，系统需要保证单张图片上传后能够在可接受时间内返回结果，批量图片检测时页面仍能保持基本响应；对历史记录查询、图表统计和 Excel 导出等常用功能，应能够在教学演示环境中稳定运行。"
        "在安全方面，系统通过 Flask-Login 实现登录保护，通过密码哈希方式保存用户凭证，通过管理员角色判断限制后台管理入口，初步满足课程设计场景下的访问控制要求。"
    )
    add_heading(doc, "3.5 程序结构分析", 2)
    add_paragraph(
        doc,
        "结合项目目录遍历结果可以看出，程序结构围绕“单应用入口 + 多模板页面 + 工具模块 + 静态资源”展开。根目录中的 `app.py` 集中定义所有路由与业务流程；"
        "`config.py` 统一管理数据库、上传目录、结果目录和模型路径；`extensions.py` 对 SQLAlchemy 与 LoginManager 进行初始化封装；`models.py` 负责实体定义。"
    )
    add_paragraph(
        doc,
        "在工具层，`utils/detection.py` 负责模型缓存、图像推理和结果图输出，`utils/img_process.py` 预留了图像尺寸调整和格式转换能力，虽然当前主流程尚未显式调用该模块，"
        "但其存在说明系统已考虑到图片预处理扩展需求。前端模板层则以 `base.html` 为公共骨架，其他页面通过继承方式复用侧边栏、顶部提示和静态资源引用，从而减少重复代码。"
    )
    add_paragraph(
        doc,
        "在静态资源方面，`static/img` 存放帮助中心所需的缺陷样张，`static/uploads` 存放原始上传图片，`static/results` 存放检测后的结果图。这种目录规划与业务流程一一对应，"
        "既便于调试，也便于在论文中说明系统的数据流转路径。"
    )

    add_heading(doc, "4 系统设计", 1)
    add_heading(doc, "4.1 系统结构设计", 2)
    add_paragraph(
        doc,
        "从总体结构上看，系统采用典型的 B/S 架构。用户通过浏览器访问系统页面，前端页面提交表单和异步请求到 Flask 服务器；Flask 根据请求类型调用数据库模型或目标检测模块处理业务，"
        "处理完成后再将 HTML 页面或 JSON 结果返回前端。系统结构可概括为表现层、业务层、数据层和算法层四部分。表现层对应 Jinja2 模板与前端脚本，业务层对应路由函数，数据层对应 SQLite 数据库，算法层对应 YOLO 检测模块。"
    )
    add_paragraph(
        doc,
        "在目录结构上，项目根目录下的 `app.py` 为统一入口；`templates/` 包含登录页、仪表盘、检测页、历史页、报表页、销售页、帮助页、用户管理页和设置页；"
        "`static/` 包含结果图、上传图和帮助中心图谱素材；`utils/` 中的 `detection.py` 负责推理，`img_process.py` 负责图像预处理扩展。"
    )
    add_paragraph(
        doc,
        "在应用启动阶段，系统会自动创建上传目录和结果目录，随后在应用上下文中执行数据库建表逻辑。如果数据库中不存在默认管理员账号，则自动创建 `admin/123456`；"
        "若库存表为空，则初始化 A 级、B 级和 C 级三类库存记录。由此可见，系统在设计时考虑了“首次启动可直接运行”的要求，减少了部署前的人工初始化步骤。"
    )
    add_heading(doc, "4.1.1 模块划分设计", 3)
    add_paragraph(
        doc,
        "表现层以 `base.html` 为公共布局模板，在左侧统一放置导航菜单和用户信息区，右侧展示各业务页面内容。这样不仅保证了系统风格统一，也使后续新增页面时只需关注局部功能区的实现。"
        "检测页、历史页、报表页、销售页和后台管理页都继承自该基础模板，体现出较好的页面复用性。"
    )
    add_paragraph(
        doc,
        "业务层主要由 `app.py` 中的路由函数组成。该文件虽然集中，但逻辑清晰，能够直观反映从请求接收、权限判断、模型调用到数据库提交的完整过程。对于课程设计来说，这种结构比过度拆分的微服务方式更便于教学讲解和答辩展示。"
    )
    add_heading(doc, "4.1.2 数据流设计", 3)
    add_paragraph(
        doc,
        "系统的数据流主要分为三条主线。第一条是检测数据流：图片上传或拍照后进入 YOLO 推理，生成结果图、缺陷类别和等级判定，再写入检测历史并同步更新库存。"
        "第二条是统计数据流：系统从检测历史与工作日志中提取数据，计算今日检测量、缺陷分布与趋势曲线，并在前端渲染图表。第三条是销售数据流：用户在销售页面提交订单后，系统读取库存数量、判断是否可出库并写入销售记录。"
    )
    add_center_note(doc, "【此处后期补充图4-1 系统总体结构图】")

    add_heading(doc, "4.2 系统顺序图设计", 2)
    add_heading(doc, "4.2.1 登录顺序图", 3)
    add_paragraph(
        doc,
        "用户访问登录页面后，系统根据用户在前端选择的登录标签页提交角色意图。后端收到用户名、密码和角色后，首先查询用户信息，再使用密码哈希校验身份，随后核对账号角色与登录入口是否一致。"
        "验证通过时，系统调用 `login_user` 写入登录状态，并新增一条工作日志记录；验证失败时，系统返回错误提示。"
    )
    add_center_note(doc, "【此处后期补充图4-2 登录顺序图】")
    add_heading(doc, "4.2.2 缺陷检测顺序图", 3)
    add_paragraph(
        doc,
        "用户在检测页面选择本地图片或通过摄像头生成图片后，前端将文件封装为表单提交至 `/detect` 路由。后端获取当前置信度阈值和模型文件路径，保存原图后调用 `process_image` 完成推理。"
        "系统解析检测框类别与置信度，将英文类别映射为中文缺陷，再根据裂纹、死节、活节等规则判定板材等级。最后系统保存结果图、写入检测历史、更新库存并将检测结果返回前端显示。"
    )
    add_center_note(doc, "【此处后期补充图4-3 缺陷检测顺序图】")
    add_heading(doc, "4.2.3 历史追溯与报表顺序图", 3)
    add_paragraph(
        doc,
        "当用户进入历史记录或日志报表页面时，系统会根据当前角色过滤查询范围。历史记录模块按时间倒序分页返回检测数据，并支持按照缺陷类型模糊搜索。"
        "日志报表模块则统计今日检测量、累计工作时长、缺陷类别占比和近 7 天趋势数据，并将结果传递给前端图表组件完成可视化。"
    )
    add_center_note(doc, "【此处后期补充图4-4 历史与报表顺序图】")
    add_heading(doc, "4.2.4 销售出库顺序图", 3)
    add_paragraph(
        doc,
        "销售页面提交订单时，系统首先读取用户输入的客户名称、等级、数量和金额，随后检查库存表中对应等级的当前数量。若库存不足，系统返回提示信息并终止提交；若库存充足，"
        "系统扣减库存数量并写入销售记录，完成出库操作。该流程体现出检测等级与库存管理之间的业务联动。"
    )
    add_center_note(doc, "【此处后期补充图4-5 销售出库顺序图】")

    add_heading(doc, "4.3 数据库设计", 2)
    add_heading(doc, "4.3.1 数据库实体设计", 3)
    add_paragraph(
        doc,
        "按照现有模型定义，用户实体负责保存账号、密码和角色信息；检测历史实体记录每次检测的文件、缺陷、等级和时间；工作日志实体记录用户登录与退出时间；销售记录实体保存客户、等级、数量和金额；"
        "系统设置实体存储置信度阈值和当前模型文件名；库存实体负责保存各等级板材数量。上述实体共同支撑系统业务运行。"
    )
    add_center_note(doc, "【此处后期补充图4-6 主要实体属性图】")
    add_heading(doc, "4.3.2 数据库 E-R 设计", 3)
    add_paragraph(
        doc,
        "系统以用户实体为中心，与检测历史和工作日志形成一对多关系；库存与销售记录通过等级字段形成业务对应关系；系统设置独立存放全局参数，不直接依附于某个业务实体。"
        "总体来看，数据库结构较为简洁，适合课程设计场景下快速实现，同时也为后续扩展订单、设备、分类和统计维度保留了空间。"
    )
    add_center_note(doc, "【此处后期补充图4-7 系统 E-R 图】")
    add_database_tables(doc)

    add_heading(doc, "5 系统实现", 1)
    add_heading(doc, "5.1 用户功能模块的实现", 2)
    add_heading(doc, "5.1.1 登录与退出", 3)
    add_paragraph(
        doc,
        "登录页面采用双标签页设计，分别对应普通用户和管理员登录入口。用户提交表单后，后端先校验用户名和密码，再检查账号角色与前端所选登录入口是否一致。"
        "该设计避免了普通用户误从管理员入口登录的问题，也便于在同一页面上实现不同角色的视觉区分。用户退出时，系统会自动补全当前工作日志的退出时间，并清空会话信息。"
    )
    add_center_note(doc, "【此处后期补充图5-1 登录界面截图】")
    add_heading(doc, "5.1.2 仪表盘与个人中心", 3)
    add_paragraph(
        doc,
        "登录成功后，系统默认进入仪表盘页面，向用户展示开始检测、查看日志和销售管理三个核心入口。页面以卡片方式呈现主要业务功能，便于快速进入。"
        "个人中心提供密码修改功能，用户输入旧密码、新密码和确认密码后，系统将校验原密码是否正确以及两次新密码是否一致，成功后强制用户重新登录。"
    )
    add_center_note(doc, "【此处后期补充图5-2 仪表盘与个人中心截图】")
    add_heading(doc, "5.1.3 公共布局与导航实现", 3)
    add_paragraph(
        doc,
        "系统所有登录后的业务页面均继承 `base.html`。基础模板中统一引入 Bootstrap、Font Awesome、Chart.js 与 Viewer.js，同时提供固定提示消息区域，用于显示登录成功、设置保存成功、库存不足等提示信息。"
        "模板左侧侧边栏按照“数据管理”“知识库”“系统管理”等分区组织菜单，不同角色在导航中呈现不同入口，管理员才可见用户管理和系统设置功能。"
    )
    add_paragraph(
        doc,
        "这种基于模板继承和条件渲染的实现方式，不仅减少了重复页面代码，也保证了不同角色使用体验的一致性。"
    )
    add_center_note(doc, "【此处后期补充图5-3 公共布局与导航截图】")

    add_heading(doc, "5.2 木材缺陷检测模块的实现", 2)
    add_heading(doc, "5.2.1 本地上传与摄像头采集", 3)
    add_paragraph(
        doc,
        "检测页面通过选项卡方式将“本地上传”和“摄像头采集”两种检测模式整合到同一界面。用户可直接上传单张或多张图片，系统支持批量检测；也可调用浏览器摄像头进行实时拍照，再将拍摄结果转换为 Blob 对象后提交到后端。"
        "这种设计兼顾了离线图片检测和现场采集检测两类使用场景。"
    )
    add_center_note(doc, "【此处后期补充图5-4 检测页面截图】")
    add_heading(doc, "5.2.2 模型推理与缺陷映射", 3)
    add_paragraph(
        doc,
        "后端在接收到文件后，会从系统设置中读取当前置信度阈值和模型文件名。如果后台已上传新的 `.pt` 文件，则优先使用新模型；否则默认使用项目根目录中的 `best.pt`。"
        "`utils/detection.py` 中通过全局变量缓存模型对象，避免重复加载带来的开销。模型输出的原始类别为英文名称，系统随后使用 `DEFECT_MAPPING` 将其转换为中文显示结果。"
    )
    add_paragraph(
        doc,
        "由于一张图片中可能存在多个检测框，系统会对类别名称进行去重并计算最大置信度，将其与结果图文件名一并返回前端。页面最终展示结果图缩略图、缺陷类型标签、等级判定和置信度信息。"
    )
    add_center_note(doc, "【此处后期补充图5-5 推理结果展示截图】")
    add_heading(doc, "5.2.3 等级判定与库存联动", 3)
    add_paragraph(
        doc,
        "项目将缺陷识别结果与板材等级划分结合起来：无缺陷判定为 A 级，存在活节判定为 B 级，存在裂纹或死节则判定为 C 级。系统在得到最终等级后，会自动查询库存表并对对应等级数量加一，"
        "从而将检测行为自动转化为入库行为。这一实现使系统不再停留在“识别演示”层面，而是向业务闭环迈进一步。"
    )
    add_center_note(doc, "【此处后期补充图5-6 等级判定与库存联动截图】")
    add_heading(doc, "5.2.4 检测结果持久化实现", 3)
    add_paragraph(
        doc,
        "每次检测完成后，系统会构造 `DetectionHistory` 实体并写入数据库，保存原始文件名、结果图文件名、缺陷类型、最高置信度、用户 ID 和等级结果。"
        "结果图文件保存至 `static/results` 目录，前端通过静态资源路径直接加载显示。"
    )
    add_paragraph(
        doc,
        "这种设计使检测结果既能够被系统页面立即消费，也能够在后续历史记录和报表分析中重复使用，是系统实现“可追溯”能力的核心。"
    )
    add_center_note(doc, "【此处后期补充图5-7 检测结果持久化截图】")

    add_heading(doc, "5.3 检测历史与报表模块的实现", 2)
    add_heading(doc, "5.3.1 历史记录查询与删除", 3)
    add_paragraph(
        doc,
        "历史记录页面采用卡片式布局呈现检测结果图、文件名、等级、缺陷类型和检测时间。系统支持按缺陷关键词搜索，并通过分页机制控制单页记录数量。"
        "管理员可以查看全量历史数据，普通用户只能查看自己的检测记录。页面同时集成 Viewer.js，实现历史结果图片的大图预览。"
    )
    add_paragraph(
        doc,
        "对于删除操作，后端会进一步校验当前用户是否拥有该记录的操作权限，避免普通用户误删他人数据。"
    )
    add_center_note(doc, "【此处后期补充图5-8 历史记录页面截图】")
    add_heading(doc, "5.3.2 日志统计与 Excel 导出", 3)
    add_paragraph(
        doc,
        "日志与报表模块负责将历史检测数据转换为可视化统计结果。系统首先统计今日检测量，再通过工作日志表汇总累计工作时长，最后基于检测历史构建缺陷类别占比和近 7 天趋势图。"
        "当前实现中，缺陷占比图使用环形图展示，趋势图使用折线图展示。管理员可查看全量数据，普通用户仅查看个人数据。"
    )
    add_paragraph(
        doc,
        "在数据导出方面，系统利用 pandas 和 openpyxl 将检测记录整理为 DataFrame，再写入内存流并以附件形式返回浏览器，从而生成 Excel 报表。该功能有利于后续归档、分析和教学展示。"
    )
    add_center_note(doc, "【此处后期补充图5-9 日志与报表页面截图】")
    add_heading(doc, "5.3.3 工作时长统计实现", 3)
    add_paragraph(
        doc,
        "系统在用户登录时自动写入登录时间，在退出时补充退出时间。日志报表页面统计累计工作时长时，会遍历已结束的工作日志记录，并对当前尚未退出的登录状态进行动态补算。"
        "最终结果会以“多少小时多少分钟”的形式显示在页面卡片中，使报表不仅关注检测量，也体现操作员的工作状态。"
    )
    add_paragraph(
        doc,
        "该功能虽然实现逻辑较为简单，但非常契合实验教学或质检值班管理场景，也是当前系统区别于单纯图像识别演示项目的一项细节设计。"
    )
    add_center_note(doc, "【此处后期补充图5-10 工作时长统计截图】")

    add_heading(doc, "5.4 销售与后台管理模块的实现", 2)
    add_heading(doc, "5.4.1 销售与库存管理", 3)
    add_paragraph(
        doc,
        "销售页面分为库存概览和订单创建两部分。库存概览以卡片形式显示 A 级、B 级和 C 级板材当前库存数量；订单创建区域允许用户输入客户名称、产品等级、数量和金额。"
        "系统在创建订单前先校验库存是否足够，库存不足时直接提示剩余数量，库存充足时自动扣减库存并新增销售记录。该页面体现出检测数据与业务管理之间的衔接。"
    )
    add_center_note(doc, "【此处后期补充图5-11 销售与库存页面截图】")
    add_heading(doc, "5.4.2 用户管理", 3)
    add_paragraph(
        doc,
        "管理员用户管理页面提供新增和删除用户功能。管理员可输入用户名、初始密码和角色完成新账号创建，系统会对密码进行哈希保存；删除用户时会进行基本限制，避免管理员误删自身账号。"
        "这一模块满足了系统在教学演示和多人操作时对账号维护的需求。"
    )
    add_center_note(doc, "【此处后期补充图5-12 用户管理页面截图】")
    add_heading(doc, "5.4.3 系统设置与模型切换", 3)
    add_paragraph(
        doc,
        "系统设置页面允许管理员动态调整置信度阈值，并上传新的 `.pt` 模型文件。保存后，系统会将阈值写入 `SystemSettings` 表，将模型文件保存到项目根目录，并把新模型文件名更新为当前生效模型。"
        "下一次执行检测时，后端将优先加载该模型。这种设计增强了系统模型试验和版本切换的灵活性。"
    )
    add_center_note(doc, "【此处后期补充图5-13 系统设置页面截图】")
    add_heading(doc, "5.4.4 初始化与默认数据实现", 3)
    add_paragraph(
        doc,
        "在 `__main__` 启动逻辑中，系统自动创建数据库表、默认管理员账号以及三类库存记录，这使项目在首次部署时不需要额外执行初始化脚本。"
        "对于毕业设计系统而言，这种“一键启动即可演示”的实现方式有助于降低部署复杂度，也便于答辩时快速展示系统完整流程。"
    )
    add_paragraph(
        doc,
        "同时，系统会自动创建上传目录与结果目录，保证图片检测流程在空目录环境下也能正常运行，体现出对运行细节的基本考虑。"
    )
    add_center_note(doc, "【此处后期补充图5-14 初始化后系统运行截图】")

    add_heading(doc, "5.5 缺陷图谱与帮助中心的实现", 2)
    add_paragraph(
        doc,
        "帮助中心页面展示活节、死节和裂纹三类典型缺陷图谱，并配套说明定义、特征和等级判定规则。页面中嵌入项目 `static/img/` 目录下的示例图片，"
        "适合作为新用户、新员工学习木材缺陷基础知识的辅助模块。相较于仅提供检测结果，帮助中心进一步增强了系统的教学属性和说明性。"
    )
    add_center_note(doc, "【此处后期补充图5-15 帮助中心页面截图】")

    add_heading(doc, "6 系统测试", 1)
    add_heading(doc, "6.1 测试目的", 2)
    add_paragraph(
        doc,
        "系统测试的主要目标是验证当前项目各功能模块是否能够按照预期完成业务流程，重点包括登录鉴权、图片检测、等级判定、历史记录保存、报表统计、Excel 导出、销售扣库和管理员设置等内容。"
        "同时，通过测试检查页面交互是否顺畅、数据库写入是否正确、角色权限是否生效。"
    )
    add_heading(doc, "6.2 测试环境", 2)
    add_paragraph(
        doc,
        "测试环境基于 Windows 开发主机，Python 运行环境为项目本地 `.venv`，后端框架为 Flask 3.0.0，数据库为 SQLite，目标检测模型为项目中配置的 `.pt` 权重文件。"
        "浏览器使用支持摄像头权限的现代浏览器，以验证拍照检测流程。"
    )
    add_paragraph(
        doc,
        "在测试过程中，分别以普通用户和管理员两种身份进入系统，对登录鉴权、检测、查询、统计、导出、销售和设置等功能进行验证。测试方式以功能测试为主，辅以业务流程检查和异常输入检查。"
    )
    add_heading(doc, "6.3 测试用例", 2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["测试编号", "测试内容", "预期结果", "测试结论"]
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
    cases = [
        ("T1", "普通用户登录", "账号密码正确且角色匹配时成功进入仪表盘", "通过"),
        ("T2", "管理员误用普通用户入口登录", "系统提示身份不匹配并拒绝登录", "通过"),
        ("T3", "上传单张图片检测", "返回结果图、缺陷类别、等级和置信度", "通过"),
        ("T4", "批量图片检测", "多张图片依次返回检测结果并写入历史", "通过"),
        ("T5", "摄像头拍照检测", "拍照后成功提交并返回检测结果", "通过"),
        ("T6", "按缺陷类型查询历史", "历史记录支持关键字搜索与分页", "通过"),
        ("T7", "导出 Excel 报表", "浏览器下载包含检测记录的表格文件", "通过"),
        ("T8", "库存不足时提交销售订单", "系统阻止出库并提示当前库存", "通过"),
        ("T9", "管理员上传新模型", "保存成功并在后续检测中优先使用新模型", "通过"),
    ]
    for row in cases:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
    doc.add_paragraph()
    add_paragraph(
        doc,
        "除正向测试外，还对空文件上传、角色入口错误、库存不足、无权限删除历史记录等异常场景进行了验证。从当前代码执行结果看，系统能够给出相应的提示信息，并阻止非法或不合理操作继续执行。"
    )
    add_heading(doc, "6.4 测试结论", 2)
    add_paragraph(
        doc,
        "从当前代码实现和功能验证结果看，系统能够完成预定的核心业务流程：用户登录与鉴权稳定可用，检测模块可以输出缺陷类型和等级判定结果，历史与报表模块能够形成可视化追溯，"
        "销售模块实现了与库存的联动，管理员也可以完成基础配置和模型替换。整体系统满足毕业设计演示和课程实践场景下的使用要求。"
    )
    add_paragraph(
        doc,
        "需要指出的是，现阶段系统仍以单机 SQLite 部署为主，且删除检测历史时只删除数据库记录，前端提示中的“物理删除文件”并未在后端真正实现；此外，模型精度与实际数据集质量密切相关，"
        "在生产环境中仍需补充更系统的性能测试、异常恢复测试和并发访问测试。"
    )

    add_heading(doc, "7 总结与展望", 1)
    add_paragraph(
        doc,
        "本文围绕当前 `WebDetectionSystem` 项目的实际程序结构，对基于 Python 的 Web 木材表面缺陷智能检测系统进行了完整整理与说明。系统以 YOLO 模型为检测核心，以 Flask 为 Web 平台基础，"
        "实现了用户登录、图片检测、摄像头采集、缺陷映射、等级判定、历史追溯、统计报表、Excel 导出、库存销售联动、帮助中心和后台设置等功能。相比只展示算法结果的演示程序，"
        "本系统进一步打通了检测、记录、统计和管理流程，更符合毕业设计对“系统性”和“工程化”的要求。"
    )
    add_paragraph(
        doc,
        "后续展望主要包括以下几个方向：其一，进一步优化模型训练数据和类别体系，提高在复杂木纹、光照变化和多缺陷叠加场景下的泛化能力；其二，将数据库从 SQLite 升级到更适合并发访问的关系型数据库，"
        "并引入前后端分离部署架构；其三，增加真正的文件物理删除、日志审计、订单统计和权限细分功能，增强系统完整性；其四，结合移动端采集、边缘设备或生产线相机，实现更贴近现场的检测部署。"
        "通过这些改进，系统有望从教学原型进一步演进为可扩展的智能质检平台。"
    )

    add_reference_section(doc)
    add_acknowledgement(doc)

    for section in doc.sections:
        footer_p = section.footer.paragraphs[0]
        add_page_number(footer_p)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
