import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOC_PATH = Path(
    r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的Web木材表面缺陷智能检测系统的设计与实现.docx"
)
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_终稿清理前备份.docx")


def set_para_text(paragraph, text):
    """保留原段落位置，仅替换文字内容，避免影响文档整体结构。"""
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph):
    """安全删除指定段落。"""
    paragraph._element.getparent().remove(paragraph._element)


def heading_level(paragraph):
    """识别 Word 标题级别，便于按整节删除。"""
    name = paragraph.style.name
    if name.startswith("Heading "):
        try:
            return int(name.split()[-1])
        except ValueError:
            return None
    return None


def delete_section_by_heading(doc, heading_text):
    """按标题删除整节内容，直到遇到同级或更高级标题为止。"""
    paras = list(doc.paragraphs)
    start = None
    level = None
    for i, p in enumerate(paras):
        if p.text.strip() == heading_text:
            start = i
            level = heading_level(p)
            break
    if start is None or level is None:
        return 0

    end = len(paras)
    for j in range(start + 1, len(paras)):
        lv = heading_level(paras[j])
        if lv is not None and lv <= level:
            end = j
            break

    for p in paras[start:end]:
        delete_paragraph(p)
    return end - start


def set_table_borders(table):
    """将普通网格表调整为论文更常用的三线表样式。"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for child in list(borders):
        borders.remove(child)

    spec = {
        "top": {"val": "single", "sz": "12", "space": "0", "color": "000000"},
        "left": {"val": "nil"},
        "bottom": {"val": "single", "sz": "12", "space": "0", "color": "000000"},
        "right": {"val": "nil"},
        "insideH": {"val": "single", "sz": "8", "space": "0", "color": "000000"},
        "insideV": {"val": "nil"},
    }

    for edge, attrs in spec.items():
        elem = OxmlElement(f"w:{edge}")
        for key, value in attrs.items():
            elem.set(qn(f"w:{key}"), value)
        borders.append(elem)


def format_table(table):
    """统一表格对齐、字体和边框，尽量贴近论文模板要求。"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10.5)
                    if row_idx == 0:
                        run.bold = True


def replace_exact_text(doc, mapping):
    """针对整段完全匹配的内容执行替换。"""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in mapping:
            set_para_text(para, mapping[text])


def replace_contains_text(doc, mapping):
    """针对段落中包含指定短语的口语化描述进行整体替换。"""
    for para in doc.paragraphs:
        text = para.text.strip()
        for key, new_text in mapping.items():
            if key in text:
                set_para_text(para, new_text)
                break


def clear_cover_placeholders(doc):
    """清理封面和英文页中的占位词，避免留下明显未完成痕迹。"""
    replacements = {
        "计算机科学与技术专业  待填写": "计算机科学与技术专业",
        "指导教师  待填写": "指导教师",
        "Computer Science and Technology  To be filled": "Computer Science and Technology",
        "Tutor  To be filled": "Tutor",
    }
    replace_exact_text(doc, replacements)


def cleanup_testing_placeholders(doc):
    """删除第六章中明显未收尾的占位说明文字。"""
    targets = {
        "后期补充说明：此处建议放置用户成功登录、图片检测成功返回结果图或历史记录正常显示的系统截图。",
        "后期补充说明：此处建议放置角色不匹配提示、未选图片提示或库存不足提示等界面截图。",
    }
    for para in list(doc.paragraphs):
        if para.text.strip() in targets:
            delete_paragraph(para)


def cleanup_toc_residuals(doc):
    """清理手工目录中与正文不一致的旧条目。"""
    targets = {
        "3.4 性能与安全需求分析\t10",
        "3.4 性能与安全需求分析",
    }
    for para in list(doc.paragraphs):
        if para.text.strip() in targets:
            delete_paragraph(para)


def remove_dashboard_block(doc):
    """删除老师不建议保留的“仪表盘与个人中心”截图及其说明代码块。"""
    paras = list(doc.paragraphs)
    start = None
    end = None
    for i, para in enumerate(paras):
        if para.text.strip().startswith("仪表盘页面主要用于集中展示系统常用功能入口"):
            start = i
            break
    if start is None:
        return 0

    for j in range(start + 1, len(paras)):
        if paras[j].text.strip() == "5.1.2 检测历史查询界面":
            end = j
            break
    if end is None:
        return 0

    for para in paras[start:end]:
        delete_paragraph(para)
    return end - start


def rewrite_colloquial_paragraphs(doc):
    """将老师容易继续指出的口语化表达改成更正式的论文语体。"""
    exact_mapping = {
        "从实际使用感受来看，系统的操作门槛并不高。登录后常用入口都已经放在左侧侧边栏中，用户一般不需要反复翻找页面。检测页支持上传和拍照两种方式，管理员相关配置也集中在表单页面里，整体路径还是比较顺手的。": (
            "从系统使用流程来看，系统整体操作步骤较为清晰。登录后，常用功能入口集中显示在左侧导航区域，用户能够较快进入检测、历史记录、日志报表和销售管理等页面。"
            "检测页面同时支持本地上传与摄像头采集，管理员相关配置也集中在系统设置页面中，能够满足毕业设计演示场景下的基本操作需求。"
        ),
        "登录模块本身不算复杂，但里面有个比较实用的小处理。普通用户和管理员共用同一个登录页，不过通过两个标签页区分入口。这样一来，前端界面不会显得分散，后端也能继续根据当前入口判断角色是否匹配。退出登录时补写工作日志结束时间，这一点又和后面的工作时长统计接了起来，所以这个模块虽然简单，前后逻辑还是连着的。": (
            "登录模块采用统一登录页面设计，普通用户与管理员通过不同标签页进入对应身份入口。"
            "后端在接收登录请求后，会结合账号信息与当前入口类型完成角色匹配校验，从而保证不同身份用户访问权限的正确性。"
            "用户退出登录时，系统还会同步补写工作日志结束时间，为后续工作时长统计提供数据支持。"
        ),
        "仪表盘页面的作用主要是把系统最常用的入口集中展示出来，避免登录后再逐个菜单查找。当前页面中放置了开始检测、查看日志和销售管理三个功能卡片，比较适合答辩演示时快速跳转。个人中心部分虽然功能不多，但保留了密码修改这一项，用户修改成功后系统会主动要求重新登录，以保证新密码立即生效。": (
            "仪表盘页面主要用于集中展示系统常用功能入口，便于用户在登录后快速进入检测、日志报表和销售管理等核心业务页面。"
            "个人中心页面提供密码修改功能，用户完成密码更新后系统会要求重新登录，以保证新密码及时生效并提高账号使用安全性。"
        ),
        "历史记录页面的实现重点在于“可查”和“可看”。页面把结果图、文件名、等级、缺陷类型和检测时间统一放在卡片中，同时允许按缺陷关键词进行筛选。由于实际检测过程中会积累较多记录，系统使用分页方式控制单页数量，避免页面一次性堆太多内容。": (
            "历史记录页面的实现重点在于检测结果的可查询性与可追溯性。页面集中展示结果图、文件名、等级、缺陷类型和检测时间等信息，"
            "并支持按照缺陷关键词进行筛选。考虑到系统运行过程中会持续积累检测记录，因此页面采用分页方式控制单页展示数量，以提升浏览效率。"
        ),
        "删除记录时，后端还额外做了一层权限判断。管理员可以查看并管理全量数据，普通用户只能操作自己的记录。这样的处理虽然简单，但能避免在多人演示时出现权限混乱的问题。": (
            "在删除记录操作中，后端增加了权限控制逻辑。管理员可以查看并管理全量检测数据，普通用户仅能操作本人产生的记录，"
            "该设计有助于保证多用户场景下的数据访问边界和操作安全性。"
        ),
        "日志与报表这部分并不只是为了把页面做得更丰富。它真正承担的是把已经保存下来的检测数据重新整理成更适合查看和汇总的统计结果。系统会计算今日检测量、缺陷类别占比和近 7 天趋势，再交给前端图表组件渲染。就目前运行情况来看，折线图和环形图都已经能满足毕业设计展示需求。": (
            "日志与报表模块的主要作用是对系统已保存的检测数据进行统计整理与可视化展示。系统能够计算今日检测量、缺陷类别占比和近 7 天检测趋势，"
            "并通过前端图表组件完成可视化呈现，从而为运行状态分析与结果汇总提供支持。"
        ),
        "数据导出功能做得比较务实。系统直接利用 pandas 和 openpyxl 整理检测记录，再生成 Excel 供浏览器下载。对于毕业设计答辩来说，这个功能的价值不只是“多了一个按钮”，而是能把系统里的数据直接转成可以查看和提交的结果文件。": (
            "数据导出功能通过 pandas 和 openpyxl 对检测记录进行整理，并生成 Excel 文件供浏览器下载。"
            "该功能能够将系统中的结构化数据直接转换为便于查看与归档的表格文件，提升了检测结果的利用效率。"
        ),
        "这部分逻辑本身不算复杂，但它让系统记录的不只是“检测了多少张图”，还顺带反映了用户的工作时长。放在教学演示里，这算是一个比较细但比较容易被注意到的设计点。": (
            "该统计逻辑在实现上相对清晰，但能够使系统记录内容从单纯的检测次数扩展到用户工作时长信息，"
            "从而提升系统日志数据的完整性，并为报表分析提供更多维度的参考。"
        ),
        "销售页面的实现思路比较直接，但和检测业务结合得很紧。页面上方先展示 A、B、C 三类库存数量，下方再提供订单录入表单。用户提交销售订单后，系统会先检查当前等级库存是否足够，只有库存满足条件时才允许出库并写入销售记录。这样做能够防止销售数据和库存数据脱节。": (
            "销售与库存管理页面在业务设计上与检测模块保持紧密联动。页面上方展示 A、B、C 三类板材库存数量，下方提供销售订单录入表单。"
            "用户提交订单后，系统会优先校验当前等级库存是否充足，仅在库存满足条件时执行出库并写入销售记录，从而保证销售数据与库存数据的一致性。"
        ),
        "帮助中心页面放置了活节、死节和裂纹三类典型样张，同时给出简短定义、外观特征和等级说明。这个页面对系统运行本身不是必须的，但对于第一次接触项目的人非常有用，特别适合在答辩时辅助说明不同缺陷之间的区别。": (
            "帮助中心页面展示了活节、死节和裂纹三类典型样张，并给出相应的定义说明、外观特征和等级判定依据。"
            "该页面有助于用户快速理解不同缺陷之间的差异，也能够为系统演示和结果说明提供辅助支持。"
        ),
        "用户管理部分主要给管理员使用。当前功能没有做得很复杂，重点是支持新增和删除账号，并在创建用户时先对密码做哈希处理。虽然页面本身比较简洁，但对于一个完整系统来说，至少把最基本的账号维护补上了。": (
            "用户管理模块主要面向管理员使用，当前实现支持新增和删除用户账号，并在创建用户时对密码进行哈希处理。"
            "该模块实现了系统最基本的账号维护功能，为后续权限控制和多用户使用提供了基础保障。"
        ),
        "系统设置页目前保留的可调项不多，核心就是检测阈值和模型文件。阈值会影响后续检测时保留哪些框，模型上传则方便管理员在不改代码的前提下替换新的 `.pt` 权重。对毕业设计项目来说，这已经足以体现系统具备一定的可维护性。": (
            "系统设置页面当前主要提供检测阈值与模型文件两类可调参数。检测阈值会影响模型推理时目标框的保留条件，"
            "模型上传功能则使管理员能够在不修改程序代码的情况下替换新的 `.pt` 权重文件，从而体现系统具备一定的可维护性与扩展性。"
        ),
        "检测页面是整个系统里交互最频繁的页面之一。当前实现把“本地上传”和“摄像头采集”放在同一个页面中，通过选项卡切换，用户不需要来回跳转页面。实际调试时，这种组织方式比拆成两个独立页面更顺手，也更适合演示。后来为了让操作更直观，页面又补充了本地图片预览，用户在点击检测前就能先确认自己选中的文件。": (
            "检测页面是系统中交互最为频繁的业务页面之一。当前实现将“本地上传”和“摄像头采集”两种方式整合在同一页面中，"
            "通过选项卡进行切换，减少了用户在不同页面之间反复跳转的操作成本。为提高操作直观性，页面还增加了本地图片预览功能，便于用户在提交检测前确认所选文件。"
        ),
        "后端检测流程的入口在 `/detect` 路由中。系统收到图片后，会先读取当前阈值和模型配置，再把原图保存到上传目录，然后调用 `process_image` 完成推理。这里有一个比较实际的问题：模型并不一定总是固定的，因此后台设置中保留了模型切换功能，下一次检测时就会优先使用管理员上传的新权重文件。": (
            "后端检测流程的入口位于 `/detect` 路由。系统接收到图片后，会先读取当前阈值和模型配置，再将原图保存到上传目录，随后调用 `process_image` 执行模型推理。"
            "考虑到系统模型文件可能存在更新需求，因此后台设置页面保留了模型切换功能，后续检测时将优先使用管理员上传的新权重文件。"
        ),
        "在这个项目里，等级判定并不是结果页上的附加说明，而是检测流程里的关键环节。系统默认把无缺陷板材记为 A 级；识别到活节时判为 B 级；若出现裂纹或死节，则直接归入 C 级。等级一旦确定，后端就会继续更新库存表，因此一次检测并不会停留在页面展示层面。": (
            "在本系统中，等级判定并不是结果展示阶段的附加说明，而是检测流程中的关键业务环节。系统默认将无缺陷板材判定为 A 级；识别到活节时判定为 B 级；"
            "若检测到裂纹或死节，则直接归入 C 级。等级结果生成后，后端会继续更新库存表，从而使一次检测操作能够与后续库存管理流程衔接。"
        ),
        "每完成一次检测，系统都会把结果写入 `DetectionHistory` 表。保存的字段包括原始文件名、结果图文件名、缺陷类型、最高置信度、操作用户和等级结果。结果图本身则单独保存在 `static/results` 目录中，前端页面通过静态路径直接访问。这种处理方式的好处是清楚：图片和结构化记录分开保存，但页面和报表又都能把它们关联起来。": (
            "每完成一次检测后，系统都会将结果写入 `DetectionHistory` 表。保存内容包括原始文件名、结果图文件名、缺陷类型、最高置信度、操作用户和等级结果等字段。"
            "结果图片则单独保存在 `static/results` 目录中，前端页面通过静态资源路径进行访问。该设计实现了图片文件与结构化记录的分离存储，同时便于页面展示与报表统计进行关联。"
        ),
        "总体来看，这篇论文的写作基础不是空想出来的一套方案，而是已经搭建并调试过的 `WebDetectionSystem` 项目。论文中的大多数页面、流程和数据结构都能在当前程序中找到对应实现，因此整篇内容和项目本身保持了较强的一致性。": (
            "总体来看，本文的研究与实现建立在已经完成开发和调试的 `WebDetectionSystem` 项目基础之上。"
            "论文中涉及的大部分页面、业务流程和数据结构均能够在当前程序中找到对应实现，因此论文内容与系统项目本身保持了较强的一致性。"
        ),
    }
    replace_exact_text(doc, exact_mapping)


def style_figure_and_table_titles(doc):
    """统一图题、表题段落的对齐与字号。"""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("图") or text.startswith("表"):
            if "-" in text[:6] or "－" in text[:6]:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(6)
                for run in para.runs:
                    run.font.name = "宋体"
                    if run._element.rPr is not None:
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10.5)


def main():
    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(DOC_PATH)

    # 1. 删除导师明确要求去掉的 3.4 小节。
    delete_section_by_heading(doc, "3.4 性能与安全需求分析")
    cleanup_toc_residuals(doc)

    # 2. 清理封面和测试章节中的明显占位内容。
    clear_cover_placeholders(doc)
    cleanup_testing_placeholders(doc)
    remove_dashboard_block(doc)

    # 3. 将明显口语化段落改成更正式的论文表达。
    rewrite_colloquial_paragraphs(doc)

    # 4. 统一数据库表与测试表样式为三线表。
    for table in doc.tables[1:]:
        format_table(table)

    # 5. 统一图表标题段落样式。
    style_figure_and_table_titles(doc)

    doc.save(DOC_PATH)
    print("done")


if __name__ == "__main__":
    main()
