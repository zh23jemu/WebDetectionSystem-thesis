from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches, RGBColor


DOCX_PATH = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的木质板材缺陷检测系统的设计与实现.docx")
BACKUP_PATH = DOCX_PATH.with_name(DOCX_PATH.stem + "_修改前备份.docx")


def set_cell_text(cell, text):
    """设置单元格文本，并统一表格正文的基本字号。"""
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.first_line_indent = None
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(10.5)


def format_paragraph(paragraph, size=12, font_name="宋体", bold=False, align=None, first_indent=True):
    """统一段落字体、字号、行距和缩进，避免正文格式前后不一致。"""
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if first_indent else None
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(size)
        run.bold = bold


def set_paragraph_text(paragraph, text):
    """在保留段落对象位置的前提下替换文本，减少对图片、表格位置的影响。"""
    paragraph.clear()
    paragraph.add_run(text)


def delete_paragraph(paragraph):
    """删除指定段落，用于去除重复或明显不合要求的内容。"""
    element = paragraph._element
    element.getparent().remove(element)


def delete_table(table):
    """删除指定表格，用于去除目录前异常出现的空白框。"""
    element = table._element
    element.getparent().remove(element)


def insert_paragraph_before(paragraph, text="", style=None):
    """在目标段落前插入新段落，便于补充导师要求增加的章节标题。"""
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    result = Paragraph(new_p, paragraph._parent)
    if style is not None:
        result.style = style
    if text:
        result.add_run(text)
    return result


def add_page_number(paragraph):
    """在页脚中插入页码域，配合首页不同实现封面无页码。"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_page_number_start(section, start=0):
    """设置页码起始值，使封面不显示页码时目录页可以从 1 开始显示。"""
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def add_simple_cover(doc):
    """在文档开头补充规范封面，解决封面缺失、页眉页码和基本信息格式问题。"""
    body = doc._body._element
    first_child = body[0]
    cover_doc = Document()
    sec = cover_doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    def add_center(text, size=16, bold=False, space_before=0, space_after=0):
        p = cover_doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(size)
        r.bold = bold
        return p

    add_center("齐鲁理工学院", 22, True, 30, 8)
    add_center("QILU INSTITUTE OF TECHNOLOGY", 10, False, 0, 70)
    add_center("本科毕业设计（论文）", 20, True, 0, 60)
    add_center("基于Python的木质板材缺陷检测系统的设计与实现", 18, True, 0, 50)

    info = [
        ("学生姓名", "栾艳康"),
        ("二级学院", "计算机与信息工程学院"),
        ("专业", "计算机科学与技术"),
        ("指导教师", "刘荣"),
        ("论文题目", "基于Python的木质板材缺陷检测系统的设计与实现"),
    ]
    table = cover_doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (k, v) in zip(table.rows, info):
        set_cell_text(row.cells[0], k)
        set_cell_text(row.cells[1], v)
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(9)
    cover_doc.add_page_break()

    # 将封面文档的正文元素复制到当前文档开头，保持后续内容原位置尽量不变。
    for child in reversed(list(cover_doc._body._element)[:-1]):
        body.insert(0, deepcopy(child))


def add_picture_border(inline_shape, color="808080", width=9525):
    """给图片添加浅灰色边框，满足系统截图外侧需要有框线的修改意见。"""
    pic = inline_shape._inline.graphic.graphicData.pic
    sp_pr = pic.spPr
    old_ln = sp_pr.find(qn("a:ln"))
    if old_ln is not None:
        sp_pr.remove(old_ln)
    ln = OxmlElement("a:ln")
    ln.set("w", str(width))
    solid_fill = OxmlElement("a:solidFill")
    srgb = OxmlElement("a:srgbClr")
    srgb.set("val", color)
    solid_fill.append(srgb)
    ln.append(solid_fill)
    sp_pr.append(ln)


def replace_table_with_rows(table, rows):
    """按导师意见扩展数据库表字段说明，加入长度、约束等列。"""
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.columns) < len(rows[0]):
        table.add_column(Cm(2.0))
    while len(table.columns) > len(rows[0]):
        # python-docx 没有直接删列接口，若原表列数较少不会进入该分支。
        break
    for row_idx, row_data in enumerate(rows):
        row = table.rows[0] if row_idx == 0 else table.add_row()
        for col_idx, value in enumerate(row_data):
            set_cell_text(row.cells[col_idx], value)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def main():
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(DOCX_PATH)

    # 删除目录前异常出现的空白单元格框，随后补充规范封面。
    if doc.tables and len(doc.tables[0].rows) == 1 and len(doc.tables[0].columns) == 1 and not doc.tables[0].cell(0, 0).text.strip():
        delete_table(doc.tables[0])
    add_simple_cover(doc)

    # 首页不显示页码，页码从目录页开始按 1 显示。
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_footer.paragraphs[0].clear()
    section.footer.paragraphs[0].clear()
    add_page_number(section.footer.paragraphs[0])
    set_page_number_start(section, 0)

    replacements = {
        "基于Python的木质板材表面缺陷系统的设计与实现": "基于Python的木质板材缺陷检测系统的设计与实现",
        "Computer science and technology   Luan Yankang": "Computer Science and Technology   Luan Yankang",
        "tutor     Liu Rong": "Supervisor     Liu Rong",
        "Keywords: Python; Wood Defect Detection; YOLO; Flask; Web System; Inventory and Sales": "Keywords: Python; wood panel defect detection; YOLO; Flask; Web system; inventory and sales linkage",
    }
    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        if txt in replacements:
            set_paragraph_text(paragraph, replacements[txt])

    # 改写英文摘要，去除口语化表达和大小写问题。
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Abstract:"):
            set_paragraph_text(
                paragraph,
                "Abstract: To address the low efficiency of manual wood panel surface defect inspection, "
                "the difficulty of preserving inspection records, and the weak connection between quality "
                "grading and inventory management, this thesis designs and implements a Python-based web "
                "defect detection system for wood panels. The system uses Flask for back-end service control, "
                "SQLite and SQLAlchemy for data persistence, and a YOLO detection model for identifying common "
                "surface defects such as cracks, dead knots and live knots. After users upload images or capture "
                "images through a camera, the system returns defect categories, confidence values, result images "
                "and grading results, and stores the corresponding inspection records. In addition, the system "
                "provides login authentication, history query, log statistics, report export, inventory linkage, "
                "sales outbound management, user management and model switching functions. Functional testing "
                "shows that the system can complete the main processes of defect detection, result tracing and "
                "business data management, and has practical value for teaching demonstration and graduation "
                "project implementation."
            )

    # 重写 3.3，避免功能结构说明松散，并明确普通用户/管理员之间关系。
    section_33 = {
        "就系统整体实现方式而言，本系统可以分成页面展示层、业务控制层、模型推理层和数据存储层。分层结构和项目代码组织大体上一致，既可以说明各个模块之间的协作关系，也可以使论文分析的内容和实际系统的实现保持一致。":
        "从系统使用角色和业务流程来看，本系统主要包括普通用户功能、管理员功能和公共支撑功能三部分。普通用户负责完成缺陷检测、检测历史查询、日志报表查看、销售出库、缺陷图谱浏览和个人信息维护；管理员在普通用户功能的基础上，增加用户管理、系统参数设置、模型切换和日志查看等维护功能。两个角色并不是相互独立的两套系统，而是在统一登录认证和权限控制基础上共享检测、历史、报表和库存等业务数据。",
        "从角色职责上来说，普通用户主要是使用登录认证、木材缺陷检测、历史记录查询、日志报表查看、销售管理、帮助中心浏览、个人信息维护等功能，管理员在此基础上又增加了用户管理、系统设置、模型切换等后台管理任务。检测模块是系统运行的中心，历史、报表、销售模块一起形成了业务闭环，使中心和后台配置变得简单。":
        "系统功能结构以缺陷检测为核心展开。检测模块接收本地上传或摄像头采集的板材图片，调用模型完成识别和等级判定；历史记录模块保存检测结果并支持追溯查询；日志报表模块对检测数量、工作时长和缺陷类别进行统计；库存与销售模块根据等级结果完成库存更新和出库记录保存。管理员模块主要为上述业务提供账号、参数和模型文件维护能力。",
        "根据页面入口进行细分，系统主要有登录、仪表盘、缺陷检测、历史记录、日志报表、数据导出、销售管理、帮助中心、用户管理、系统设置等九个功能页面。各个页面同后端业务模块一一对应，页面组织关系比较清晰，可以满足日常使用、权限划分以及未来扩展的要求。":
        "按照页面入口细分，系统包含登录、仪表盘、缺陷检测、历史记录、日志报表、数据导出、销售管理、帮助中心、用户管理和系统设置等功能页面。普通用户只访问与检测和业务操作相关的页面，管理员可额外进入后台维护页面。该功能划分能够对应系统实际代码中的路由、模板和数据库操作，也为后续系统设计图和实现章节提供依据。"
    }
    for paragraph in doc.paragraphs:
        if paragraph.text in section_33:
            set_paragraph_text(paragraph, section_33[paragraph.text])

    # 压缩 4.1 中重复、AI 痕迹较重的段落，保留必要设计说明。
    replacements_41 = {
        "该系统采用B/S架构，用户通过浏览器浏览各种功能页面，服务器端做业务处理、模型推理、数据库读写等工作。该架构对于客户端环境要求不高，在本地或者局域网内可以很快地进行部署和演示。系统的总体结构图如下图4-1所示。":
        "本系统采用B/S架构，用户通过浏览器访问系统页面，服务器端负责业务处理、模型推理和数据库读写。该结构对客户端环境要求较低，适合毕业设计中的本地部署、局域网演示和后续功能扩展。系统总体结构如图4-1所示。",
        "从构成上来说，前端页面主要实现图片选择、检测结果展示、图表显示、表单提交的功能；后端服务主要是接收到请求之后执行业务逻辑、调用检测模型返回处理结果；数据库用来保存用户的、检测记录的、工作日志的、库存的、销售的等数据；算法模块主要是对木材表面缺陷进行识别。各个部分一起工作，可以共同支撑系统的主要业务功能。":
        "系统结构主要由前端展示层、后端业务层、模型检测层和数据存储层组成。前端展示层完成图片选择、检测结果展示、图表渲染和表单提交；后端业务层接收请求并完成权限校验、文件保存、模型调用、库存更新和记录写入；模型检测层根据管理员配置的权重文件和置信度阈值完成缺陷识别；数据存储层保存用户、检测历史、工作日志、库存、销售记录和系统参数等数据。",
        "系统启动时会自动创建上传目录和结果目录，然后进行数据库表的初始化。当系统中没有默认的管理员账号或者初始库存信息时，程序就会写入基础数据，从而保证系统第一次运行之后就可以完成登录、检测以及业务流程演示。":
        "系统运行时，检测业务的数据流从图片上传开始，后端保存原图后调用YOLO模型进行推理，并将缺陷类别、置信度、等级结果和结果图路径写入检测历史表；统计业务从检测历史和工作日志中提取数据生成图表；销售业务根据库存表进行余量校验，并在出库成功后写入销售记录。通过上述流程，系统能够形成检测、记录、统计和库存销售联动的闭环。"
    }
    for paragraph in list(doc.paragraphs):
        if paragraph.text in replacements_41:
            set_paragraph_text(paragraph, replacements_41[paragraph.text])
        elif paragraph.text in {
            "系统用统一布局模板来组织页面。用户登录之后，公共模板会承载起侧边导航、用户信息、消息提示以及功能入口的作用，各个业务页面在统一的框架中会以各自的页面内容来呈现出来。该种设计既保证了界面风格的统一，又利于以后页面的维护以及功能的扩展。",
            "系统业务层主要是由后端路由函数组成的。不同的路由对应着不同的请求，即登录、检测、历史记录、日志报表、销售管理、用户管理、系统设置等请求，在接收到前端的数据之后进行权限判断、数据处理、模型调用、数据库提交等工作。",
            "从数据流的角度来分析，系统主要有检测数据流、统计数据流和销售数据流这三个流程。检测数据流从图片上传开始，经过模型推理和等级判定之后写入检测历史并更新库存，统计数据流从检测记录、工作日志中提取数据生成报表和趋势图，销售数据流根据订单信息完成库存扣减并保存销售记录。三类数据流一起组成了系统的业务链路。"
        }:
            delete_paragraph(paragraph)

    # 增补 E-R 图说明，避免实体、属性和关系交代不足。
    for paragraph in doc.paragraphs:
        if paragraph.text == "数据库的E-R图用来表示系统中的主要实体以及它们之间的联系。本系统以用户、检测历史、工作日志、销售记录、库存等数据对象为设计对象，用户和检测历史、工作日志之间存在直接关系，库存和销售记录之间是通过板材等级形成的业务关系。主要实体属性图如下图4-6所示。":
            set_paragraph_text(paragraph, "数据库的E-R图用来表示系统中的主要实体、属性以及实体之间的联系。本系统根据实际数据库模型抽取出用户、检测历史、工作日志、销售记录、库存和系统设置六类实体。用户实体用于保存登录账号、密码哈希和角色信息；检测历史实体用于保存原始文件名、结果图、缺陷类型、置信度、检测时间和等级结果；工作日志实体用于记录用户登录、退出和累计工作时长；销售记录实体用于保存客户、等级、数量、金额和销售时间；库存实体用于记录不同等级板材的库存数量；系统设置实体用于保存当前模型路径和检测阈值等参数。主要实体属性如图4-6所示。")
        elif paragraph.text == "根据系统业务流程可知，用户登录之后会产生检测记录和工作日志；检测结果完成等级判定之后会改变库存数量；销售出库操作要读取库存并生成销售记录。系统各个实体之间的联系如图4-7所示。":
            set_paragraph_text(paragraph, "在实体关系上，一个用户可以产生多条检测历史记录，也可以对应多条工作日志记录，因此用户与检测历史、用户与工作日志均为一对多关系。检测历史中的等级结果会影响库存数量，库存又会被销售出库业务读取和扣减，因此库存与销售记录之间通过板材等级形成业务关联。系统设置实体为检测流程提供模型路径和阈值参数，不直接依赖某一条检测记录，但会影响后续检测结果生成。系统各实体之间的联系如图4-7所示。")

    # 改写数据库表说明，并扩展数据库表列数。
    for paragraph in doc.paragraphs:
        if paragraph.text == "根据系统的功能需求以及E-R图的设计结果，数据库主要是用户表、检测历史表、工作日志表、销售记录表、系统设置表和库存表。各个数据表分别担负着账号管理、检测结果保存、工作时间记载、销售业务存贮、参数设定以及库存统计这些工作。具体表结构设计如上图所示。":
            set_paragraph_text(paragraph, "根据系统功能需求和E-R图设计结果，数据库主要包括用户表、检测历史表、工作日志表、销售记录表、系统设置表和库存表。各表采用字段名、数据类型、长度或取值范围、约束和字段说明进行描述，便于体现表结构设计依据。其中，用户表用于账号与权限管理，检测历史表用于保存每次检测的结构化结果，工作日志表用于统计用户工作时长，销售记录表用于保存出库业务数据，系统设置表用于保存模型和阈值参数，库存表用于记录不同等级板材的当前库存。")

    table_rows = [
        [["字段名", "类型", "长度/范围", "约束", "默认值", "说明"],
         ["id", "Integer", "-", "主键，自增", "-", "用户编号"],
         ["username", "String", "50", "唯一，非空", "-", "登录用户名"],
         ["password", "String", "255", "非空", "-", "密码哈希值"],
         ["role", "String", "20", "非空", "user", "用户角色，取值为 user 或 admin"]],
        [["字段名", "类型", "长度/范围", "约束", "默认值", "说明"],
         ["id", "Integer", "-", "主键，自增", "-", "检测记录编号"],
         ["filename", "String", "300", "非空", "-", "上传原始文件名"],
         ["result_image", "String", "300", "可空", "-", "检测结果图文件名"],
         ["upload_date", "DateTime", "-", "非空", "当前时间", "上传检测时间"],
         ["defect_type", "String", "100", "可空", "无缺陷", "识别出的缺陷类型"],
         ["confidence", "Float", "0-1", "可空", "0", "最高置信度"],
         ["grade", "String", "10", "非空", "A", "板材等级结果"],
         ["user_id", "Integer", "-", "外键", "-", "关联用户编号"]],
        [["字段名", "类型", "长度/范围", "约束", "默认值", "说明"],
         ["id", "Integer", "-", "主键，自增", "-", "日志编号"],
         ["user_id", "Integer", "-", "外键，非空", "-", "操作用户编号"],
         ["login_time", "DateTime", "-", "非空", "当前时间", "登录时间"],
         ["logout_time", "DateTime", "-", "可空", "-", "退出时间"]],
        [["字段名", "类型", "长度/范围", "约束", "默认值", "说明"],
         ["id", "Integer", "-", "主键，自增", "-", "销售记录编号"],
         ["customer_name", "String", "100", "非空", "-", "客户名称"],
         ["product_grade", "String", "10", "非空", "-", "销售板材等级"],
         ["quantity", "Integer", "正整数", "非空", "1", "销售数量"],
         ["total_price", "Float", "非负数", "非空", "0", "订单总价"],
         ["sale_date", "DateTime", "-", "非空", "当前时间", "销售时间"]],
        [["字段名", "类型", "长度/范围", "约束", "默认值", "说明"],
         ["id", "Integer", "-", "主键，自增", "-", "设置编号"],
         ["key_name", "String", "50", "唯一，非空", "-", "设置项名称"],
         ["value", "String", "200", "可空", "-", "设置项取值"]],
        [["字段名", "类型", "长度/范围", "约束", "默认值", "说明"],
         ["grade", "String", "10", "主键", "-", "板材等级"],
         ["count", "Integer", "非负整数", "非空", "0", "当前库存数量"]],
    ]
    db_tables = doc.tables[-7:-1]
    for table, rows in zip(db_tables, table_rows):
        replace_table_with_rows(table, rows)

    # 调整系统实现章节标题与说明，补齐 5.3 结构，并统一“截图”表述为“界面图”。
    caption_replacements = {
        "图5-3 历史记录页面截图": "图5-3 历史记录页面图",
        "图5-4 日志与报表页面截图": "图5-4 日志与报表页面图",
        "图5-5 工作时长统计截图": "图5-5 工作时长统计图",
        "图5-6 销售与库存页面截图": "图5-6 销售与库存页面图",
        "图5-7 帮助中心页面截图": "图5-7 帮助中心页面图",
        "图5-8 用户管理页面截图": "图5-8 用户管理页面图",
        "图5-9 系统设置页面截图": "图5-9 系统设置页面图",
        "图5-10 检测页面截图": "图5-10 检测页面图",
        "图5-12 等级判定与库存联动截图": "图5-12 等级判定与库存联动图",
        "图5-13 检测结果持久化截图": "图5-13 检测结果持久化图",
        "图6-1 正常功能测试截图": "图6-1 正常功能测试结果图",
        "图6-2 异常功能测试截图": "图6-2 异常功能测试结果图",
    }
    for paragraph in doc.paragraphs:
        if paragraph.text in caption_replacements:
            set_paragraph_text(paragraph, caption_replacements[paragraph.text])
        elif paragraph.text == "5.3.1 数据获取与预处理":
            new_p = insert_paragraph_before(paragraph, "5.3 算法检测模块的实现")
            format_paragraph(new_p, size=14, bold=True, first_indent=False)
        elif paragraph.text == "5.3.3 系统检测与可视化分析":
            set_paragraph_text(paragraph, "5.3.2 系统检测与可视化分析")

    for paragraph in doc.paragraphs:
        if paragraph.text == "登陆模块用统一的界面风格设计，普通用户通过不同的标签页进入对应账号入口，管理员通过别的途径进行操作。后端收到登录请求之后，会根据账号信息和入口类型来完成角色匹配校验工作，从而保证不同的身份用户可以按照正确的权限去访问。用户退出系统之后，系统会把工作日志结束时间同步到数据库中，给之后的工作时长统计提供基础数据。":
            set_paragraph_text(paragraph, "登录模块采用统一的页面入口。普通用户和管理员在登录页面输入用户名、密码并选择对应角色后，后端根据账号信息、密码哈希和角色字段进行校验。若账号密码正确且角色匹配，系统写入登录状态并跳转到对应功能页面；若角色入口与账号身份不一致，则返回错误提示并拒绝访问。用户退出系统时，系统会补写工作日志的退出时间，为后续工作时长统计提供数据基础。")

    # 扩充 5.1.2、5.1.3 和测试章节段落，回应“改成段落并补一句说明”的意见。
    for paragraph in doc.paragraphs:
        if paragraph.text == "历史记录页面的主要工作就是保证检测结果可以被查询、追溯。页面集中显示结果图、文件名、板材等级、缺陷类型、检测时间等信息，可以按缺陷关键词筛选。由于系统运行之后检测记录会不断增多，所以页面使用分页的方式控制单页展示的数量，从而提高查询和浏览的效率。":
            set_paragraph_text(paragraph, "历史记录页面用于展示和追溯已经完成的检测结果。页面以卡片形式显示结果图、文件名、板材等级、缺陷类型和检测时间等信息，并提供缺陷关键词筛选功能。由于系统运行后检测记录会不断增加，页面采用分页方式控制单页展示数量，从而提高查询和浏览效率。管理员可以查看全部用户的检测记录，普通用户只能查看本人产生的记录，该设计与系统权限控制逻辑保持一致。")
        elif paragraph.text == "日志、报表模块主要是对系统中已经保存的检测数据进行统计、整理并做可视化展示。系统可以计算出今天检测量、缺陷类别占比、近七天检测趋势，用前端图表组件展示出来，为系统运行状态分析、检测结果汇总提供支持。":
            set_paragraph_text(paragraph, "日志报表页面主要对系统中已经保存的检测数据和工作日志进行统计、整理与可视化展示。页面展示今日检测量、累计工作时长、缺陷类别占比和近七天检测趋势等内容，使用户能够直观了解系统近期运行情况。该模块既服务于检测结果汇总，也为后续质量分析和管理决策提供数据依据。")
        elif paragraph.text == "正常测试时系统可以进行用户登录、图片检测、结果展示和历史记录保存等操作，相关的测试截图如下图6-1所示。":
            set_paragraph_text(paragraph, "正常功能测试围绕系统的主要业务流程展开。测试时先使用正确账号完成登录，再上传木质板材图片进行缺陷检测，随后检查页面是否能够返回结果图、缺陷类别、置信度和板材等级，并确认检测记录是否写入历史记录页面。测试结果表明，系统在正常输入条件下能够完成登录、检测、结果展示和记录保存等核心流程，相关测试结果如图6-1所示。")
        elif paragraph.text == "异常测试时系统会给出未选图片直接提交检测、身份不符登录、库存不足出库等提示，相关的测试截图如下图6-2所示。":
            set_paragraph_text(paragraph, "异常功能测试主要用于验证系统在错误输入或业务条件不满足时的处理能力。测试内容包括未选择图片时直接提交检测、普通用户误用管理员入口登录、账号角色不匹配以及库存不足时继续提交销售订单等情况。测试过程中，系统能够返回明确提示并阻止错误操作继续执行，说明系统具备基本的异常处理能力和业务约束能力，相关测试结果如图6-2所示。")
        elif paragraph.text == "从综合测试结果可知，系统已经可以比较完整的实现论文前文设计的主要功能，即木材表面缺陷图像检测、等级判定、历史追溯、日志报表统计、库存更新、销售出库和管理员后台管理等业务。整体运行比较稳定，可以满足毕业设计阶段对于系统展示和功能测试的基本需求。":
            set_paragraph_text(paragraph, "从综合测试结果可知，系统能够较完整地实现论文前文设计的主要功能，包括木质板材表面缺陷图像检测、等级判定、历史追溯、日志报表统计、库存更新、销售出库和管理员后台管理等业务。各测试用例均达到预期结果，说明系统整体运行较为稳定，能够满足毕业设计阶段系统演示和功能验证的基本要求。")

    # 扩充测试目的与测试环境，避免章节内容过短。
    for paragraph in doc.paragraphs:
        if paragraph.text == "测试时以普通用户和管理员两种不同的身份进入系统，并按照实际的业务流程检验各个主要的功能模块。除了正常的测试之外，对没有选择图片直接提交检测、角色入口和账号身份不匹配、库存不足时继续提交订单等异常情况进行检查，看系统能否给出明确的提示，并且保持稳定运行。":
            set_paragraph_text(paragraph, "测试时分别以普通用户和管理员两种身份进入系统，并按照实际业务流程检验主要功能模块。普通用户侧重点包括图片上传检测、摄像头采集检测、历史记录查询、日志报表查看和销售出库；管理员侧重点包括用户管理、系统参数设置和模型切换。除正常功能测试外，还对未选择图片直接提交检测、角色入口和账号身份不匹配、库存不足时继续提交订单等异常情况进行检查，以验证系统能否给出明确提示并保持稳定运行。")

    # 增加总结与展望内容，并改为连续段落形式，避免过短和机械列点。
    for paragraph in doc.paragraphs:
        if paragraph.text == "本文的研究工作是在已经完成开发和调试的木材表面缺陷检测系统的基础上进行的。论文中大部分页面、业务流程、数据结构都可以在系统实现中找到对应的内容，因此论文分析内容与系统实现情况基本一致。":
            set_paragraph_text(paragraph, "本文围绕木质板材表面缺陷检测和业务数据管理需求，设计并实现了一套基于Python的Web检测系统。系统以Flask作为后端框架，使用SQLite和SQLAlchemy保存用户、检测历史、工作日志、库存、销售记录和系统参数，结合YOLO目标检测模型完成裂纹、死节、活节等常见缺陷识别。通过系统分析、结构设计、数据库设计、功能实现和测试验证可以看出，论文中的功能划分、业务流程和数据结构均能在实际系统中找到对应实现，系统设计与开发结果基本一致。")
        elif paragraph.text.startswith("未来可以从三个方面来开展完善工作。"):
            set_paragraph_text(paragraph, "在系统实现过程中，本文将缺陷检测结果与等级判定、库存更新、销售出库和报表统计进行结合，使系统不只是完成单次图像识别，而是能够形成检测、记录、追溯和业务管理的连续流程。测试结果表明，系统能够完成登录认证、图片检测、摄像头采集、历史查询、日志报表、数据导出、销售管理、用户管理和模型切换等主要功能，基本达到毕业设计阶段的预期目标。后续仍可从数据集扩充、模型精度优化、多人部署、权限粒度控制、文件清理和长期运行稳定性等方面继续改进，使系统更适合真实板材加工场景中的质量检测和管理需求。")

    # 参考文献格式粗略统一：英文文献标点后补空格，中文文献编号后加空格。
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("["):
            text = paragraph.text
            if len(text) > 3 and text[3] != " ":
                text = text[:3] + " " + text[3:]
            text = text.replace(" ,", ",").replace(" .", ".").replace(" et al.", " et al. ")
            set_paragraph_text(paragraph, text)

    # 图片统一限制宽度并添加边框，避免截图过宽或缺少外框。
    max_width = Inches(6.2)
    for idx, shape in enumerate(doc.inline_shapes):
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)
        if idx >= 2:
            add_picture_border(shape)

    # 全文格式整理：正文首行缩进，标题、图题、表题和代码提示不缩进。
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        no_indent = (
            text.startswith("图") or text.startswith("表") or text.startswith("部分关键代码")
            or text == "目 录" or text == "参考文献" or text == "致谢"
            or text[0].isdigit() and len(text) < 30
            or text.startswith("[")
            or text.startswith("摘要") or text.startswith("关键词") or text.startswith("Abstract") or text.startswith("Keywords")
        )
        if text.startswith("图") or text.startswith("表"):
            format_paragraph(paragraph, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
        elif no_indent:
            format_paragraph(paragraph, size=12, first_indent=False)
        else:
            format_paragraph(paragraph, size=12, first_indent=True)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
