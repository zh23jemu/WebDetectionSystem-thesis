from pathlib import Path
import shutil

from docx import Document


CURRENT = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的木质板材缺陷检测系统的设计与实现.docx")
BACKUP = Path(r"C:\Coding\WebDetectionSystem-thesis\WebDetectionSystem\基于Python的木质板材缺陷检测系统的设计与实现_修改前备份.docx")


def replace_text_keep_runs(paragraph, new_text):
    """只替换段落文本内容，尽量保留原段落的字体、缩进、对齐、行距等格式。"""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def set_run_texts(paragraph, texts):
    """按原有 run 数量分配文本，避免把局部加粗段落错误变成整段加粗。"""
    while len(paragraph.runs) < len(texts):
        paragraph.add_run("")
    for idx, text in enumerate(texts):
        paragraph.runs[idx].text = text
    for idx in range(len(texts), len(paragraph.runs)):
        paragraph.runs[idx].text = ""


def main():
    # 先以备份文件为模板恢复原有版式，再在其上做最小文本修改。
    shutil.copy2(BACKUP, CURRENT)
    doc = Document(CURRENT)

    replacements = {
        "基于Python的木质板材表面缺陷系统的设计与实现": "基于Python的木质板材缺陷检测系统的设计与实现",
        "Computer science and technology   Luan Yankang": "Computer Science and Technology   Luan Yankang",
        "tutor     Liu Rong": "Supervisor     Liu Rong",
        "Keywords: Python; Wood Defect Detection; YOLO; Flask; Web System; Inventory and Sales": "Keywords: Python; wood panel defect detection; YOLO; Flask; Web system; inventory and sales linkage",
        "5.3.3 系统检测与可视化分析": "5.3.2 系统检测与可视化分析",
        "5.3.3 系统检测与可视化分析\t26": "5.3.2 系统检测与可视化分析\t26",
    }

    long_replacements = {
        "就系统整体实现方式而言，本系统可以分成页面展示层、业务控制层、模型推理层和数据存储层。分层结构和项目代码组织大体上一致，既可以说明各个模块之间的协作关系，也可以使论文分析的内容和实际系统的实现保持一致。":
        "从系统使用角色和业务流程来看，本系统主要包括普通用户功能、管理员功能和公共支撑功能三部分。普通用户负责完成缺陷检测、检测历史查询、日志报表查看、销售出库、缺陷图谱浏览和个人信息维护；管理员在普通用户功能的基础上，增加用户管理、系统参数设置、模型切换和日志查看等维护功能。两个角色在统一登录认证和权限控制基础上共享检测、历史、报表和库存等业务数据。",
        "从角色职责上来说，普通用户主要是使用登录认证、木材缺陷检测、历史记录查询、日志报表查看、销售管理、帮助中心浏览、个人信息维护等功能，管理员在此基础上又增加了用户管理、系统设置、模型切换等后台管理任务。检测模块是系统运行的中心，历史、报表、销售模块一起形成了业务闭环，使中心和后台配置变得简单。":
        "系统功能结构以缺陷检测为核心展开。检测模块接收本地上传或摄像头采集的板材图片，调用模型完成识别和等级判定；历史记录模块保存检测结果并支持追溯查询；日志报表模块对检测数量、工作时长和缺陷类别进行统计；库存与销售模块根据等级结果完成库存更新和出库记录保存。管理员模块主要为上述业务提供账号、参数和模型文件维护能力。",
        "根据页面入口进行细分，系统主要有登录、仪表盘、缺陷检测、历史记录、日志报表、数据导出、销售管理、帮助中心、用户管理、系统设置等九个功能页面。各个页面同后端业务模块一一对应，页面组织关系比较清晰，可以满足日常使用、权限划分以及未来扩展的要求。":
        "按照页面入口细分，系统包含登录、仪表盘、缺陷检测、历史记录、日志报表、数据导出、销售管理、帮助中心、用户管理和系统设置等功能页面。普通用户只访问与检测和业务操作相关的页面，管理员可额外进入后台维护页面。该功能划分能够对应系统实际代码中的路由、模板和数据库操作，也为后续系统设计图和实现章节提供依据。",
        "数据库的E-R图用来表示系统中的主要实体以及它们之间的联系。本系统以用户、检测历史、工作日志、销售记录、库存等数据对象为设计对象，用户和检测历史、工作日志之间存在直接关系，库存和销售记录之间是通过板材等级形成的业务关系。主要实体属性图如下图4-6所示。":
        "数据库的E-R图用来表示系统中的主要实体、属性以及实体之间的联系。本系统根据实际数据库模型抽取出用户、检测历史、工作日志、销售记录、库存和系统设置六类实体。用户实体用于保存登录账号、密码哈希和角色信息；检测历史实体用于保存原始文件名、结果图、缺陷类型、置信度、检测时间和等级结果；工作日志实体用于记录用户登录、退出和累计工作时长；销售记录实体用于保存客户、等级、数量、金额和销售时间；库存实体用于记录不同等级板材的库存数量；系统设置实体用于保存当前模型路径和检测阈值等参数。主要实体属性如图4-6所示。",
        "根据系统业务流程可知，用户登录之后会产生检测记录和工作日志；检测结果完成等级判定之后会改变库存数量；销售出库操作要读取库存并生成销售记录。系统各个实体之间的联系如图4-7所示。":
        "在实体关系上，一个用户可以产生多条检测历史记录，也可以对应多条工作日志记录，因此用户与检测历史、用户与工作日志均为一对多关系。检测历史中的等级结果会影响库存数量，库存又会被销售出库业务读取和扣减，因此库存与销售记录之间通过板材等级形成业务关联。系统设置实体为检测流程提供模型路径和阈值参数，不直接依赖某一条检测记录，但会影响后续检测结果生成。系统各个实体之间的联系如图4-7所示。",
        "根据系统的功能需求以及E-R图的设计结果，数据库主要是用户表、检测历史表、工作日志表、销售记录表、系统设置表和库存表。各个数据表分别担负着账号管理、检测结果保存、工作时间记载、销售业务存贮、参数设定以及库存统计这些工作。具体表结构设计如上图所示。":
        "根据系统功能需求和E-R图设计结果，数据库主要包括用户表、检测历史表、工作日志表、销售记录表、系统设置表和库存表。各表分别承担账号管理、检测结果保存、工作时间记录、销售业务存储、参数设置和库存统计等任务。表结构设计既要与系统实际模型保持一致，也要说明字段含义、数据类型和业务用途。",
        "历史记录页面的主要工作就是保证检测结果可以被查询、追溯。页面集中显示结果图、文件名、板材等级、缺陷类型、检测时间等信息，可以按缺陷关键词筛选。由于系统运行之后检测记录会不断增多，所以页面使用分页的方式控制单页展示的数量，从而提高查询和浏览的效率。":
        "历史记录页面用于展示和追溯已经完成的检测结果。页面集中显示结果图、文件名、板材等级、缺陷类型和检测时间等信息，并提供缺陷关键词筛选功能。由于系统运行后检测记录会不断增加，页面采用分页方式控制单页展示数量，从而提高查询和浏览效率。管理员可以查看全部用户的检测记录，普通用户只能查看本人产生的记录，该设计与系统权限控制逻辑保持一致。",
        "日志、报表模块主要是对系统中已经保存的检测数据进行统计、整理并做可视化展示。系统可以计算出今天检测量、缺陷类别占比、近七天检测趋势，用前端图表组件展示出来，为系统运行状态分析、检测结果汇总提供支持。":
        "日志报表页面主要对系统中已经保存的检测数据和工作日志进行统计、整理与可视化展示。页面展示今日检测量、累计工作时长、缺陷类别占比和近七天检测趋势等内容，使用户能够直观了解系统近期运行情况。该模块既服务于检测结果汇总，也为后续质量分析和管理决策提供数据依据。",
        "正常测试时系统可以进行用户登录、图片检测、结果展示和历史记录保存等操作，相关的测试截图如下图6-1所示。":
        "正常功能测试围绕系统的主要业务流程展开。测试时先使用正确账号完成登录，再上传木质板材图片进行缺陷检测，随后检查页面是否能够返回结果图、缺陷类别、置信度和板材等级，并确认检测记录是否写入历史记录页面。测试结果表明，系统在正常输入条件下能够完成登录、检测、结果展示和记录保存等核心流程，相关测试结果如图6-1所示。",
        "异常测试时系统会给出未选图片直接提交检测、身份不符登录、库存不足出库等提示，相关的测试截图如下图6-2所示。":
        "异常功能测试主要用于验证系统在错误输入或业务条件不满足时的处理能力。测试内容包括未选择图片时直接提交检测、普通用户误用管理员入口登录、账号角色不匹配以及库存不足时继续提交销售订单等情况。测试过程中，系统能够返回明确提示并阻止错误操作继续执行，说明系统具备基本的异常处理能力和业务约束能力，相关测试结果如图6-2所示。",
        "本文的研究工作是在已经完成开发和调试的木材表面缺陷检测系统的基础上进行的。论文中大部分页面、业务流程、数据结构都可以在系统实现中找到对应的内容，因此论文分析内容与系统实现情况基本一致。":
        "本文围绕木质板材表面缺陷检测和业务数据管理需求，设计并实现了一套基于Python的Web检测系统。系统以Flask作为后端框架，使用SQLite和SQLAlchemy保存用户、检测历史、工作日志、库存、销售记录和系统参数，结合YOLO目标检测模型完成裂纹、死节、活节等常见缺陷识别。通过系统分析、结构设计、数据库设计、功能实现和测试验证可以看出，论文中的功能划分、业务流程和数据结构均能在实际系统中找到对应实现，系统设计与开发结果基本一致。",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text in replacements:
            replace_text_keep_runs(paragraph, replacements[text])
        elif text in long_replacements:
            replace_text_keep_runs(paragraph, long_replacements[text])
        elif text.startswith("Abstract:"):
            replace_text_keep_runs(
                paragraph,
                "Abstract: To address the low efficiency of manual wood panel surface defect inspection, the difficulty of preserving inspection records, and the weak connection between quality grading and inventory management, this thesis designs and implements a Python-based web defect detection system for wood panels. The system uses Flask for back-end service control, SQLite and SQLAlchemy for data persistence, and a YOLO detection model for identifying common surface defects such as cracks, dead knots and live knots. After users upload images or capture images through a camera, the system returns defect categories, confidence values, result images and grading results, and stores the corresponding inspection records. In addition, the system provides login authentication, history query, log statistics, report export, inventory linkage, sales outbound management, user management and model switching functions. Functional testing shows that the system can complete the main processes of defect detection, result tracing and business data management, and has practical value for teaching demonstration and graduation project implementation."
            )

    # 恢复英文题名页里被文本替换影响的 run 结构，保持和备份版式一致。
    if len(doc.paragraphs) > 66:
        set_run_texts(doc.paragraphs[62], ["Computer Science and Technology   ", "Luan Yankang"])
        set_run_texts(doc.paragraphs[63], ["Supervisor", "    ", " Liu Rong"])
        set_run_texts(
            doc.paragraphs[64],
            [
                "Abstract:",
                " To address the low efficiency of manual wood panel surface defect inspection, the difficulty of preserving inspection records, and the weak connection between quality grading and inventory management, this thesis designs and implements a Python-based web defect detection system for wood panels. The system uses Flask for back-end service control, SQLite and SQLAlchemy for data persistence, and a YOLO detection model for identifying common surface defects such as cracks, dead knots and live knots. After users upload images or capture images through a camera, the system returns defect categories, confidence values, result images and grading results, and stores the corresponding inspection records. In addition, the system provides login authentication, history query, log statistics, report export, inventory linkage, sales outbound management, user management and model switching functions. Functional testing shows that the system can complete the main processes of defect detection, result tracing and business data management, and has practical value for teaching demonstration and graduation project implementation.",
            ],
        )
        set_run_texts(
            doc.paragraphs[66],
            ["Keywords:", " Python; wood panel defect detection; YOLO; Flask; Web system; inventory and sales linkage"],
        )

    # 只修改表格文字，不改表格列数、边框、字体和行高。
    table_updates = [
        [["字段名", "类型", "说明"],
         ["id", "Integer", "主键，自增"],
         ["username", "String(50)", "用户名，唯一，非空"],
         ["password", "String(255)", "密码哈希值，非空"],
         ["role", "String(20)", "角色，user/admin"]],
        [["字段名", "类型", "说明"],
         ["id", "Integer", "主键，自增"],
         ["filename", "String(300)", "原始文件名"],
         ["result_image", "String(300)", "结果图文件名"],
         ["upload_date", "DateTime", "上传检测时间"],
         ["defect_type", "String(100)", "缺陷类型"],
         ["confidence", "Float", "最高置信度"],
         ["user_id", "Integer", "关联用户 ID"],
         ["grade", "String(10)", "板材等级"]],
        [["字段名", "类型", "说明"],
         ["id", "Integer", "主键，自增"],
         ["user_id", "Integer", "操作用户 ID"],
         ["login_time", "DateTime", "登录时间"],
         ["logout_time", "DateTime", "退出时间"]],
        [["字段名", "类型", "说明"],
         ["id", "Integer", "主键，自增"],
         ["customer_name", "String(100)", "客户名称"],
         ["product_grade", "String(10)", "销售等级"],
         ["quantity", "Integer", "销售数量"],
         ["total_price", "Float", "订单总价"],
         ["sale_date", "DateTime", "销售时间"]],
        [["字段名", "类型", "说明"],
         ["id", "Integer", "主键，自增"],
         ["key_name", "String(50)", "设置键名，唯一"],
         ["value", "String(200)", "设置值"]],
        [["字段名", "类型", "说明"],
         ["grade", "String(10)", "等级主键"],
         ["count", "Integer", "库存数量，非负整数"]],
    ]
    db_tables = doc.tables[-7:-1]
    for table, rows in zip(db_tables, table_updates):
        for r_idx, row_data in enumerate(rows):
            if r_idx >= len(table.rows):
                break
            for c_idx, value in enumerate(row_data):
                if c_idx < len(table.rows[r_idx].cells):
                    table.rows[r_idx].cells[c_idx].text = value

    doc.save(CURRENT)


if __name__ == "__main__":
    main()
